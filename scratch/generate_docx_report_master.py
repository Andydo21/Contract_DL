import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_master_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # Helper to add headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        return p
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(79, 70, 229)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_bullet(bold_prefix, text_content, indent_level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
        p.paragraph_format.space_after = Pt(3)
        r_bold = p.add_run(bold_prefix)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(30, 41, 59)
        
        r_text = p.add_run(text_content)
        r_text.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # ------------------ COVER PAGE ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(48)
    title_run = title_p.add_run("BÁO CÁO KỸ THUẬT PHÂN TÍCH TỔNG HỢP Master")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_p2.add_run("HỆ THỐNG QUẢN LÝ HỢP ĐỒNG THÔNG MINH")
    title_run2.font.name = 'Arial'
    title_run2.font.size = Pt(24)
    title_run2.font.bold = True
    title_run2.font.color.rgb = RGBColor(79, 70, 229)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Tài liệu đặc tả kỹ thuật chi tiết dạng văn bản thuần: Database Entities, Document Processor, Contract Services, AI Subsystems và Hyperledger Fabric Blockchain Flow")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run("Tác giả: ").bold = True
    p_meta.add_run("Đỗ Đăng An (dodangan)\n")
    p_meta.add_run("Phiên bản phân hệ: ").bold = True
    p_meta.add_run("v1.0.9\n")
    p_meta.add_run("Công nghệ lõi: ").bold = True
    p_meta.add_run("Django, FastAPI, PyTorch (Qwen 3B NF4), PyMuPDF + PaddleOCR, Hyperledger Fabric Gateway\n")
    p_meta.paragraph_format.line_spacing = 1.3
    
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep_run = p_sep.add_run("―" * 50)
    p_sep_run.font.color.rgb = RGBColor(226, 232, 240)
    
    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 1: TỔNG QUAN VÀ Ý NGHĨA LUỒNG CHI TIẾT
    # =========================================================================
    add_heading_1("Chương 1. Chi Tiết Ý Nghĩa và Quy Trình các Luồng Nghiệp Vụ")
    
    doc.add_paragraph(
        "Hệ thống Quản lý Hợp đồng được thiết kế nhằm đồng bộ hóa ba phân hệ chính thành một pipeline thống nhất. "
        "Dưới đây là phân tích chi tiết ý nghĩa và cách vận hành của từng luồng nghiệp vụ cốt lõi dưới dạng văn bản thuần phân cấp."
    )
    
    add_heading_2("1.1 Luồng Tải lên & Phân Tích Rủi Ro Hợp Đồng (End-to-End)")
    doc.add_paragraph(
        "Ý nghĩa thực tiễn: Giúp số hóa hoàn toàn các văn bản hợp đồng giấy hoặc tệp tin số, tự động bóc tách các điều khoản "
        "và phân tích phát hiện các rủi ro pháp lý tiềm ẩn dưới vai trò của một luật sư đại diện nghiêm khắc. Luồng này bảo vệ "
        "quyền lợi của doanh nghiệp (đặc biệt là bên mua/bên thuê) tránh các điều khoản bất lợi về bồi thường, phạt vi phạm, hoặc đơn phương chấm dứt."
    )
    doc.add_paragraph(
        "Quy trình xử lý chi tiết 10 bước:\n"
        "1. Người dùng upload tệp tin (PDF/DOCX) hoặc dán văn bản thô qua dashboard giao diện SPA.\n"
        "2. ContractService tiếp nhận, lưu trữ thông tin hợp đồng thô ở trạng thái DRAFT.\n"
        "3. encrypt_pdf() mã hóa đối xứng dữ liệu file bằng thuật toán AES-256-GCM để bảo mật an toàn dữ liệu lưu trữ vật lý trên máy chủ.\n"
        "4. DocumentService giải mã file trong RAM, đọc văn bản. Nếu là Scanned PDF hoặc hình ảnh, hệ thống chạy PaddleOCR tiếng Việt để nhận diện ký tự quang học.\n"
        "5. Nội dung text được chia tách thành các trang và lưu vào ContractContext.\n"
        "6. ClauseExtractService gọi API /api/v1/extract_clauses trên AI Service. Nếu AI Service offline, tự động kích hoạt Regex ClauseSplitter cục bộ để bóc tách điều khoản (Clause) dự phòng.\n"
        "7. Hệ thống tự động bóc tách các thực thể cơ bản (Bên A, Bên B, giá trị, hiệu lực) qua các heuristics và lưu vào ExtractedEntity.\n"
        "8. Gửi payload gồm các điều khoản và luật rủi ro (RiskRule) sang /api/v1/analyze để chạy mô hình AI Qwen2.5-3B-Instruct (lượng hóa 4-bit NF4) thực thi suy luận tìm rủi ro.\n"
        "9. Nhận kết quả JSON, regex lọc khối markdown, cập nhật điểm rủi ro tổng hợp (overall_score), và tạo các bản ghi RiskFinding.\n"
        "10. Trạng thái hợp đồng được cập nhật thành ACTIVE, sẵn sàng cho chuyên gia đánh giá và ký duyệt."
    )
    
    # Insert diagram flow E2E
    img_e2e_path = r"d:\Django_project\RiskDL\scratch\flow_e2e.png"
    if os.path.exists(img_e2e_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(img_e2e_path, width=Inches(6.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = p_cap.add_run("Hình 1.1 Sơ đồ tuần tự của luồng phân tích rủi ro đầu-cuối trong Hệ thống Quản lý Hợp đồng.")
        caption.font.italic = True
        caption.font.size = Pt(9.5)
        caption.font.color.rgb = RGBColor(100, 116, 139)
        
    add_heading_2("1.2 Luồng Đăng Ký Danh Tính Lên Blockchain (Identity Registry)")
    doc.add_paragraph(
        "Ý nghĩa thực tiễn: Đảm bảo tính xác thực pháp lý của các bên tham gia giao kết hợp đồng. Thông tin định danh của doanh nghiệp "
        "và tài khoản người dùng được lưu trữ bất biến trên sổ cái Blockchain Hyperledger Fabric, ngăn chặn hành vi giả mạo danh tính."
    )
    doc.add_paragraph(
        "Quy trình xử lý:\n"
        "1. Quản trị viên nhập thông tin Công ty (tên, mã số thuế) hoặc Người dùng (username, role) qua trang quản trị danh tính.\n"
        "2. Django tạo bản ghi tương ứng trong CSDL và gửi yêu cầu đăng ký sang blockchain-service qua API /company/register/ hoặc /user/register/.\n"
        "3. EnterpriseRegistryService gọi cổng REST Fabric Gateway (/company/store hoặc /user/store) để thực thi Chaincode lưu trữ dữ liệu.\n"
        "4. Nếu Fabric Gateway online: Trả về tx_hash, block_number và block_hash cập nhật ngược lại vào Django CSDL để lưu bằng chứng.\n"
        "5. Nếu Fabric Gateway offline: Tự động chuyển sang cơ chế Simulated Anchoring, tạo tx_hash giả lập dạng 0x... để đảm bảo tiến trình nghiệp vụ cục bộ không bị nghẽn."
    )
    
    add_heading_2("1.3 Luồng Ký Duyệt Workflow Bằng Chữ Ký Số (Digital Signature)")
    doc.add_paragraph(
        "Ý nghĩa thực tiễn: Thiết lập quy trình phê duyệt điện tử an toàn, chống từ chối (non-repudiation). Chữ ký số liên kết trực tiếp "
        "giữa danh tính người ký (thông qua chứng thư số công cộng x509) và mã băm toàn vẹn của nội dung hợp đồng."
    )
    doc.add_paragraph(
        "Quy trình xử lý:\n"
        "1. Người dùng phê duyệt một bước trong Workflow (ví dụ Legal Review hoặc Manager Approval).\n"
        "2. Hệ thống kiểm tra thời hạn và trạng thái hiệu lực của SignatureCertificate tương ứng với người dùng.\n"
        "3. Tính toán signature_hash bằng cách băm khóa cá nhân của người dùng với document_hash lưu trong HashProof của phiên bản hợp đồng.\n"
        "4. Gọi SignatureService.verify_and_sign để tạo và lưu bản ghi DigitalSignature gắn liền với HashProof và SignatureCertificate.\n"
        "5. Giao dịch ký được gửi lên ledger Blockchain để đảm bảo vết phê duyệt không thể bị xóa bỏ hoặc thay đổi."
    )

    add_heading_2("1.4 Luồng Đối Soát Toàn Vẹn Tài Liệu (Integrity Verification)")
    doc.add_paragraph(
        "Ý nghĩa thực tiễn: Phát hiện ngay lập tức mọi hành vi sửa đổi trái phép tệp tin hợp đồng lưu trên ổ cứng vật lý của máy chủ."
    )
    doc.add_paragraph(
        "Quy trình xử lý:\n"
        "1. Quản lý yêu cầu đối soát một phiên bản hợp đồng cụ thể.\n"
        "2. Hệ thống đọc file đã mã hóa trên disk, giải mã tạm thời trong RAM bằng khóa AES-256-GCM để lấy bytes văn bản gốc.\n"
        "3. Tính toán lại mã băm SHA-256 của văn bản gốc giải mã đó (current_hash).\n"
        "4. Gọi VerificationService.verify_proof gửi truy vấn lên sổ cái Blockchain qua Fabric Gateway API /verify để lấy mã băm gốc đã neo giữ.\n"
        "5. So sánh current_hash với document_hash gốc lưu trên sổ cái Blockchain. Trả về kết quả xác thực thành công (Toàn vẹn) hoặc thất bại (Hợp đồng đã bị sửa đổi trái phép)."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 2: CHI TIẾT TOÀN BỘ CÁC THỰC THỂ CSDL (ENTITIES / MODELS)
    # =========================================================================
    add_heading_1("Chương 2. Đặc Tả Chi Tiết Toàn Bộ các Thực Thể (Database Entities)")
    
    doc.add_paragraph(
        "Hệ thống Quản lý Hợp đồng thiết kế hệ thống CSDL chuẩn hóa cao, liên kết chặt chẽ giữa quản lý hợp đồng, "
        "kết quả phân tích AI và bảo chứng sổ cái Blockchain. Dưới đây là sơ đồ kiến trúc thực thể và mô tả chi tiết."
    )
    
    # Insert diagram flow Entity
    img_ent_path = r"d:\Django_project\RiskDL\scratch\flow_entity.png"
    if os.path.exists(img_ent_path):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(img_ent_path, width=Inches(6.0))
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption2 = p_cap2.add_run("Hình 2.1 Sơ đồ thực thể CSDL (ERD) và các liên kết quan hệ trong Hệ thống Quản lý Hợp đồng.")
        caption2.font.italic = True
        caption2.font.size = Pt(9.5)
        caption2.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph(
        "Dưới đây là mô tả chi tiết của từng bảng thực thể trong toàn bộ các ứng dụng hệ thống dạng văn bản thuần phân cấp:"
    )
    
    entities = [
        ("Company (contracts)", "Lưu thông tin các doanh nghiệp sử dụng hệ thống. Các trường: company_name (tên công ty), tax_code (mã số thuế), status (trạng thái ACTIVE/INACTIVE), tx_hash, block_number, block_hash (thông tin định danh neo trên Blockchain Hyperledger Fabric)."),
        ("Permission (contracts)", "Bảng định nghĩa các quyền hạn trong hệ thống. Các trường: permission_code (mã quyền, ví dụ VIEW_CONTRACT, ANALYZE_CONTRACT), description (mô tả quyền)."),
        ("Role (contracts)", "Vai trò của người dùng. Các trường: role_name (tên vai trò, ví dụ MANAGER, LAWYER, EMPLOYEE), description, permissions (liên kết nhiều-nhiều với bảng Permission)."),
        ("User (contracts)", "Thông tin tài khoản người dùng hệ thống. Các trường: username, email, company, role, status, tx_hash, block_number, block_hash (danh tính người dùng được neo giữ bất biến trên Blockchain)."),
        ("Tag (contracts)", "Nhãn phân loại hợp đồng (ví dụ: 'Thanh toán', 'Bất động sản'). Các trường: tag_name."),
        ("Contract (contracts)", "Thực thể chính quản lý thông tin hợp đồng. Các trường: company, contract_code (mã hợp đồng độc nhất), title (tiêu đề), contract_type (loại hợp đồng), start_date, end_date, contract_value (giá trị), status (DRAFT/ANALYZING/ACTIVE/APPROVED), tags (M2M)."),
        ("ContractVersion (contracts)", "Quản lý các phiên bản sửa đổi của hợp đồng. Các trường: contract (khóa ngoại đến Contract), version_number (số thứ tự phiên bản v1, v2...), file_hash (mã băm toàn vẹn), change_summary (tóm tắt các điểm thay đổi)."),
        ("ContractFile (contracts)", "Lưu trữ thông tin tệp tin đã mã hóa AES-256-GCM. Các trường: version, file_name, file_path, file_size, mime_type, uploaded_at."),
        ("ContractContext (contracts)", "Lưu trữ nội dung văn bản thô bóc tách theo từng trang để tìm kiếm ngữ cảnh nhanh. Các trường: version, context_type (thường là raw_text), source (tên file hoặc số trang), content (văn bản thô), relevance_score."),
        ("ContextEmbedding (contracts)", "Lưu trữ vector ID của văn bản phục vụ tìm kiếm RAG ngữ nghĩa. Các trường: context, vector_id, embedding_model."),
        ("RiskRule / Risk (contracts)", "Định nghĩa danh mục các luật rủi ro cần kiểm tra. Các trường: rule_code, rule_name (tên rủi ro), description (mô tả), severity (mức độ rủi ro HIGH/MEDIUM/LOW)."),
        ("AIAnalysis (contracts)", "Ghi nhận kết quả của một phiên bản phân tích rủi ro AI. Các trường: version, model_name, overall_score (điểm số rủi ro tổng hợp 0-100), risk_level (mức rủi ro tổng hợp), summary (tóm tắt kết quả)."),
        ("RiskFinding (contracts)", "Chi tiết từng lỗi rủi ro được phát hiện trong hợp đồng. Các trường: analysis, clause (điều khoản vi phạm), rule (luật rủi ro bị vi phạm), risk_score (điểm rủi ro), risk_level, explanation (giải thích chi tiết bằng tiếng Việt), recommendation (đề xuất cụ thể để giảm thiểu), disadvantaged_party (bên gặp bất lợi)."),
        ("Review (contracts)", "Ý kiến phê duyệt thủ công của chuyên gia pháp lý. Các trường: analysis, user (người duyệt), note (ghi chú), decision (quyết định phê duyệt), reviewed_at."),
        ("ContractParty (contracts)", "Các bên trực tiếp giao kết trong hợp đồng. Các trường: contract, party_name, tax_code, email, phone, party_type (ví dụ Bên mua, Bên bán)."),
        ("Notification (contracts)", "Thông báo gửi đến người dùng. Các trường: user, title, message, is_read."),
        ("AuditLog (contracts)", "Nhật ký kiểm toán ghi nhận lịch sử thao tác của người dùng. Các trường: user, contract, action, ip_address, created_at."),
        ("ContractSummary (ai_extract)", "Bản ghi tóm tắt điều hành (Executive Summary) sinh ra bởi AI cho phiên bản hợp đồng. Các trường: version, summary (nội dung tóm tắt), model_id, created_at, updated_at."),
        ("Clause (ai_extract)", "Chi tiết các điều khoản được phân tách từ văn bản hợp đồng. Các trường: version, context, clause_type (loại điều khoản), clause_title (tiêu đề điều khoản), clause_content (nội dung chi tiết)."),
        ("ExtractedEntity (ai_extract)", "Thực thể pháp lý (tên công ty, mã số thuế, ngày hiệu lực, giá trị) trích xuất được từ điều khoản. Các trường: clause, entity_type, entity_value, normalized_value, confidence_score."),
        ("SignatureCertificate (blockchain)", "Chứng thư số công cộng của người dùng trên Blockchain. Các trường: user_id, serial_number, issuer, valid_from, valid_to, status, certificate_pem, public_key, signature_algorithm, revoked, revoked_at."),
        ("HashProof (blockchain)", "Lưu trữ bằng chứng băm của hợp đồng trên Blockchain. Các trường: version_id, hash_algorithm (mặc định SHA-256), document_hash (mã băm hợp đồng), file_size, previous_hash (mã băm phiên bản liền trước), merkle_root (mã băm liên kết chuỗi), verified, verified_at."),
        ("BlockchainNetwork (blockchain)", "Cấu hình kết nối mạng Blockchain. Các trường: network_name, chain_type, rpc_endpoint, status."),
        ("SmartContract (blockchain)", "Thông tin Smart Contract/Chaincode triển khai trên mạng Blockchain. Các trường: network, contract_address, contract_name, version, deployed_at."),
        ("BlockchainTransaction (blockchain)", "Bản ghi chi tiết các giao dịch tương tác trực tiếp lên sổ cái Blockchain. Các trường: proof, network, smart_contract, tx_hash, block_hash, block_number, gas_fee, status, created_at, tx_type, sender, endorser, channel_name, chaincode_name, fabric_tx_id, confirmation_time, latency, retry_count."),
        ("BlockchainAudit (blockchain)", "Nhật ký kiểm toán ghi nhận lịch sử thay đổi trạng thái trước/sau của các tài nguyên Blockchain. Các trường: transaction, action, resource, before_state, after_state, status, created_at."),
        ("DigitalSignature (blockchain)", "Mã chữ ký số phê duyệt hợp đồng. Các trường: certificate, hashproof, signature (giá trị chữ ký số), algorithm, created_at, verified, verified_at."),
        ("VerificationHistory (blockchain)", "Lịch sử các lần thực thi đối soát tính toàn vẹn của hợp đồng. Các trường: version_id, verify_time, verify_result (thành công/thất bại), reason (nguyên nhân lỗi nếu có), user_id."),
        ("SmartCabinet (contracts)", "Thông tin thiết lập tủ lưu trữ hồ sơ thông minh. Các trường: cabinet_code (mã tủ độc nhất), location (vị trí đặt tủ), status (trạng thái khóa cơ học: LOCKED, UNLOCKED)."),
        ("CabinetAccessLog (contracts)", "Nhật ký kiểm toán truy cập mở khóa tủ thông minh. Các trường: cabinet (khóa ngoại SmartCabinet), user (người mở khóa), opened_at (thời điểm mở), closed_at (thời điểm đóng), status (trạng thái mở tủ: SUCCESS, DENIED, TIMEOUT)."),
        ("CabinetDocumentEvent (contracts)", "Sự kiện cho vào/rút ra các tập hồ sơ vật lý trong phiên mở tủ để chống đánh tráo. Các trường: access_log (khóa ngoại CabinetAccessLog), contract_version (khóa ngoại ContractVersion), event_type (CHECK_IN/CHECK_OUT), rfid_tag_uid (mã thẻ RFID gắn trên bìa hồ sơ để định danh), timestamp (thời điểm giao dịch), is_valid (cờ kiểm tra chéo Blockchain chống đánh tráo tài liệu: True/False).")
    ]
    
    for ename, edesc in entities:
        add_bullet("Thực thể " + ename + ": ", edesc, indent_level=0)

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 3: PHÂN HỆ XỬ LÝ TÀI LIỆU CHUYÊN SÂU (DOCUMENT_PROCESSOR)
    # =========================================================================
    add_heading_1("Chương 3. Đặc Tả Kỹ Thuật Phân Hệ Xử Lý Tài Liệu (document_processor)")
    
    doc.add_paragraph(
        "document_processor là phân hệ độc lập đảm nhận nhiệm vụ đọc, trích xuất text, OCR tiếng Việt và cắt lát điều khoản. "
        "Dưới đây là đặc tả chi tiết của các lớp dịch vụ chính và chức năng đi kèm:"
    )
    
    add_heading_2("3.1 Lớp điều phối chính DocumentService (document_processor/services/document_service.py)")
    
    funcs_ds = [
        ("process(file_path: str, split_clauses: bool = False) -> DocumentOutput",
         "Hàm điều phối toàn bộ quy trình xử lý văn bản.\n"
         "  - Logic xử lý:\n"
         "    * Bước 1: Gọi FileLoader.load(file_path) để kiểm thử định dạng và kích thước tệp tin.\n"
         "    * Bước 2: Gọi phương thức trích xuất tương ứng theo định dạng tệp (DOCX, PDF, image, txt).\n"
         "    * Bước 3: Chuẩn hóa văn bản thông qua TextNormalizer.\n"
         "    * Bước 4: Nếu split_clauses = True, gọi ClauseSplitter.split() để chia tách các điều khoản.\n"
         "    * Bước 5: Kiểm định sinh warnings bằng bộ Validator nội bộ.\n"
         "    * Bước 6: Đóng gói kết quả trả về dưới dạng mô hình DocumentOutput.\n"
         "  - Ví dụ giá trị trả về: Một thực thể DocumentOutput chứa document_info (file_name='hd.pdf', ocr_used=True), danh sách pages, clauses và cảnh báo warnings."),
        
        ("_extract(load_result: LoadResult) -> Tuple[List[PageData], bool]",
         "Xác định phương pháp trích xuất văn bản thích hợp.\n"
         "  - Logic xử lý: Phân nhánh dựa vào trường load_result.file_type. Trả về danh sách PageData và cờ hiệu ocr_used.\n"
         "  - Ví dụ giá trị trả về: ([PageData(page_number=1, text='...', source='pdf')], False)."),
         
        ("_extract_txt(file_path: str) -> Tuple[List[PageData], bool]",
         "Trích xuất văn bản từ file text thô.\n"
         "  - Logic xử lý: Đọc nhị phân, giải mã bằng UTF-8 (bỏ qua ký tự lỗi).\n"
         "  - Ví dụ giá trị trả về: ([PageData(page_number=1, text='...', source='txt')], False)."),
         
        ("_extract_docx(file_path: str) -> Tuple[List[PageData], bool]",
         "Trích xuất văn bản từ file Word (.docx).\n"
         "  - Logic xử lý: Sử dụng thư viện python-docx thông qua DocxExtractor để gom toàn bộ đoạn văn (paragraphs).\n"
         "  - Ví dụ giá trị trả về: ([PageData(page_number=1, text='Điều khoản bảo mật...', source='docx')], False)."),
         
        ("_extract_pdf(file_path: str) -> Tuple[List[PageData], bool]",
         "Trích xuất văn bản từ file PDF, tự động nhận diện trang scan để kích hoạt OCR.\n"
         "  - Logic xử lý: Duyệt qua từng trang PDF bằng PyMuPDF. Nếu số lượng ký tự có nghĩa trên trang < 50 ký tự, đưa trang đó vào danh sách chạy OCR.\n"
         "  - Ví dụ giá trị trả về: ([PageData(page_number=1, text='...', source='pdf')], True)."),
         
        ("_extract_image(file_path: str) -> Tuple[List[PageData], bool]",
         "Trích xuất văn bản từ file hình ảnh (PNG, JPG).\n"
         "  - Logic xử lý: Gọi OCREngine để quét chữ viết quang học trên tệp ảnh.\n"
         "  - Ví dụ giá trị trả về: ([PageData(page_number=1, text='Cộng hòa xã hội...', source='ocr')], True)."),
         
        ("_ocr_pdf_pages(extractor: PdfExtractor, file_path: str, page_numbers: List[int]) -> List[PageData]",
         "Thực hiện OCR riêng lẻ cho các trang PDF được xác định là trang quét (Scanned Page).\n"
         "  - Logic xử lý: Render trang PDF thành định dạng ảnh tạm thời trong RAM, truyền bytes dữ liệu vào PaddleOCR để nhận diện chữ tiếng Việt.\n"
         "  - Ví dụ giá trị trả về: Danh sách các đối tượng PageData với trường source='ocr'."),
         
        ("_validate(pages: List[PageData], clauses: List[ClauseData], load_result: LoadResult, split_clauses: bool = False) -> List[str]",
         "Kiểm tra tính hợp lệ kỹ thuật của tài liệu sau xử lý.\n"
         "  - Logic xử lý: Tích lũy các lỗi cảnh báo như: trang rỗng, độ tự tin nhận diện chữ OCR thấp (< 60%), không tìm thấy tiêu đề điều khoản nào trong văn bản.\n"
         "  - Ví dụ giá trị trả về: ['Trang rỗng sau normalize: [3]', 'OCR confidence thấp (<0.6) tại trang: [1]'].")
    ]
    
    for fname, fdesc in funcs_ds:
        add_bullet(fname + ":\n", fdesc, indent_level=0)
        
    add_heading_2("3.2 Lớp chia tách điều khoản ClauseSplitter (document_processor/splitter/clause_splitter.py)")
    
    funcs_cs_split = [
        ("split(pages: List[PageData]) -> List[ClauseData]",
         "Phân tách toàn bộ văn bản của hợp đồng thành các điều khoản pháp lý riêng biệt.\n"
         "  - Logic xử lý:\n"
         "    * Bước 1: Ghép văn bản của tất cả các trang kèm theo marker phân tách trang đặc biệt.\n"
         "    * Bước 2: Quét biểu thức chính quy (Regex) tiếng Anh và tiếng Việt để tìm các tiêu đề điều khoản (ví dụ: 'Điều 1.', 'Article II').\n"
         "    * Bước 3: Cắt lát nội dung văn bản giữa hai tiêu đề điều khoản liên tiếp.\n"
         "    * Bước 4: Dựa vào marker phân tách trang để suy ngược số trang bắt đầu (start_page) và số trang kết thúc (end_page) của từng điều khoản.\n"
         "  - Ví dụ giá trị trả về: Danh sách các đối tượng ClauseData (id='Điều 1', title='Điều 1. Đối tượng hợp đồng', content='Bên A cam kết bán cho bên B...', start_page=1, end_page=2)."),
         
        ("_merge_pages(pages: List[PageData]) -> Tuple[str, List[Tuple[int, int]]]",
         "Ghép các trang văn bản lại với nhau thành một chuỗi duy nhất để xử lý chia tách xuyên trang.\n"
         "  - Logic xử lý: Chèn ký tự marker đặc biệt có dạng '\\x00PAGE:{page_num}\\x00' vào ranh giới giữa các trang.\n"
         "  - Ví dụ giá trị trả về: ('\\x00PAGE:1\\x00\\nNội dung trang 1...\\n\\x00PAGE:2\\x00\\nNội dung trang 2...', [(1, 0), (2, 452)])."),
         
        ("_extract_raw_clauses(merged_text: str) -> List[_RawClause]",
         "Sử dụng Regex khớp tiêu đề điều khoản để trích xuất danh sách điều khoản thô.\n"
         "  - Logic xử lý: Quét Regex biểu thức chính quy kết hợp mẫu tiếng Việt 'Điều <số>' và tiếng Anh 'Article/Section <số>', loại bỏ các ký tự marker trang phân tách khỏi nội dung thô.\n"
         "  - Ví dụ giá trị trả về: Danh sách các đối tượng _RawClause chứa title và nội dung body thô."),
         
        ("_assign_pages(raw_clauses: List[_RawClause], page_boundaries: List[Tuple[int, int]], pages: List[PageData]) -> List[ClauseData]",
         "Tính toán chính xác số trang bắt đầu và kết thúc của điều khoản.\n"
         "  - Logic xử lý: Tra cứu vị trí index ký tự bắt đầu của điều khoản thô trong chuỗi văn bản ghép, đối soát với mảng page_boundaries để xác định số trang tương ứng.\n"
         "  - Ví dụ giá trị trả về: Danh sách ClauseData hoàn chỉnh.")
    ]
    
    for fname, fdesc in funcs_cs_split:
        add_bullet(fname + ":\n", fdesc, indent_level=0)
        
    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 4: ĐẶC TẢ CHI TIẾT TOÀN BỘ CÁC HÀM DỊCH VỤ DJANGO & BLOCKCHAIN (SERVICES)
    # =========================================================================
    add_heading_1("Chương 4. Đặc Tả Chi Tiết Lớp Dịch Vụ Django và Phân Hệ Blockchain")
    
    doc.add_paragraph(
        "Dưới đây là phần mô tả đặc tả chi tiết của từng hàm dịch vụ trong Django Backend và Blockchain Fabric Service. "
        "Mỗi hàm được nêu rõ thuộc Service Class nào, tác động trực tiếp đến bảng dữ liệu nào và các trường bị ảnh hưởng kèm theo ví dụ cụ thể."
    )
    
    add_heading_2("4.1 Lớp dịch vụ ContractService (contracts/services.py)")
    
    funcs_cs_django = [
        ("list_all_contracts(company=None) -> List[Dict]",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Đọc bảng Contract và AIAnalysis.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Không ghi đè hay cập nhật trường nào (Hàm chỉ đọc dữ liệu).\n"
         "  - Ví dụ giá trị đọc ra: Contract.contract_code = 'HD-2026-001', Contract.title = 'Hợp đồng mua bán thiết bị CNTT', AIAnalysis.overall_score = 80.00."),
        
        ("get_contract_details(contract_id, version_id=None) -> Dict",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Đọc Contract, ContractVersion, ContractFile, Clause, ExtractedEntity, Review, AIAnalysis, ContractContext và ContractSummary.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Nếu hợp đồng chưa có phiên bản nào, tự động tạo mới bản ghi phiên bản mặc định:\n"
         "    * ContractVersion.version_number (Bảng ContractVersion): Lưu số phiên bản. Ví dụ: 1.\n"
         "    * ContractVersion.change_summary (Bảng ContractVersion): Lưu tóm tắt. Ví dụ: 'Khởi tạo phiên bản mặc định'.\n"
         "    * ContractVersion.contract_id (Bảng ContractVersion): Ví dụ: 12."),
        
        ("create_and_analyze_contract(code, title, contract_type, start_date, end_date, contract_value, file_obj=None, raw_content=None, company=None) -> Contract",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi trong Contract, ContractVersion, và ContractFile.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng Contract:\n"
         "    * Contract.contract_code: Lưu mã hợp đồng. Ví dụ: 'HD-MB-2026-009'.\n"
         "    * Contract.title: Lưu tiêu đề. Ví dụ: 'Hợp đồng mua bán hệ thống máy chủ'.\n"
         "    * Contract.contract_type: Lưu loại hợp đồng. Ví dụ: 'Mua bán'.\n"
         "    * Contract.start_date: Lưu ngày bắt đầu. Ví dụ: '2026-07-20'.\n"
         "    * Contract.end_date: Lưu ngày hết hạn. Ví dụ: '2027-07-20'.\n"
         "    * Contract.contract_value: Lưu trị giá. Ví dụ: 1200000000.00 (1.2 tỷ VNĐ).\n"
         "    * Contract.status: Lưu trạng thái ban đầu. Ví dụ: 'DRAFT'.\n"
         "  - Bảng ContractVersion:\n"
         "    * ContractVersion.version_number: Lưu phiên bản. Ví dụ: 1.\n"
         "    * ContractVersion.change_summary: Lưu lý do thay đổi. Ví dụ: 'Initial version'.\n"
         "  - Bảng ContractFile:\n"
         "    * ContractFile.file_name: Lưu tên tệp. Ví dụ: 'hop_dong_server.pdf'.\n"
         "    * ContractFile.file_path: Lưu đường dẫn mã hóa AES-256-GCM. Ví dụ: '/media/contracts/hop_dong_server_enc.pdf'.\n"
         "    * ContractFile.file_size: Lưu dung lượng tệp. Ví dụ: 541024 (bytes)."),
        
        ("create_new_version(contract_id, file_obj=None, raw_content=None, change_summary='') -> ContractVersion",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi trong ContractVersion và ContractFile.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng ContractVersion:\n"
         "    * ContractVersion.version_number: Lưu số phiên bản tăng dần. Ví dụ: 2.\n"
         "    * ContractVersion.change_summary: Lý do thay đổi. Ví dụ: 'Sửa điều khoản phạt chậm trả theo Biên bản thương thảo'.\n"
         "  - Bảng ContractFile:\n"
         "    * ContractFile.file_name: Tên tệp phiên bản mới. Ví dụ: 'hop_dong_server_v2.pdf'.\n"
         "    * ContractFile.file_path: Đường dẫn mã hóa mới. Ví dụ: '/media/contracts/hop_dong_server_v2_enc.pdf'."),
        
        ("analyze_contract(contract_id, version_id=None) -> Contract",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Cập nhật trạng thái bảng Contract, thực hiện lệnh DELETE dọn dẹp các bản ghi cũ của Clause, AIAnalysis và RiskFinding liên quan.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng Contract:\n"
         "    * Contract.status: Cập nhật trạng thái xử lý phân tích. Ví dụ: 'ANALYZING'."),
        
        ("manual_extract_contract(contract_id, version_id=None) -> Contract",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Cập nhật Contract, xóa Clause cũ và tạo mới các bản ghi Clause và ExtractedEntity.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng Contract:\n"
         "    * Contract.status: Trạng thái hợp đồng quay lại DRAFT. Ví dụ: 'DRAFT'.\n"
         "  - Bảng Clause:\n"
         "    * Clause.clause_title: Tiêu đề điều khoản chia tách bằng Regex. Ví dụ: 'Điều 4. Phạt vi phạm hợp đồng'.\n"
         "    * Clause.clause_content: Nội dung điều khoản. Ví dụ: 'Nếu Bên B chậm giao hàng thì chịu phạt 0.5% giá trị hợp đồng vi phạm cho mỗi tuần chậm trễ...'.\n"
         "  - Bảng ExtractedEntity:\n"
         "    * ExtractedEntity.entity_type: Lưu loại thực thể. Ví dụ: 'ACTION'.\n"
         "    * ExtractedEntity.entity_value: Giá trị thô trích xuất. Ví dụ: 'chậm giao hàng'.\n"
         "    * ExtractedEntity.normalized_value: Giá trị chuẩn hóa. Ví dụ: 'LATE_DELIVERY'.\n"
         "    * ExtractedEntity.confidence_score: Độ tin cậy heuristics. Ví dụ: 0.90."),
        
        ("submit_expert_review(analysis_id, comment, final_risk_level) -> Review",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi Review, cập nhật trạng thái Contract, tạo mới AuditLog.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng Review:\n"
         "    * Review.note: Ghi chú nhận xét của chuyên gia. Ví dụ: 'Đã giảm mức phạt vi phạm xuống 8%, điều khoản chấp nhận được'.\n"
         "    * Review.decision: Mức độ rủi ro quyết định bởi chuyên gia. Ví dụ: 'LOW'.\n"
         "  - Bảng Contract:\n"
         "    * Contract.status: Cập nhật trạng thái duyệt. Ví dụ: 'APPROVED'.\n"
         "  - Bảng AuditLog:\n"
         "    * AuditLog.action: Lưu vết kiểm toán. Ví dụ: 'REVIEW_SUBMITTED'."),
        
        ("push_to_workflow(contract_id, version_id=None) -> Dict",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Cập nhật Contract, tạo mới AuditLog.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng Contract:\n"
         "    * Contract.status: Trạng thái chờ workflow duyệt. Ví dụ: 'PENDING_WORKFLOW'.\n"
         "  - Bảng AuditLog:\n"
         "    * AuditLog.action: Hành động kiểm toán. Ví dụ: 'PUSHED_TO_WORKFLOW'."),
        
        ("get_workflow_status(contract_id, version_id=None) -> Dict/None",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Không lưu (chỉ thực hiện gọi API để đọc tiến trình workflow).\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị: Không có."),
        
        ("extract_and_save_clauses_via_processor(version, force_rule_based=False) -> None",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Xóa và tạo mới bản ghi trong ContractContext và Clause.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng ContractContext:\n"
         "    * ContractContext.context_type: Loại ngữ cảnh. Ví dụ: 'raw_text'.\n"
         "    * ContractContext.content: Đoạn văn bản bóc tách thô. Ví dụ: 'ĐIỀU 1: ĐỐI TƯỢNG VÀ NỘI DUNG HỢP ĐỒNG...'.\n"
         "    * ContractContext.relevance_score: Độ tin cậy trích xuất. Ví dụ: 0.98.\n"
         "  - Bảng Clause:\n"
         "    * Clause.clause_title: Ví dụ: 'Điều 1. Phạm vi công việc'.\n"
         "    * Clause.clause_content: Ví dụ: 'Bên B cam kết cung cấp dịch vụ phát triển phần mềm theo mô tả kỹ thuật...'."),
        
        ("_run_ai_analysis_via_api(contract, version) -> None",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Tạo mới AIAnalysis, RiskFinding, ExtractedEntity, và cập nhật Contract.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng AIAnalysis:\n"
         "    * AIAnalysis.model_name: Tên mô hình suy luận. Ví dụ: 'Qwen2.5-3B-Instruct (Fine-tuned)'.\n"
         "    * AIAnalysis.overall_score: Điểm số rủi ro tổng hợp. Ví dụ: 68.50.\n"
         "    * AIAnalysis.risk_level: Mức độ rủi ro. Ví dụ: 'MEDIUM'.\n"
         "    * AIAnalysis.summary: Đoạn tóm tắt. Ví dụ: 'Hợp đồng có rủi ro trung bình do quy định bồi thường thiệt hại vô hạn với bên mua'.\n"
         "  - Bảng RiskFinding:\n"
         "    * RiskFinding.risk_level: Ví dụ: 'HIGH'.\n"
         "    * RiskFinding.explanation: Ví dụ: 'Khoản 2 Điều 6 yêu cầu Bên B bồi thường toàn bộ thiệt hại gián tiếp, không giới hạn giá trị hợp đồng, gây rủi ro tài chính lớn cho doanh nghiệp'.\n"
         "    * RiskFinding.recommendation: Ví dụ: 'Bổ sung giới hạn bồi thường thiệt hại tối đa bằng 100% giá trị hợp đồng thực tế'.\n"
         "    * RiskFinding.disadvantaged_party: Bên gặp bất lợi. Ví dụ: 'Bên B'.\n"
         "  - Bảng Contract:\n"
         "    * Contract.status: Cập nhật trạng thái sau phân tích. Ví dụ: 'ANALYZED'."),
        
        ("_extract_and_save_basic_entities(clause) -> None",
         "Service Class: ContractService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi ExtractedEntity.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng ExtractedEntity:\n"
         "    * ExtractedEntity.entity_type: Loại thực thể bóc tách. Ví dụ: 'PARTY'.\n"
         "    * ExtractedEntity.entity_value: Tên thực thể thô. Ví dụ: 'Công ty Cổ phần DevCore'.\n"
         "    * ExtractedEntity.normalized_value: Giá trị chuẩn hóa. Ví dụ: 'CÔNG TY CỔ PHẦN DEVCORE'.\n"
         "    * ExtractedEntity.confidence_score: Độ tin cậy. Ví dụ: 0.95.")
    ]
    
    for fname, fdesc in funcs_cs_django:
        add_bullet(fname + ":\n", fdesc, indent_level=0)

    add_heading_2("4.2 Lớp dịch vụ RiskService và AnalysisHistoryService (contracts/services.py)")
    
    funcs_risk = [
        ("RiskService.list_all_risks() -> List[Dict]",
         "Service Class: RiskService\n"
         "Bảng CSDL bị tác động: Đọc bảng RiskRule và RiskFinding.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Không lưu (chỉ Read).\n"
         "  - Ví dụ giá trị đọc ra: RiskRule.rule_name = 'Phạt vi phạm vượt trần', RiskRule.severity = 'HIGH'."),
        
        ("RiskService.create_new_risk(name, description, severity_level) -> RiskRule",
         "Service Class: RiskService\n"
         "Bảng CSDL bị tác động: Tạo mới RiskRule, tạo mới AuditLog.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng RiskRule:\n"
         "    * RiskRule.rule_name: Lưu tên luật rủi ro. Ví dụ: 'Điều khoản bảo mật một chiều'.\n"
         "    * RiskRule.description: Mô tả luật rủi ro. Ví dụ: 'Chỉ yêu cầu một bên bảo mật thông tin mà không có điều khoản bảo mật song phương'.\n"
         "    * RiskRule.severity: Lưu mức độ nghiêm trọng. Ví dụ: 'MEDIUM'.\n"
         "  - Bảng AuditLog:\n"
         "    * AuditLog.action: Lưu vết kiểm toán. Ví dụ: 'RISK_CREATED'."),
        
        ("AnalysisHistoryService.list_all_analyses(company=None) -> List[Dict]",
         "Service Class: AnalysisHistoryService\n"
         "Bảng CSDL bị tác động: Đọc bảng AIAnalysis, Contract, và RiskFinding.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Không lưu (chỉ Read).")
    ]
    
    for fname, fdesc in funcs_risk:
        add_bullet(fname + ":\n", fdesc, indent_level=0)

    add_heading_2("4.3 Lớp dịch vụ phân hệ AI (ai_extract/services.py)")
    
    funcs_ai_django = [
        ("SummarizeService.summarize_version(version) -> Dict",
         "Service Class: SummarizeService\n"
         "Bảng CSDL bị tác động: Tạo mới hoặc cập nhật bản ghi ContractSummary.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng ContractSummary:\n"
         "    * ContractSummary.summary: Lưu đoạn tóm tắt tiếng Việt của hợp đồng do LLM sinh ra. Ví dụ: 'Hợp đồng mua bán máy chủ Dell PowerEdge ký ngày 20/07/2026 giữa Công ty TechVibe (Bên B) và Công ty DevCore (Bên A), tổng trị giá 1.2 tỷ VNĐ, thời gian thực hiện 30 ngày...'.\n"
         "    * ContractSummary.model_id: Lưu định danh mô hình. Ví dụ: 'phamthanhfd/contract-analysis-qwen2.5-3b'."),
        
        ("ExtractEntityService.extract_version(version, re_extract=False) -> List[Dict]",
         "Service Class: ExtractEntityService\n"
         "Bảng CSDL bị tác động: Xóa các thực thể cũ và tạo mới các bản ghi ExtractedEntity.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng ExtractedEntity:\n"
         "    * ExtractedEntity.entity_type: Ví dụ: 'TAX_CODE'.\n"
         "    * ExtractedEntity.entity_value: Lưu mã số thuế bóc tách. Ví dụ: '0101234567'.\n"
         "    * ExtractedEntity.normalized_value: Ví dụ: '0101234567'.\n"
         "    * ExtractedEntity.confidence_score: Ví dụ: 0.98."),
        
        ("ClauseExtractService.extract_version(version, re_extract=False, force_rule_based=False) -> Dict",
         "Service Class: ClauseExtractService\n"
         "Bảng CSDL bị tác động: Xóa và tạo mới các bản ghi trong ContractContext và Clause.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng ContractContext:\n"
         "    * ContractContext.content: Lưu văn bản thô bóc tách từ file PDF sau giải mã (hoặc OCR). Ví dụ: 'ĐIỀU 1. GIÁ CẢ VÀ PHƯƠNG THỨC THANH TOÁN...'.\n"
         "  - Bảng Clause:\n"
         "    * Clause.clause_title: Ví dụ: 'Điều 2. Thời hạn giao nhận hàng'.\n"
          "    * Clause.clause_content: Ví dụ: 'Bên B bàn giao toàn bộ máy chủ và giấy tờ hướng dẫn sử dụng cho Bên A trước ngày 20/08/2026...'.\n"
          "    * Clause.clause_type: Ví dụ: 'DELIVERY_TERM'.")
    ]
    
    for fname, fdesc in funcs_ai_django:
        add_bullet(fname + ":\n", fdesc, indent_level=0)

    add_heading_2("4.4 Lớp dịch vụ phân hệ Blockchain (blockchain_service/blockchain/services.py)")
    
    funcs_bc = [
        ("CertificateService.register_certificate(user_id, serial_number, issuer, valid_days=365) -> SignatureCertificate",
         "Service Class: CertificateService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi SignatureCertificate.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng SignatureCertificate:\n"
         "    * SignatureCertificate.user_id: Khóa ngoại liên kết người dùng. Ví dụ: 5.\n"
         "    * SignatureCertificate.serial_number: Mã chứng thư số duy nhất. Ví dụ: 'CA-2026-991823'.\n"
         "    * SignatureCertificate.issuer: Đơn vị cấp phát chứng thư số. Ví dụ: 'Root CA của Hệ thống'.\n"
         "    * SignatureCertificate.valid_from: Ngày hiệu lực. Ví dụ: '2026-07-19 15:00:00'.\n"
         "    * SignatureCertificate.valid_to: Ngày hết hạn. Ví dụ: '2027-07-19 15:00:00'.\n"
         "    * SignatureCertificate.status: Trạng thái chứng thư. Ví dụ: 'ACTIVE'."),
        
        ("ProofService.generate_proof(version_id, content=None, contract_code='CODE', version_number=1) -> HashProof",
         "Service Class: ProofService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi HashProof.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng HashProof:\n"
         "    * HashProof.version_id: Ví dụ: 2.\n"
         "    * HashProof.document_hash: Mã SHA-256 của file văn bản gốc. Ví dụ: '8f3c713b19c2...'.\n"
         "    * HashProof.previous_hash: Mã băm phiên bản liền trước. Ví dụ: '4e9a112c3b88...'.\n"
         "    * HashProof.merkle_root: Gốc Merkle liên kết lịch sử. Ví dụ: '6d7b8f9e0a1b...'."),
        
        ("BlockchainAnchorService.anchor_proof(proof_id, network_id=1, smart_contract_id=1) -> BlockchainTransaction",
         "Service Class: BlockchainAnchorService\n"
         "Bảng CSDL bị tác động: Cập nhật HashProof, tạo mới BlockchainTransaction và BlockchainAudit.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng HashProof:\n"
         "    * HashProof.verified: Đánh dấu trạng thái đối soát. Ví dụ: True.\n"
         "    * HashProof.verified_at: Ngày đối soát. Ví dụ: '2026-07-19 15:05:00'.\n"
         "  - Bảng BlockchainTransaction:\n"
         "    * BlockchainTransaction.tx_hash: Mã băm giao dịch neo giữ. Ví dụ: '0x9b3f8a7e0c2d...'.\n"
         "    * BlockchainTransaction.block_number: Số khối. Ví dụ: 12058.\n"
         "    * BlockchainTransaction.block_hash: Băm khối. Ví dụ: '0xabc123...'.\n"
         "    * BlockchainTransaction.status: Ví dụ: 'CONFIRMED'.\n"
         "  - Bảng BlockchainAudit:\n"
         "    * BlockchainAudit.action: Hành động trên blockchain. Ví dụ: 'Hash Anchored'."),
        
        ("VerificationService.verify_proof(version_id, content=None) -> Dict",
         "Service Class: VerificationService\n"
         "Bảng CSDL bị tác động: Đọc HashProof, tạo mới bản ghi VerificationHistory.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng VerificationHistory:\n"
         "    * VerificationHistory.version_id: Ví dụ: 2.\n"
         "    * VerificationHistory.verify_result: Kết quả xác minh toàn vẹn. Ví dụ: True.\n"
         "    * VerificationHistory.reason: Mô tả chi tiết. Ví dụ: 'Mã băm tập tin 8f3c713b19c2... trùng khớp hoàn toàn với mã băm lưu giữ trên sổ cái Hyperledger Fabric khối 12058.'."),
        
        ("SignatureService.verify_and_sign(step_id, user_id, certificate_id, signature_hash) -> Dict",
         "Service Class: SignatureService\n"
         "Bảng CSDL bị tác động: Tạo mới bản ghi DigitalSignature.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng DigitalSignature:\n"
         "    * DigitalSignature.certificate_id: Chứng thư sử dụng. Ví dụ: 3.\n"
         "    * DigitalSignature.hashproof_id: Ví dụ: 12.\n"
         "    * DigitalSignature.signature: Chữ ký số mã hóa. Ví dụ: 'MIIBIjANBgkqhkiG9w0BAQEFAAOCAQ8AMIIBCgKCAQEA...' (chuỗi PEM chữ ký).\n"
         "    * DigitalSignature.verified: Trạng thái xác thực chữ ký. Ví dụ: True."),
        
        ("EnterpriseRegistryService.register_company(company_id, company_name, tax_code) -> BlockchainTransaction",
         "Service Class: EnterpriseRegistryService\n"
         "Bảng CSDL bị tác động: Tạo mới BlockchainTransaction và BlockchainAudit.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng BlockchainTransaction:\n"
         "    * BlockchainTransaction.tx_hash: Lưu băm giao dịch ghi định danh công ty. Ví dụ: '0x1c8b9d...'.\n"
         "    * BlockchainTransaction.status: Ví dụ: 'CONFIRMED'."),
        
        ("EnterpriseRegistryService.register_user(user_id, username, company_id, role) -> BlockchainTransaction",
         "Service Class: EnterpriseRegistryService\n"
         "Bảng CSDL bị tác động: Tạo mới BlockchainTransaction và BlockchainAudit.\n"
         "Các trường dữ liệu bị tác động & Ví dụ giá trị:\n"
         "  - Bảng BlockchainTransaction:\n"
         "    * BlockchainTransaction.tx_hash: Lưu băm giao dịch định danh người dùng. Ví dụ: '0x2d9e0f...'.\n"
         "    * BlockchainTransaction.status: Ví dụ: 'CONFIRMED'.")
    ]
    
    for fname, fdesc in funcs_bc:
        add_bullet(fname + ":\n", fdesc, indent_level=0)

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 5: CÁC KỸ THUẬT AI ĐÃ TRIỂN KHAI TRONG DỰ ÁN
    # =========================================================================
    add_heading_1("Chương 5. Các Kỹ Thuật Trí Tuệ Nhân Tạo (AI) Đã Triển Khai")
    
    doc.add_paragraph(
        "Nhằm xây dựng một hệ thống phân tích điều khoản hợp đồng thông minh, tự động và đạt độ chính xác pháp lý cao, "
        "dự án Quản lý Hợp đồng đã nghiên cứu và ứng dụng thành công chuỗi kỹ thuật Trí tuệ nhân tạo (AI) hiện đại. "
        "Dưới đây là chi tiết các kỹ thuật AI đã được triển khai thực tế trong hệ thống:"
    )

    ai_techniques = [
        ("Prompt Engineering (Kỹ nghệ Gợi ý):", " Định nghĩa System Prompt rất chi tiết trong `ai_service/main.py` để ép mô hình đóng vai trò là một Luật sư đại diện nghiêm khắc, bảo vệ quyền lợi Bên mua/Bên thuê. Quy định định dạng JSON đầu ra bắt buộc và thiết lập thang điểm rủi ro rõ ràng."),
        ("Knowledge Base (Kho tri thức cục bộ):", " Cơ sở dữ liệu hợp đồng tải lên hệ thống (lưu trữ văn bản gốc trong DB và file hệ thống đã mã hóa AES-256-GCM) đóng vai trò là kho tri thức cục bộ để đối chiếu, phân tích và trích xuất."),
        ("Chunking (Cắt lát văn bản):", " Hợp đồng dài được phân tách tự động (hoặc thủ công bằng biểu thức chính quy Regex ClauseSplitter) thành từng điều khoản độc lập (`Clause`) trước khi lưu trữ và đưa vào xử lý vector, giúp tránh tràn cửa sổ ngữ cảnh của mô hình."),
        ("Embedding (Nhúng vector ngữ nghĩa):", " Sử dụng mô hình SentenceTransformers (như `Multilingual-E5-Base`) để chuyển các đoạn điều khoản hợp đồng thành vector nhúng dense 768 chiều, làm nền tảng cho việc so khớp ngữ nghĩa pháp lý."),
        ("Vector Database (Cơ sở dữ liệu Vector):", " Sử dụng ChromaDB để lưu trữ các vector nhúng của điều khoản hợp đồng và thực hiện truy vấn độ tương tự Cosine (Cosine Similarity) nhằm tìm kiếm ngữ nghĩa nhanh chóng giữa các hợp đồng."),
        ("Structured Output (Đầu ra cấu trúc):", " Ràng buộc mô hình trả về JSON thuần túy thông qua Pydantic Schema trong FastAPI (`AnalyzeResponse`, `FindingOutput`) kết hợp với hàm xử lý biểu thức chính quy `clean_and_parse_json` để lọc sạch các thẻ markdown trước khi phân tích cú pháp."),
        ("Fine-tuning (Tinh chỉnh mô hình):", " Đã thực hiện huấn luyện QLoRA (lượng tử hóa 4-bit thông qua thư viện `Unsloth` / `PEFT`) trên nền mô hình `Qwen2.5-3B-Instruct` sử dụng tập dữ liệu CUAD và dữ liệu thực tế mẫu để chuyên biệt hóa khả năng phân loại và phát hiện rủi ro hợp đồng (lưu trữ adapter tại `Doan2108/contract-risk-qwen2.5-3b-fix1`)."),
        ("Evaluation (Đánh giá chất lượng):", " Đo lường chất lượng mô hình sau huấn luyện thông qua chỉ số tổn thất (Validation Loss), độ hỗn loạn (Perplexity đạt 1.24) và chạy tập test case thực tế (đạt độ chính xác phân loại 80%)."),
        ("Guardrails (Hàng rào bảo vệ):", " Sử dụng tầng kiểm duyệt dữ liệu đầu ra JSON để validate các trường bắt buộc, lọc danh mục rủi ro hợp lệ và xử lý fallback để hệ thống hoạt động ổn định kể cả khi mô hình trả về lỗi.")
    ]

    for title, desc in ai_techniques:
        add_bullet(title, desc, indent_level=0)

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 6: MA TRẬN CRUD MASTER
    # =========================================================================
    add_heading_1("Chương 6. Ma Trận CRUD Tổng Hợp (CRUD Master Matrix)")
    
    doc.add_paragraph(
        "Dưới đây là ma trận mô tả chi tiết các tác động CRUD (Create, Read, Update, Delete) của các hoạt động nghiệp vụ chính lên CSDL của Hệ thống Quản lý Hợp đồng:"
    )
    
    crud_matrix = [
        ("Tạo Hợp đồng mới", "Hàm dịch vụ thực thi: CS.create_and_analyze_contract.\nTác động CSDL: Tạo mới bản ghi Contract (C), tạo mới ContractVersion (C), tạo mới ContractFile (C) để lưu trữ tệp AES-256-GCM, và bóc tách các trang thô lưu vào ContractContext (C)."),
        
        ("Tạo phiên bản mới", "Hàm dịch vụ thực thi: CS.create_new_version.\nTác động CSDL: Tạo mới ContractVersion (C) tăng số thứ tự v+1, tạo mới ContractFile (C) chứa file phiên bản mới, và bóc tách lưu trữ ContractContext (C) tương ứng."),
        
        ("Tách điều khoản tự động/AI", "Hàm dịch vụ thực thi: AES.extract_version.\nTác động CSDL: Đọc ContractContext (R) để lấy văn bản thô, xóa các điều khoản cũ Clause (D) cùng ExtractedEntity (D) cũ, tạo mới các bản ghi điều khoản Clause (C) và các thực thể ExtractedEntity (C)."),
        
        ("Tách điều khoản thủ công/Regex", "Hàm dịch vụ thực thi: CS.manual_extract_contract.\nTác động CSDL: Đọc ContractContext (R), xóa Clause (D) và ExtractedEntity (D) của phiên bản hiện tại, tạo mới Clause (C) và ExtractedEntity (C) cục bộ."),
        
        ("Phân tích Rủi ro AI", "Hàm dịch vụ thực thi: CS.analyze_contract.\nTác động CSDL: Cập nhật trạng thái Contract (U), đọc các điều khoản Clause (R), tạo mới AIAnalysis (C) và các bản ghi phát hiện rủi ro RiskFinding (C). Xóa các bản ghi phân tích cũ AIAnalysis (D) và RiskFinding (D) nếu phân tích lại."),
        
        ("Đánh giá của chuyên gia", "Hàm dịch vụ thực thi: CS.submit_expert_review.\nTác động CSDL: Cập nhật trạng thái Contract (U) thành APPROVED, đọc kết quả phân tích AIAnalysis (R), tạo mới bản ghi Review (C) lưu ý kiến đánh giá chuyên gia."),
        
        ("Đẩy lên Workflow phê duyệt", "Hàm dịch vụ thực thi: CS.push_to_workflow.\nTác động CSDL: Cập nhật trạng thái Contract (U) thành PENDING_WORKFLOW, đọc ContractVersion (R) để chuyển dữ liệu sang service phê duyệt."),
        
        ("Đăng ký Doanh nghiệp lên BC", "Hàm dịch vụ thực thi: BC.register_company.\nTác động CSDL: Cập nhật tx_hash và thông tin khối vào Company (U), tạo mới bản ghi BlockchainTransaction (C) và BlockchainAudit (C) để ghi vết giao dịch."),
        
        ("Đăng ký Người dùng lên BC", "Hàm dịch vụ thực thi: BC.register_user.\nTác động CSDL: Cập nhật tx_hash và thông tin khối vào User (U), tạo mới bản ghi BlockchainTransaction (C) và BlockchainAudit (C) trên database."),
        
        ("Neo bằng chứng lên Blockchain", "Hàm dịch vụ thực thi: BC.anchor_proof.\nTác động CSDL: Cập nhật trạng thái verified và thông tin đối soát vào HashProof (U), tạo mới bản ghi BlockchainTransaction (C) và BlockchainAudit (C)."),
        
        ("Đối soát Hợp đồng", "Hàm dịch vụ thực thi: BC.verify_proof.\nTác động CSDL: Đọc HashProof (R) và BlockchainTransaction (R), tạo mới bản ghi lịch sử VerificationHistory (C) ghi nhận kết quả đối soát toàn vẹn."),
        
        ("Ký duyệt số bước Workflow", "Hàm dịch vụ thực thi: BC.verify_and_sign.\nTác động CSDL: Đọc thông tin SignatureCertificate (R) và HashProof (R) để đối chiếu khóa công khai, tạo mới bản ghi DigitalSignature (C) lưu chữ ký số.")
    ]
    
    for action, detail in crud_matrix:
        add_bullet(action + ": ", "\n" + detail, indent_level=0)

    # Save document
    output_path = r"d:\Django_project\RiskDL\Bao_Cao_He_Thong_Quan_Ly_Hop_Dong_Master.docx"
    doc.save(output_path)
    print(f"Master Text Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_master_report()
