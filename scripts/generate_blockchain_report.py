import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color definitions (Deep Teal theme for Blockchain)
    c_primary = RGBColor(0, 90, 90)     # Deep Teal
    c_secondary = RGBColor(0, 130, 130)  # Medium Slate Teal
    c_text = RGBColor(51, 51, 51)        # Off-black
    c_gray = RGBColor(128, 128, 128)     # Gray
    
    # Configure default style font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = c_text
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    def add_custom_heading(text, level, space_before=12, space_after=6):
        h = doc.add_heading('', level=level)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        
        run = h.add_run(text)
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = c_primary
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = c_secondary
        elif level == 3:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = c_gray
        return h

    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr.append(borders)

    def add_code_block(code_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        
        cell = table.cell(0, 0)
        set_cell_background(cell, "F5F9F9")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        set_table_borders(table, color="B2D8D8", sz="6")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        run = p.add_run(code_text.strip())
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(40, 40, 40)
        
        # Add spacing after table
        empty_p = doc.add_paragraph()
        empty_p.paragraph_format.space_before = Pt(2)
        empty_p.paragraph_format.space_after = Pt(2)
        empty_p.paragraph_format.line_spacing = 1.0

    def add_explanation_bullet(code_line, source_info, explanation):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        
        run_code = p.add_run(code_line)
        run_code.font.name = 'Consolas'
        run_code.font.size = Pt(8.5)
        run_code.font.bold = True
        run_code.font.color.rgb = RGBColor(0, 90, 90)
        
        if source_info:
            p.add_run(" [")
            run_src = p.add_run(source_info)
            run_src.italic = True
            run_src.font.size = Pt(8.5)
            p.add_run("]")
            
        p.add_run(" : ")
        
        run_exp = p.add_run(explanation)
        run_exp.font.name = 'Calibri'
        run_exp.font.size = Pt(10)

    # --- TITLE PAGE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(100)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("BÁO CÁO TÍCH HỢP HỆ THỐNG BLOCKCHAIN HYPERLEDGER FABRIC CHI TIẾT")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = c_primary

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(100)
    run_sub = p_sub.add_run(
        "Hướng Dẫn Chi Tiết Hợp Nhất Giữa Luồng Hoạt Động Nghiệp Vụ Thực Tế\n"
        "Và Giải Thích Ý Nghĩa Từng Dòng Lệnh, Từng Biến Số Trong Mã Nguồn\n"
        "Bổ Sung Hướng Dẫn Cài Đặt, Vận Hành Và Xem Dữ Liệu Trên Giao Diện Hyperledger Explorer"
    )
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = c_secondary

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(0)
    run_meta = p_meta.add_run(
        "Hệ thống quản lý hợp đồng thông minh RiskDL\n"
        "Bản mô tả tích hợp chuỗi khối bảo đảm tính toàn vẹn và chống chối bỏ"
    )
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = c_gray

    doc.add_page_break()

    # --- SECTION 1: IDENTITY REGISTRY ---
    add_custom_heading("1. LUỒNG ĐĂNG KÝ DANH TÍNH DOANH NGHIỆP & NHÂN SỰ (IDENTITY REGISTRY)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("• Ví dụ thực tế: ").bold = True
    p.add_run(
        "Doanh nghiệp mới là \"Công ty Cổ phần Giải pháp Công nghệ ABC\" đăng ký tham gia vào hệ thống RiskDL với Mã số thuế là '0102030405'. "
        "Sau đó, quản trị viên tạo tài khoản nhân sự cho nhân viên tên \"Nguyễn Văn A\" với vai trò là \"MANAGER\" để phê duyệt hợp đồng. "
        "Thông tin định danh này phải được neo vĩnh viễn lên sổ cái Hyperledger Fabric để phục vụ việc xác thực nguồn gốc chứng thư và chữ ký số sau này."
    )

    p_flow = doc.add_paragraph()
    p_flow.add_run("• Mô tả luồng hoạt động trực quan:").bold = True
    p_flow_steps = (
        "\n1. [RiskDL Web App] : Quản trị viên nhập thông tin Công ty ABC và Nhân viên Nguyễn Văn A trên giao diện web và bấm Lưu."
        "\n2. [Django Client] : Web App lưu thông tin vào cơ sở dữ liệu (PostgreSQL) cục bộ, đồng thời gửi một yêu cầu HTTP POST chứa thông tin ID, Tên, Mã số thuế sang Blockchain Service."
        "\n3. [Django Blockchain Service] : Nhận gói tin, ghi nhận nhật ký kiểm toán (Audit Trail) và chuyển tiếp dữ liệu qua REST API sang Node.js Fabric Gateway."
        "\n4. [Node.js Fabric Gateway] : Sử dụng chứng thư bảo mật gRPC TLS để gọi giao dịch Smart Contract 'StoreCompany' / 'StoreUser' trên Peer của mạng Fabric."
        "\n5. [Hyperledger Fabric Ledger] : Smart Contract nhận mảng byte JSON, lưu cặp Key-Value định danh (\"COMPANY_X\", \"USER_Y\") vào World State. Giao dịch được xác nhận (committed) và trả về Tx ID cùng mã băm block thực tế để lưu trữ chéo vào cơ sở dữ liệu."
    )
    p_flow.add_run(p_flow_steps)

    # Django Client Code
    add_custom_heading("1.1 Phía Django Web App (Client) - Gọi API Đăng ký", level=2)
    code_django_id = """
# Nguồn: contracts/views.py (api_register_company)
import requests
import json
from django.http import JsonResponse

_BC_URL = "http://blockchain-service:8000"

def api_register_company(request):
    body = json.loads(request.body)
    name = body.get('company_name')
    tax_code = body.get('tax_code')
    
    # Tạo thực thể Company cục bộ
    company = Company.objects.create(company_name=name, tax_code=tax_code, status='ACTIVE')
    
    # Gửi dữ liệu định danh sang Blockchain Service
    resp = requests.post(f"{_BC_URL}/company/register/", json={
        'company_id': company.id,
        'company_name': name,
        'tax_code': tax_code,
        'status': 'ACTIVE',
        'sender': 'System'
    }, timeout=20)
    
    if resp.status_code == 200:
        data = resp.json()
        company.tx_hash = data.get('tx_hash')
        company.block_number = data.get('block_number')
        company.block_hash = data.get('block_hash')
        company.save()
    return JsonResponse({'status': 'success', 'company_id': company.id})
    """
    add_code_block(code_django_id)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("_BC_URL = \"http://...\"", "Biến cấu hình", "Địa chỉ cổng mạng dẫn đến dịch vụ Blockchain phụ trợ để chương trình biết nơi gửi dữ liệu.")
    add_explanation_bullet("def api_register_company(request)", "Khai báo hàm", "Hàm thực hiện đăng ký thông tin công ty. Biến 'request' chứa toàn bộ thông tin yêu cầu gửi lên từ trình duyệt của người dùng.")
    add_explanation_bullet("body = json.loads(request.body)", "Nhận từ trình duyệt", "Giải nén gói dữ liệu dạng văn bản nhận được từ trình duyệt và gán vào biến 'body' dưới dạng một danh mục dữ liệu dễ tra cứu.")
    add_explanation_bullet("name = body.get('company_name')", "Lấy thông tin", "Lấy giá trị tên công ty từ gói dữ liệu và gán vào biến 'name' (Ví dụ: \"Công ty ABC\").")
    add_explanation_bullet("tax_code = body.get('tax_code')", "Lấy thông tin", "Lấy giá trị mã số thuế từ gói dữ liệu và gán vào biến 'tax_code' (Ví dụ: \"0102030405\").")
    add_explanation_bullet("company = Company.objects.create(...)", "Lưu cơ sở dữ liệu", "Tạo một hàng dữ liệu mới trong bảng Công ty của cơ sở dữ liệu cục bộ với trạng thái hoạt động ban đầu là 'ACTIVE' (đang hoạt động) và lưu vào biến đại diện 'company'.")
    add_explanation_bullet("resp = requests.post(...)", "Giao tiếp mạng HTTP", "Gửi toàn bộ thông tin của công ty sang dịch vụ Blockchain phụ trợ thông qua đường truyền mạng, kết quả phản hồi nhận về được gán vào biến 'resp'.")
    add_explanation_bullet("timeout=20", "Tham số kết nối", "Giới hạn thời gian chờ tối đa là 20 giây, nếu quá thời gian này mà mạng không phản hồi thì sẽ ngắt kết nối để tránh treo hệ thống.")
    add_explanation_bullet("if resp.status_code == 200", "Kiểm tra kết quả", "Kiểm tra nếu kết quả trả về từ Blockchain Service là mã 200 (thành công rực rỡ).")
    add_explanation_bullet("data = resp.json()", "Nhận phản hồi", "Dịch dữ liệu nhận về từ định dạng mạng sang dạng biến để đọc được và gán vào biến 'data'.")
    add_explanation_bullet("company.tx_hash = data.get('tx_hash')", "Cập nhật dữ liệu", "Lấy mã biên nhận giao dịch (tx_hash - là một chuỗi mã số dài duy nhất chứng minh giao dịch đã lên chuỗi) và lưu vào thuộc tính của công ty.")
    add_explanation_bullet("company.block_number = data.get('block_number')", "Cập nhật dữ liệu", "Lấy số hiệu khối (block_number - số thứ tự trang sổ sách chứa giao dịch này trên Blockchain) và lưu vào công ty.")
    add_explanation_bullet("company.block_hash = data.get('block_hash')", "Cập nhật dữ liệu", "Lấy mã bảo mật của toàn bộ trang sổ sách (block_hash) và lưu vào công ty.")
    add_explanation_bullet("company.save()", "Lưu cơ sở dữ liệu", "Ghi lại toàn bộ các thông tin chuỗi khối vừa nhận được vào cơ sở dữ liệu cục bộ để hoàn tất.")
    add_explanation_bullet("return JsonResponse(...)", "Trả kết quả", "Gửi phản hồi thông báo thành công về lại trình duyệt cho người dùng.")

    # Django Blockchain Service Code
    add_custom_heading("1.2 Phía Django Blockchain Service (Middle Tier) - Xử lý Định danh", level=2)
    code_service_id = """
# Nguồn: blockchain_service/blockchain/services.py (EnterpriseRegistryService)
class EnterpriseRegistryService:
    @staticmethod
    def register_company(company_id, company_name, tax_code, status="ACTIVE", sender="System"):
        # Gọi tiếp sang Node.js Fabric Gateway
        resp = requests.post("http://fabric-gateway:5000/company/store", json={
            "companyId": str(company_id),
            "companyName": company_name,
            "taxCode": tax_code,
            "status": status
        }, timeout=15)
        
        data = resp.json()
        tx_hash = data.get("tx_hash")
        block_number = data.get("block_number")
        block_hash = data.get("block_hash")
        
        # Ghi nhận Transaction và tạo Audit Trail
        tx = BlockchainTransaction.objects.create(
            tx_hash=tx_hash, block_hash=block_hash, block_number=block_number, status="CONFIRMED"
        )
        BlockchainAudit.objects.create(
            transaction=tx, action="Register Company", resource=f"Company:{company_id}",
            before_state="UNREGISTERED", after_state="REGISTERED", status="SUCCESS"
        )
        return tx
    """
    add_code_block(code_service_id)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("company_id", "Tham số đầu vào", "Mã định danh (số ID tự tăng) của công ty trong hệ thống.")
    add_explanation_bullet("company_name", "Tham số đầu vào", "Chuỗi ký tự tên công ty (Ví dụ: \"Công ty ABC\").")
    add_explanation_bullet("tax_code", "Tham số đầu vào", "Mã số thuế của công ty.")
    add_explanation_bullet("status", "Tham số đầu vào", "Trạng thái hoạt động của công ty trên blockchain (mặc định là \"ACTIVE\").")
    add_explanation_bullet("sender", "Tham số đầu vào", "Tên tác nhân thực hiện yêu cầu này (mặc định là \"System\" - hệ thống tự động).")
    add_explanation_bullet("resp = requests.post(...)", "Giao tiếp mạng HTTP", "Lệnh gọi mạng gửi gói thông tin công ty sang Node.js Fabric Gateway ở cổng 5000, kết quả gán vào biến 'resp'.")
    add_explanation_bullet("data = resp.json()", "Nhận phản hồi", "Giải nén gói phản hồi từ Gateway và lưu vào biến 'data'.")
    add_explanation_bullet("tx_hash, block_number, block_hash", "Biến trung gian", "Các biến lưu trữ biên nhận giao dịch, số trang sổ và chữ ký mã hóa của trang sổ từ blockchain.")
    add_explanation_bullet("tx = BlockchainTransaction.objects.create(...)", "Lưu cơ sở dữ liệu", "Tạo một dòng dữ liệu biên nhận giao dịch mới trong bảng Lịch sử Giao dịch Blockchain của cơ sở dữ liệu phụ trợ và lưu vào biến 'tx'.")
    add_explanation_bullet("BlockchainAudit.objects.create(...)", "Lưu cơ sở dữ liệu", "Ghi chép một dòng nhật ký kiểm toán (Audit Log) để lưu vết lịch sử: trước đó công ty chưa đăng ký (UNREGISTERED), nay đã đăng ký thành công (REGISTERED).")

    # Node.js Gateway Code
    add_custom_heading("1.3 Phía Node.js Fabric Gateway (Bridge API) - Gửi Giao dịch", level=2)
    code_node_id = """
# Nguồn: fabric/gateway/server.js (POST /company/store)
app.post('/company/store', async (req, res) => {
    const { companyId, companyName, taxCode, status: companyStatus } = req.body;
    
    // Thực hiện gọi Smart Contract StoreCompany
    const commit = await contract.submitAsync('StoreCompany', {
        arguments: [String(companyId), companyName, taxCode, companyStatus]
    });
    
    const status = await commit.getStatus();
    const txId = status.transactionId;
    const blockInfo = await getBlockInfoByTxId(txId);
    
    res.json({
        status: 'CONFIRMED',
        tx_hash: txId,
        block_number: blockInfo.blockNumber,
        block_hash: blockInfo.blockHash
    });
});
    """
    add_code_block(code_node_id)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("req", "Tham số đầu vào", "Biến chứa dữ liệu yêu cầu gửi đến từ phía Django Blockchain Service.")
    add_explanation_bullet("res", "Tham số đầu vào", "Biến dùng để gửi phản hồi kết quả về lại cho Django Blockchain Service.")
    add_explanation_bullet("companyId, companyName, taxCode, companyStatus", "Tách biến dữ liệu", "Các biến tách ra từ nội dung yêu cầu (req.body) đại diện cho thông tin công ty cần lưu.")
    add_explanation_bullet("commit = await contract.submitAsync(...)", "Gọi Blockchain Gateway", "Gửi bất đồng bộ yêu cầu thực thi StoreCompany lên Peer để lấy chữ ký đồng thuận (endorsement) và gửi sang Orderer xếp khối, trả về đối tượng giao dịch lưu vào biến 'commit'.")
    add_explanation_bullet("status = await commit.getStatus()", "Gọi Blockchain Gateway", "Đợi phản hồi xác thực khối từ các Peer (Committer) để chắc chắn giao dịch đã được ghi vĩnh viễn vào sổ cái, trả về trạng thái gán vào biến 'status'.")
    add_explanation_bullet("txId = status.transactionId", "Lấy dữ liệu", "Mã định danh giao dịch duy nhất của đợt đăng ký này được trích xuất từ trạng thái giao dịch và gán vào biến 'txId'.")
    add_explanation_bullet("blockInfo = await getBlockInfoByTxId(txId)", "Truy vấn System Chaincode", "Gọi hàm phụ để tìm kiếm số hiệu khối và mã băm khối chứa giao dịch này, thông tin trả về lưu vào biến 'blockInfo'.")
    add_explanation_bullet("res.json(...)", "Trả kết quả", "Trả về dữ liệu kết quả thành công dưới dạng JSON cho Django Blockchain Service.")

    # Smart Contract Code
    add_custom_heading("1.4 Phía Smart Contract (Go Chaincode) - Ghi Sổ cái", level=2)
    code_go_id = """
// Nguồn: fabric/chaincode/contract_verify.go (StoreCompany)
func (s *SmartContract) StoreCompany(ctx contractapi.TransactionContextInterface, companyID string, companyName string, taxCode string, status string) error {
	record := CompanyRecord{
		CompanyID:   companyID,
		CompanyName: companyName,
		TaxCode:     taxCode,
		Status:      status,
	}

	recordBytes, err := json.Marshal(record)
	if err != nil {
		return err
	}

	return ctx.GetStub().PutState("COMPANY_"+companyID, recordBytes)
}
    """
    add_code_block(code_go_id)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("s *SmartContract", "Đối tượng Smart Contract", "Đối tượng đại diện cho Hợp đồng thông minh hiện hành đang chạy.")
    add_explanation_bullet("ctx", "Tham số ngữ cảnh", "Tham số ngữ cảnh giao dịch của mạng lưới, giúp giao tiếp với sổ cái của nút mạng Peer.")
    add_explanation_bullet("companyID, companyName, taxCode, status", "Tham số đầu vào", "Các tham số đầu vào đại diện cho thông tin công ty do Gateway gửi tới.")
    add_explanation_bullet("record = CompanyRecord{...}", "Khởi tạo đối tượng", "Tạo một đối tượng cấu trúc dữ liệu Công ty để chuẩn bị lưu trữ và gán vào biến 'record'.")
    add_explanation_bullet("recordBytes = json.Marshal(record)", "Chuyển đổi dữ liệu", "Chuyển đổi đối tượng sang chuỗi byte dạng JSON để lưu trữ, kết quả gán vào biến 'recordBytes'. Biến 'err' dùng để bắt lỗi nếu quá trình chuyển đổi thất bại.")
    add_explanation_bullet("ctx.GetStub().PutState(...)", "Gọi cơ sở dữ liệu chuỗi", "Ghi chuỗi byte dữ liệu vào sổ cái World State dưới dạng khóa chính là COMPANY_ + ID của công ty. Khóa này giúp tìm kiếm thông tin công ty nhanh chóng sau này.")

    doc.add_page_break()

    # --- SECTION 2: DOCUMENT ANCHORING ---
    add_custom_heading("2. LUỒNG NEO VẾT BẰNG CHỨNG DỮ LIỆU HỢP ĐỒNG (DOCUMENT ANCHORING)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("• Ví dụ thực tế: ").bold = True
    p.add_run(
        "Nhân viên soạn thảo xong phiên bản 2 của \"Hợp đồng Thuê Văn phòng\". "
        "Hệ thống cần băm nội dung hợp đồng này thành mã băm bảo mật, liên kết nó với mã băm của phiên bản 1 bằng thuật toán Merkle Root "
        "và neo (anchor) mã băm này lên Hyperledger Fabric để đảm bảo tệp tin không bị sửa đổi trái phép."
    )

    p_flow = doc.add_paragraph()
    p_flow.add_run("• Mô tả luồng hoạt động trực quan:").bold = True
    p_flow_steps = (
        "\n1. [RiskDL Web App] : Người dùng bấm nút 'Neo lên Blockchain' cho hợp đồng v2."
        "\n2. [Django Client] : Đọc tệp PDF v2 từ thư mục vật lý, dùng module mã hóa để giải mã tệp thô tạm thời trong bộ nhớ RAM, chuyển nội dung thô cùng ID phiên bản trước sang Blockchain Service."
        "\n3. [Django Blockchain Service] : Tính SHA-256 của v2. Đồng thời lấy mã băm v1 đã lưu trong DB và thực hiện phép băm kép: SHA-256(Hash_v2 + Hash_v1) tạo gốc Merkle Root của chuỗi phiên bản. Lưu HashProof cục bộ và gửi sang Node.js Gateway."
        "\n4. [Node.js Fabric Gateway] : Nhận yêu cầu và kích hoạt giao dịch 'StoreHash' chứa ProofID, Document Hash và Merkle Root lên Blockchain."
        "\n5. [Hyperledger Fabric Ledger] : Smart Contract kiểm tra tính tồn tại của ProofID, sau đó ghi thông tin băm vào sổ cái qua PutState. Trả về giao dịch thành công."
    )
    p_flow.add_run(p_flow_steps)

    # Django Client Code
    add_custom_heading("2.1 Phía Django Web App (Client) - Giải mã PDF và Gọi Băm", level=2)
    code_django_anchor = """
# Nguồn: contracts/views.py (_bc_version_payload)
from .crypto_utils import decrypt_pdf

def _bc_version_payload(version):
    content = None
    file_obj = version.files.first()
    if file_obj and file_obj.file_path:
        full_path = os.path.join(settings.MEDIA_ROOT, file_obj.file_path.lstrip(settings.MEDIA_URL))
        with open(full_path, 'rb') as f:
            encrypted_data = f.read()
        # Giải mã PDF trong bộ nhớ RAM
        decrypted_data = decrypt_pdf(encrypted_data)
        content = decrypted_data.decode('utf-8', errors='ignore')
        
    prev_version = version.contract.versions.filter(version_number=version.version_number - 1).first()
    prev_version_id = prev_version.id if prev_version else None

    return {
        "version_id": version.id, "content": content,
        "contract_code": version.contract.contract_code,
        "version_number": version.version_number,
        "previous_version_id": prev_version_id
    }
    """
    add_code_block(code_django_anchor)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("def _bc_version_payload(version)", "Khai báo hàm", "Hàm chuẩn bị gói dữ liệu phiên bản hợp đồng gửi đi. Biến 'version' đại diện cho đối tượng phiên bản hợp đồng hiện tại.")
    add_explanation_bullet("file_obj = version.files.first()", "Đọc dữ liệu", "Lấy file đính kèm đầu tiên của phiên bản hợp đồng này gán vào biến 'file_obj'.")
    add_explanation_bullet("full_path = os.path.join(...)", "Tìm đường dẫn", "Xác định đường dẫn vật lý đầy đủ của file hợp đồng trên ổ đĩa máy chủ và gán vào biến 'full_path'.")
    add_explanation_bullet("with open(full_path, 'rb') as f", "Đọc file", "Mở file hợp đồng ở chế độ đọc nhị phân đại diện bởi biến 'f'.")
    add_explanation_bullet("encrypted_data = f.read()", "Đọc dữ liệu", "Đọc toàn bộ nội dung file đã bị mã hóa và lưu vào biến 'encrypted_data'.")
    add_explanation_bullet("decrypted_data = decrypt_pdf(encrypted_data)", "Giải mã tệp tin", "Sử dụng thuật toán giải mã để đưa nội dung file về dạng nhị phân thô trong bộ nhớ RAM và lưu vào biến 'decrypted_data'.")
    add_explanation_bullet("content = decrypted_data.decode(...)", "Chuyển đổi dữ liệu", "Chuyển mã nhị phân thô sang dạng chuỗi văn bản UTF-8 đọc được và gán vào biến 'content'.")
    add_explanation_bullet("prev_version = version.contract.versions.filter(...)", "Tìm dữ liệu", "Truy tìm phiên bản liền trước (Ví dụ: v1) của hợp đồng này và gán vào biến 'prev_version'.")
    add_explanation_bullet("prev_version_id = prev_version.id ...", "Lấy dữ liệu", "Lấy mã ID của phiên bản trước gán vào biến 'prev_version_id'. Nếu không có phiên bản trước, biến này sẽ là rỗng (None).")

    # Django Blockchain Service Code
    add_custom_heading("2.2 Phía Django Blockchain Service (Middle Tier) - Tạo Merkle Root & Neo chuỗi", level=2)
    code_service_anchor = """
# Nguồn: blockchain_service/blockchain/services.py (ProofService)
import hashlib

class ProofService:
    @staticmethod
    def generate_proof(version_id, content=None, previous_version_id=None):
        content_bytes = content.encode('utf-8')
        doc_hash = hashlib.sha256(content_bytes).hexdigest()
        
        previous_hash = None
        if previous_version_id:
            prev_proof = HashProof.objects.get(version_id=previous_version_id)
            previous_hash = prev_proof.document_hash

        # Thuật toán Merkle Root liên kết chuỗi phiên bản
        if previous_hash:
            merkle_root = hashlib.sha256((doc_hash + previous_hash).encode('utf-8')).hexdigest()
        else:
            merkle_root = hashlib.sha256(doc_hash.encode('utf-8')).hexdigest()

        proof = HashProof.objects.create(
            version_id=version_id, document_hash=doc_hash,
            previous_hash=previous_hash, merkle_root=merkle_root, verified=False
        )
        return proof
    """
    add_code_block(code_service_anchor)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("version_id", "Tham số đầu vào", "Mã số định danh phiên bản hợp đồng.")
    add_explanation_bullet("content", "Tham số đầu vào", "Chuỗi văn bản thô nội dung hợp đồng.")
    add_explanation_bullet("previous_version_id", "Tham số đầu vào", "Mã số định danh của phiên bản liền trước.")
    add_explanation_bullet("doc_hash = hashlib.sha256(content_bytes).hexdigest()", "Mã hóa băm", "Thực hiện tính toán mã băm bảo mật SHA-256 (tạo chuỗi 64 ký tự đại diện duy nhất cho nội dung tệp) gán vào biến 'doc_hash'.")
    add_explanation_bullet("prev_proof = HashProof.objects.get(...)", "Tìm dữ liệu", "Lấy dữ liệu băm của phiên bản trước từ cơ sở dữ liệu gán vào biến 'prev_proof'.")
    add_explanation_bullet("previous_hash = prev_proof.document_hash", "Lấy dữ liệu", "Lấy mã băm của phiên bản trước gán vào biến 'previous_hash'.")
    add_explanation_bullet("merkle_root = hashlib.sha256(...)", "Mã hóa băm kép", "Thực hiện giải thuật băm nối chuỗi (Merkle Chain): băm kết hợp mã băm hiện tại và mã băm cũ để sinh ra một mã gốc gán vào biến 'merkle_root'. Nếu bản cũ bị đổi, mã gốc này sẽ lập tức sai lệch.")
    add_explanation_bullet("proof = HashProof.objects.create(...)", "Lưu cơ sở dữ liệu", "Tạo một bản ghi bằng chứng băm mới trong cơ sở dữ liệu cục bộ với trạng thái xác thực ban đầu là False (chưa neo blockchain thành công), kết quả trả về lưu vào biến 'proof'.")

    # Node.js Gateway Code
    add_custom_heading("2.3 Phía Node.js Fabric Gateway (Bridge API) - Ghi mã băm", level=2)
    code_node_anchor = """
# Nguồn: fabric/gateway/server.js (POST /anchor)
app.post('/anchor', async (req, res) => {
    const { proofId, documentHash, merkleRoot } = req.body;
    const timestamp = new Date().toISOString();
    
    // Gọi phương thức ghi StoreHash của Chaincode
    const commit = await contract.submitAsync('StoreHash', {
        arguments: [String(proofId), documentHash, merkleRoot, timestamp]
    });
    
    const status = await commit.getStatus();
    const txId = status.transactionId;
    const blockInfo = await getBlockInfoByTxId(txId);
    
    res.json({
        status: 'CONFIRMED', tx_hash: txId,
        block_number: blockInfo.blockNumber, block_hash: blockInfo.blockHash
    });
});
    """
    add_code_block(code_node_anchor)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("proofId, documentHash, merkleRoot", "Tách biến dữ liệu", "Các biến chứa mã số định danh phiên bản, mã băm tài liệu và gốc liên kết phiên bản truyền sang từ yêu cầu mạng.")
    add_explanation_bullet("timestamp = new Date().toISOString()", "Lấy thời gian thực", "Tạo nhãn thời gian hiện tại theo định dạng chuẩn quốc tế gán vào biến 'timestamp'.")
    add_explanation_bullet("commit = await contract.submitAsync(...)", "Gọi Blockchain Gateway", "Giao dịch ghi StoreHash được gửi bất đồng bộ lên chuỗi khối và gán vào biến 'commit'.")
    add_explanation_bullet("status = await commit.getStatus()", "Gọi Blockchain Gateway", "Chờ phản hồi xác thực khối từ các Peer trên blockchain và gán kết quả trạng thái vào biến 'status'.")
    add_explanation_bullet("txId = status.transactionId", "Lấy dữ liệu", "Lấy mã định danh giao dịch chuỗi khối gán vào biến 'txId'.")
    add_explanation_bullet("blockInfo = await getBlockInfoByTxId(txId)", "Truy vấn System Chaincode", "Truy tìm số hiệu khối và mã băm khối chứa giao dịch gán vào biến 'blockInfo'.")

    # Smart Contract Code
    add_custom_heading("2.4 Phía Smart Contract (Go Chaincode) - StoreHash", level=2)
    code_go_anchor = """
// Nguồn: fabric/chaincode/contract_verify.go (StoreHash)
func (s *SmartContract) StoreHash(ctx contractapi.TransactionContextInterface, proofID string, documentHash string, merkleRoot string, timestamp string) error {
	exists, err := s.HashExists(ctx, proofID)
	if exists {
		return fmt.Errorf("the hash for proof %s already exists", proofID)
	}

	record := HashRecord{
		ProofID:      proofID,
		DocumentHash: documentHash,
		MerkleRoot:   merkleRoot,
		Timestamp:    timestamp,
	}
	recordBytes, err := json.Marshal(record)
	return ctx.GetStub().PutState(proofID, recordBytes)
}
    """
    add_code_block(code_go_anchor)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("exists, err := s.HashExists(ctx, proofID)", "Gọi Blockchain API", "Kiểm tra xem mã số định danh proofID đã tồn tại trên sổ cái chưa, kết quả Đúng/Sai gán vào biến 'exists'. Biến 'err' dùng để lưu lỗi nếu có.")
    add_explanation_bullet("if exists", "Kiểm tra điều kiện", "Nếu mã băm đã tồn tại từ trước, báo lỗi và dừng giao dịch để bảo toàn tính bất biến của lịch sử.")
    add_explanation_bullet("record := HashRecord{...}", "Khởi tạo đối tượng", "Tạo đối tượng HashRecord chứa bằng chứng băm tài liệu gán vào biến 'record'.")
    add_explanation_bullet("recordBytes, err := json.Marshal(record)", "Chuyển đổi dữ liệu", "Mã hóa đối tượng sang chuỗi byte JSON gán vào biến 'recordBytes' để ghi xuống cơ sở dữ liệu.")
    add_explanation_bullet("ctx.GetStub().PutState(proofID, recordBytes)", "Gọi cơ sở dữ liệu chuỗi", "Ghi mảng byte JSON trực tiếp vào cơ sở dữ liệu World State của Peer với khóa chính là proofID.")

    doc.add_page_break()

    # --- SECTION 3: INTEGRITY VERIFICATION ---
    add_custom_heading("3. LUỒNG KIỂM ĐỊNH TÍNH TOÀN VẸN TÀI LIỆU HỢP ĐỒNG (INTEGRITY VERIFICATION)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("• Ví dụ thực tế: ").bold = True
    p.add_run(
        "Đối tác gửi lại bản PDF hợp đồng thuê văn phòng v2. "
        "Người dùng tải bản PDF này lên RiskDL để thực hiện đối soát. "
        "Hệ thống sẽ băm nội dung tài liệu này và so sánh trực tiếp với bản ghi băm gốc đã neo trên Hyperledger Fabric từ trước. "
        "Nếu có bất kỳ một ký tự nào bị sửa (ví dụ: thay đổi giá thuê từ $1000 thành $100), hệ thống sẽ phát hiện ra ngay lập tức."
    )

    p_flow = doc.add_paragraph()
    p_flow.add_run("• Mô tả luồng hoạt động trực quan:").bold = True
    p_flow_steps = (
        "\n1. [RiskDL Web App] : Người dùng click vào chức năng 'Kiểm định Hợp đồng' trên giao diện."
        "\n2. [Django Client] : Đọc tệp PDF hiện tại, giải mã bằng khóa đối xứng và gửi nội dung thô cùng ID sang Blockchain Service."
        "\n3. [Django Blockchain Service] : Tính toán mã băm SHA-256 thực tế hiện tại. Gọi API `/verify` của Node.js Gateway."
        "\n4. [Node.js Fabric Gateway] : Gọi hàm `GetHash` dạng truy vấn đọc (evaluateTransaction) để lấy dữ liệu gốc từ World State mà không qua Orderer (giúp phản hồi siêu tốc)."
        "\n5. [Django Blockchain Service & Client] : So sánh mã băm thực tế hiện tại với mã băm nhận về từ sổ cái. Nếu trùng khớp -> Xác nhận tệp toàn vẹn. Nếu sai lệch -> Cảnh báo tệp tin bị giả mạo."
    )
    p_flow.add_run(p_flow_steps)

    # Django Client Code
    add_custom_heading("3.1 Phía Django Web App (Client) - Yêu cầu Kiểm định", level=2)
    code_django_verify = """
# Nguồn: contracts/views.py (api_blockchain_verify_proof)
def api_blockchain_verify_proof(request, contract_id):
    body = json.loads(request.body) if request.body else {}
    contract, version, err = _get_version_or_latest(contract_id, body.get('version_id'))
    
    # Gửi payload văn bản thô sang Blockchain Service
    resp = requests.post(f"{_BC_URL}/proofs/verify/", json=_bc_version_payload(version), timeout=15)
    result = resp.json()
    
    # Trả về kết quả so sánh băm
    return JsonResponse({
        'contract_code': contract.contract_code,
        'verified': result.get('verified'),
        'proof_hash': result.get('proof_hash'),
        'current_hash': result.get('current_hash'),
        'blockchain_anchored': result.get('blockchain_anchored')
    })
    """
    add_code_block(code_django_verify)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("contract_id", "Tham số đầu vào", "Mã số định danh của hợp đồng cần kiểm tra.")
    add_explanation_bullet("contract, version, err = _get_version_or_latest(...)", "Tìm dữ liệu", "Truy tìm thông tin hợp đồng và phiên bản tương ứng gán vào các biến 'contract' và 'version'. Biến 'err' dùng để chứa lỗi nếu không tìm thấy.")
    add_explanation_bullet("resp = requests.post(...)", "Giao tiếp mạng HTTP", "Gửi gói tin chứa văn bản thô của hợp đồng sang Blockchain Service để thực hiện đối chiếu và gán phản hồi vào biến 'resp'.")
    add_explanation_bullet("result = resp.json()", "Nhận phản hồi", "Dịch phản hồi nhận được sang dạng biến dễ dùng gán vào biến 'result'.")
    add_explanation_bullet("return JsonResponse({...})", "Trả kết quả", "Gửi phản hồi báo kết quả xác minh về lại trình duyệt để hiển thị màu xanh (Hợp lệ) hoặc màu đỏ (Đã bị chỉnh sửa) cho người dùng.")

    # Django Blockchain Service Code
    add_custom_heading("3.2 Phía Django Blockchain Service (Middle Tier) - Tính băm & So sánh", level=2)
    code_service_verify = """
# Nguồn: blockchain_service/blockchain/services.py (VerificationService)
class VerificationService:
    @staticmethod
    def verify_proof(version_id, content=None):
        proof = HashProof.objects.filter(version_id=version_id).first()
        content_bytes = content.encode('utf-8')
        current_hash = hashlib.sha256(content_bytes).hexdigest()

        # 1. So sánh cục bộ
        if current_hash != proof.document_hash:
            return {"verified": False, "message": "Document hash mismatch."}

        # 2. Truy vấn thực tế từ mạng Hyperledger Fabric qua Gateway
        blockchain_anchored = False
        resp = requests.post("http://fabric-gateway:5000/verify", json={"version_id": str(version_id)})
        
        if resp.status_code == 200:
            data = resp.json()
            if data.get("verified") and data.get("proof_hash") == current_hash:
                blockchain_anchored = True

        return {
            "verified": blockchain_anchored,
            "proof_hash": proof.document_hash,
            "current_hash": current_hash,
            "blockchain_anchored": blockchain_anchored
        }
    """
    add_code_block(code_service_verify)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("proof = HashProof.objects.filter(...).first()", "Tìm cơ sở dữ liệu", "Tìm bản ghi bằng chứng băm đã lưu trữ cục bộ cho phiên bản này và gán vào biến 'proof'.")
    add_explanation_bullet("current_hash = hashlib.sha256(...).hexdigest()", "Mã hóa băm", "Tính toán mã băm SHA-256 thực tế của tài liệu tải lên và gán vào biến 'current_hash'.")
    add_explanation_bullet("if current_hash != proof.document_hash", "Kiểm tra điều kiện", "Kiểm định cấp 1: So đối chiếu mã băm thực tế hiện tại với mã băm lưu tại cơ sở dữ liệu cục bộ. Nếu sai lệch báo lỗi ngay.")
    add_explanation_bullet("resp = requests.post(...)", "Giao tiếp mạng HTTP", "Kiểm định cấp 2: Gọi sang Node.js Gateway gửi yêu cầu truy xuất dữ liệu gốc trên blockchain để đảm bảo cơ sở dữ liệu cục bộ không bị thay đổi, gán phản hồi vào biến 'resp'.")
    add_explanation_bullet("blockchain_anchored = True", "Gán trạng thái", "Đặt biến cờ hiệu 'blockchain_anchored' thành True để xác nhận bằng chứng đã được neo giữ an toàn trên chuỗi khối.")

    # Node.js Gateway Code
    add_custom_heading("3.3 Phía Node.js Fabric Gateway (Bridge API) - Đọc Sổ cái", level=2)
    code_node_verify = """
# Nguồn: fabric/gateway/server.js (POST /verify)
app.post('/verify', async (req, res) => {
    const { version_id } = req.body;
    
    // Gọi phương thức evaluate (chỉ đọc, không ghi)
    const resultBytes = await contract.evaluateTransaction('GetHash', String(version_id));
    const record = JSON.parse(Buffer.from(resultBytes).toString());

    res.json({
        verified: true,
        proof_hash: record.documentHash,
        blockchain_anchored: true
    });
});
    """
    add_code_block(code_node_verify)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("version_id", "Tách biến dữ liệu", "Lấy mã số định danh phiên bản hợp đồng truyền sang từ yêu cầu mạng.")
    add_explanation_bullet("resultBytes = await contract.evaluateTransaction(...)", "Gọi Blockchain Gateway", "Thực hiện truy vấn đọc dữ liệu GetHash từ chuỗi khối. Lệnh này gửi yêu cầu đến Peer chạy chaincode lấy dữ liệu trạng thái mà không gửi giao dịch qua Orderer (giúp phản hồi siêu tốc) và gán mảng byte nhận được vào biến 'resultBytes'.")
    add_explanation_bullet("record = JSON.parse(...)", "Giải mã dữ liệu", "Chuyển mảng byte nhận được về dạng chuỗi văn bản và giải nén thành đối tượng dữ liệu dễ đọc gán vào biến 'record'.")
    add_explanation_bullet("res.json({...})", "Trả kết quả", "Phản hồi kết quả mã băm đã được lưu vĩnh viễn trên sổ cái chuỗi khối về cho Django Blockchain Service.")

    # Smart Contract Code
    add_custom_heading("3.4 Phía Smart Contract (Go Chaincode) - GetHash", level=2)
    code_go_verify = """
// Nguồn: fabric/chaincode/contract_verify.go (GetHash)
func (s *SmartContract) GetHash(ctx contractapi.TransactionContextInterface, proofID string) (*HashRecord, error) {
	recordBytes, err := ctx.GetStub().GetState(proofID)
	if err != nil {
		return nil, fmt.Errorf("failed to read from world state: %v", err)
	}
	if recordBytes == nil {
		return nil, fmt.Errorf("the hash for proof %s does not exist", proofID)
	}

	var record HashRecord
	err = json.Unmarshal(recordBytes, &record)
	return &record, nil
}
    """
    add_code_block(code_go_verify)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("proofID", "Tham số đầu vào", "Mã số định danh của bằng chứng băm cần truy vấn.")
    add_explanation_bullet("recordBytes, err := ctx.GetStub().GetState(proofID)", "Gọi cơ sở dữ liệu chuỗi", "Gọi hàm của nút mạng Peer để đọc dữ liệu thô dạng byte của khóa proofID từ cơ sở dữ liệu World State, kết quả gán vào biến 'recordBytes' và biến chứa lỗi 'err'.")
    add_explanation_bullet("if recordBytes == nil", "Kiểm tra điều kiện", "Nếu không tìm thấy khóa nào tương ứng trên sổ cái, báo lỗi dừng truy vấn.")
    add_explanation_bullet("err = json.Unmarshal(recordBytes, &record)", "Giải mã dữ liệu", "Giải nén chuỗi byte JSON thô thành cấu trúc đối tượng HashRecord gán vào biến 'record'.")

    doc.add_page_break()

    # --- SECTION 4: SIGNING AND HISTORY ---
    add_custom_heading("4. LUỒNG KÝ DUYỆT SỐ & TRUY VẤN LỊCH SỬ GIAO DỊCH (SIGNING & HISTORY QUERY)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("• Ví dụ thực tế: ").bold = True
    p.add_run(
        "Trưởng phòng Nguyễn Văn A tiến hành duyệt hợp đồng. "
        "Hệ thống kiểm tra chứng thư số 'CERT-ABC-0001-4928' của anh A xem có hợp lệ hay không, "
        "sau đó tạo một chữ ký số liên kết trực tiếp tài khoản anh A với mã băm của hợp đồng. "
        "Ngoài ra, khi có tranh chấp, kiểm toán viên có thể bấm truy vấn lịch sử để hiển thị toàn bộ nhật ký thay đổi của hợp đồng vĩnh viễn trên blockchain."
    )

    p_flow = doc.add_paragraph()
    p_flow.add_run("• Mô tả luồng hoạt động trực quan:").bold = True
    p_flow_steps = (
        "\n1. [Ký duyệt số] : Django Web gửi yêu cầu ký duyệt kèm chữ ký số và ID chứng thư số của nhân viên."
        "\n2. [Django Blockchain Service] : Xác minh chứng thư số trong DB xem còn hạn dùng và có đúng của User không. Nếu hợp lệ, ghi nhận DigitalSignature liên kết với HashProof của hợp đồng."
        "\n3. [Truy vấn lịch sử] : Người dùng chọn xem lịch sử giao dịch chuỗi khối của một phiên bản hợp đồng."
        "\n4. [Node.js Gateway -> Peer] : Gọi hàm `GetHistory` của Chaincode. Peer thực hiện quét cơ sở dữ liệu lịch sử giao dịch (HistoryDB) của Hyperledger Fabric."
        "\n5. [Kết quả] : Trả về danh sách tất cả các Transaction ID, dấu thời gian block và thông tin ghi nhận trong lịch sử, biểu diễn sơ đồ vòng đời phiên bản hợp đồng."
    )
    p_flow.add_run(p_flow_steps)

    # Signature Service Code
    add_custom_heading("4.1 Phía Django Blockchain Service (Middle Tier) - Xác minh & Ký duyệt", level=2)
    code_service_sign = """
# Nguồn: blockchain_service/blockchain/services.py (SignatureService)
class SignatureService:
    @staticmethod
    def verify_and_sign(step_id, user_id, certificate_id, signature_hash):
        cert = SignatureCertificate.objects.filter(id=certificate_id).first()
        if not cert or cert.user_id != user_id:
            raise ValueError("Certificate does not belong to the user")
        if cert.status != "ACTIVE" or cert.revoked or cert.valid_to < timezone.now():
            raise ValueError("Certificate is invalid, revoked, or expired")

        # Tìm bằng chứng băm hợp đồng mới nhất để ký duyệt
        proof = HashProof.objects.order_by('-generated_at').first()
        
        sig = DigitalSignature.objects.create(
            certificate=cert, hashproof=proof, signature=signature_hash,
            algorithm=cert.signature_algorithm, verified=True, verified_at=timezone.now()
        )
        return {"status": "success", "signature_id": sig.id}
    """
    add_code_block(code_service_sign)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("step_id, user_id, certificate_id, signature_hash", "Tham số đầu vào", "Các tham số đầu vào đại diện cho ID bước phê duyệt, ID người dùng, ID chứng thư số sử dụng và chuỗi băm của chữ ký.")
    add_explanation_bullet("cert = SignatureCertificate.objects.filter(...).first()", "Tìm cơ sở dữ liệu", "Tìm chứng thư số của người dùng trong cơ sở dữ liệu và gán vào biến 'cert'.")
    add_explanation_bullet("cert.valid_to < timezone.now()", "Kiểm tra thời gian", "Kiểm tra xem chứng thư đã hết hạn sử dụng hay chưa so với thời gian thực tại.")
    add_explanation_bullet("proof = HashProof.objects.order_by(...).first()", "Tìm cơ sở dữ liệu", "Lấy bằng chứng băm hợp đồng mới nhất vừa được neo chuỗi để làm đối tượng ký và gán vào biến 'proof'.")
    add_explanation_bullet("sig = DigitalSignature.objects.create(...)", "Lưu cơ sở dữ liệu", "Tạo một bản ghi Chữ ký số liên kết chứng thư số và bằng chứng băm hợp đồng vào cơ sở dữ liệu cục bộ để làm bằng chứng pháp lý chống chối bỏ, kết quả lưu vào biến 'sig'.")

    # Node.js Gateway Code
    add_custom_heading("4.2 Phía Node.js Fabric Gateway (Bridge API) - Truy vấn Lịch sử", level=2)
    code_node_history = """
# Nguồn: fabric/gateway/server.js (GET /history/:proofId)
app.get('/history/:proofId', async (req, res) => {
    const { proofId } = req.params;
    
    // Gọi evaluateTransaction để chạy hàm GetHistory
    const resultBytes = await contract.evaluateTransaction('GetHistory', String(proofId));
    const history = JSON.parse(Buffer.from(resultBytes).toString());

    res.json({ history });
});
    """
    add_code_block(code_node_history)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("proofId", "Tách biến dữ liệu", "Mã số định danh bằng chứng băm cần tra cứu lịch sử được lấy từ tham số đường dẫn URL.")
    add_explanation_bullet("resultBytes = await contract.evaluateTransaction(...)", "Gọi Blockchain Gateway", "Gửi yêu cầu truy vấn GetHistory đến Peer để duyệt cơ sở dữ liệu lịch sử các giao dịch ghi nhận khóa proofId này và gán kết quả vào biến 'resultBytes'.")
    add_explanation_bullet("history = JSON.parse(...)", "Giải mã dữ liệu", "Giải nén chuỗi kết quả nhận được thành đối tượng mảng lịch sử gán vào biến 'history'.")

    # Smart Contract Code
    add_custom_heading("4.3 Phía Smart Contract (Go Chaincode) - GetHistory", level=2)
    code_go_history = """
// Nguồn: fabric/chaincode/contract_verify.go (GetHistory)
func (s *SmartContract) GetHistory(ctx contractapi.TransactionContextInterface, proofID string) ([]HistoryQueryResult, error) {
	resultsIterator, err := ctx.GetStub().GetHistoryForKey(proofID)
	if err != nil {
		return nil, err
	}
	defer resultsIterator.Close()

	var records []HistoryQueryResult
	for resultsIterator.HasNext() {
		queryResponse, err := resultsIterator.Next()
		
		var record HashRecord
		if !queryResponse.IsDelete {
			_ = json.Unmarshal(queryResponse.Value, &record)
		}

		txTimestamp := time.Unix(queryResponse.Timestamp.Seconds, int64(queryResponse.Timestamp.Nanos)).Format(time.RFC3339)

		historicalRecord := HistoryQueryResult{
			TxID:      queryResponse.TxId,
			Value:     &record,
			Timestamp: txTimestamp,
			IsDelete:  queryResponse.IsDelete,
		}
		records = append(records, historicalRecord)
	}
	return records, nil
}
    """
    add_code_block(code_go_history)
    p = doc.add_paragraph("Giải thích ý nghĩa từng dòng lệnh và các biến số:")
    add_explanation_bullet("resultsIterator, err := ctx.GetStub().GetHistoryForKey(proofID)", "Gọi cơ sở dữ liệu chuỗi", "Đọc lịch sử thay đổi của một khóa trên Sổ cái chuỗi khối. Hàm này truy vấn cơ sở dữ liệu lịch sử ghi nhận tất cả giao dịch sửa đổi khóa proofID từ trước tới nay, kết quả lưu vào biến duyệt 'resultsIterator'.")
    add_explanation_bullet("resultsIterator.HasNext()", "Vòng lặp", "Vòng lặp chạy qua từng giao dịch lịch sử được tìm thấy.")
    add_explanation_bullet("queryResponse, err := resultsIterator.Next()", "Duyệt dữ liệu", "Lấy thông tin giao dịch lịch sử tiếp theo gán vào biến 'queryResponse'.")
    add_explanation_bullet("txTimestamp := time.Unix(...)", "Xử lý thời gian", "Đọc thông tin giây và nano giây được mạng lưới đồng thuận đóng dấu khi tạo khối và chuyển đổi về chuỗi ký tự thời gian chuẩn ISO 8601 gán vào biến 'txTimestamp'.")
    add_explanation_bullet("historicalRecord := HistoryQueryResult{...}", "Khởi tạo đối tượng", "Tạo đối tượng chứa thông tin lịch sử bao gồm TxID (Mã giao dịch), Value (Giá trị bằng chứng băm tại thời điểm đó), Timestamp (Thời gian ghi nhận) và IsDelete (Cờ báo xóa) gán vào biến 'historicalRecord'.")

    doc.add_page_break()

    # --- NEW SECTION 5: HYPERLEDGER EXPLORER ---
    add_custom_heading("5. HƯỚNG DẪN CÀI ĐẶT & SỬ DỤNG GIAO DIỆN KIỂM GIÁM (HYPERLEDGER EXPLORER)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("5.1 Giao diện Hyperledger Explorer là gì và nó chứa cái gì?").bold = True
    p = doc.add_paragraph(
        "Hyperledger Explorer là một công cụ giám sát trực quan (Web Dashboard) dành riêng cho mạng lưới Blockchain Hyperledger Fabric. "
        "Công cụ này hoạt động giống như một trình duyệt khối (tương tự Etherscan của Ethereum hay Blockchain.com của Bitcoin), "
        "giúp người dùng quản trị và kiểm toán viên dễ dàng quan sát, tìm kiếm dữ liệu mà không cần phải viết code."
    )
    p = doc.add_paragraph("Giao diện Web UI của Explorer (truy cập tại cổng 8090) chứa các mục thông tin chính sau:")
    add_explanation_bullet("Dashboard (Trang tổng quan)", "", "Hiển thị thống kê tổng số lượng Khối (Blocks) đã tạo, tổng số Giao dịch (Transactions), số lượng nút mạng đang chạy (Peers) và số lượng Hợp đồng thông minh (Chaincodes) đã cài đặt.")
    add_explanation_bullet("Blocks (Danh sách Khối)", "", "Hiển thị danh sách các khối trong chuỗi theo thứ tự thời gian. Mỗi khối chứa số thứ tự khối, mã băm tiêu đề khối, mã băm dữ liệu và số lượng giao dịch chứa trong khối đó.")
    add_explanation_bullet("Transactions (Danh sách Giao dịch)", "", "Hiển thị chi tiết từng giao dịch được gửi lên chuỗi. Người dùng có thể tìm kiếm theo Transaction ID, xem thời gian ghi sổ, tổ chức ký số giao dịch (ví dụ: Org1MSP) và tên hàm được gọi (ví dụ: StoreHash, StoreUser).")
    add_explanation_bullet("Chaincodes & Network", "", "Hiển thị danh sách các hợp đồng thông minh đang chạy trên kênh và sơ đồ kết nối vật lý của các nút Peer trong mạng lưới.")

    p = doc.add_paragraph()
    p.add_run("5.2 Khi đẩy dữ liệu lên Fabric, Explorer ghi nhận cụ thể cái gì và lấy dữ liệu từ đâu?").bold = True
    p = doc.add_paragraph(
        "Khi hệ thống RiskDL thực hiện các hành động ghi dữ liệu lên Blockchain (như Đăng ký công ty, đăng ký Nhân viên, hay Neo mã băm hợp đồng):\n"
        "1. Dữ liệu được đóng gói thành giao dịch và ghi nhận vào sổ cái dưới dạng Khối (Block). Mỗi giao dịch chứa các tham số đầu vào dưới dạng mảng JSON (ví dụ: mã số thuế, tên công ty, hoặc mã băm SHA-256 của file PDF hợp đồng).\n"
        "2. Explorer có một dịch vụ ứng dụng chạy ngầm (Node.js background process) kết nối liên tục với nút Peer (peer0.org1.example.com) thông qua giao thức truyền tin gRPC TLS ở cổng 7051.\n"
        "3. Mỗi khi nút Peer ghi nhận một khối mới được đồng thuận, Peer sẽ phát ra sự kiện (Block Event). Ứng dụng ngầm của Explorer lập tức lắng nghe thấy, tự động tải khối dữ liệu nhị phân thô đó về, giải mã cấu trúc Protobuf và lưu trữ các thông tin này vào một cơ sở dữ liệu PostgreSQL phụ trợ chuyên biệt đặt tên là 'fabricexplorer'.\n"
        "4. Giao diện Web hiển thị trên trình duyệt sẽ truy vấn trực tiếp dữ liệu từ cơ sở dữ liệu PostgreSQL này để vẽ biểu đồ và hiển thị danh sách cho người dùng xem một cách nhanh chóng, thay vì phải quét lại toàn bộ blockchain thô từ đầu."
    )

    p = doc.add_paragraph()
    p.add_run("5.3 Kiểm tra thông tin người dùng (User) đã đăng ký ở đâu trên sổ cái?").bold = True
    p = doc.add_paragraph(
        "Khi một nhân viên (User) được đăng ký thành công trên RiskDL, thông tin sẽ được lưu trữ dưới dạng cặp Key-Value trên World State với khóa chính là USER_<userID> (ví dụ: USER_1).\n"
        "Để xem thông tin người dùng này trên sổ cái, chúng ta có hai cách thực hiện:"
    )
    add_explanation_bullet("Cách 1: Xem qua Giao dịch trên Hyperledger Explorer", "", "Vào tab 'Transactions' trên giao diện Explorer, tìm giao dịch có tên hàm gọi là 'StoreUser'. Bấm vào Transaction ID đó để mở bảng chi tiết. Tại mục 'Writeset' (Tập tin ghi sổ), giao diện sẽ hiển thị rõ khóa 'USER_1' và toàn bộ thông tin nhân viên dưới dạng JSON (Username, CompanyID, Role, Status) được ghi nhận vĩnh viễn trên sổ cái.")
    add_explanation_bullet("Cách 2: Truy vấn Cơ sở dữ liệu World State (CouchDB Fauxton)", "", "Do nút Peer sử dụng CouchDB để lưu trạng thái hiện tại (World State), chúng ta có thể truy cập trực tiếp vào giao diện quản trị CouchDB Fauxton tại địa chỉ http://localhost:5984/_utils/ trên máy chủ chạy node. Vào database tên là 'contracts-channel_contractverifychaincode', bạn sẽ thấy danh sách tất cả các tài liệu có khóa bắt đầu bằng tiền tố 'USER_' chứa thông tin nhân sự thô đang được lưu giữ.")

    p = doc.add_paragraph()
    p.add_run("5.4 Hướng dẫn thiết lập và cài đặt Hyperledger Explorer").bold = True
    p = doc.add_paragraph(
        "Hệ thống Explorer được cài đặt và tích hợp trực tiếp vào tệp docker-compose.yml của dự án. "
        "Dưới đây là phần cấu hình cài đặt và giải thích hoạt động:"
    )

    code_docker_explorer = """
# Nguồn cấu hình: docker-compose.yml (explorer-db & explorer)
services:
  explorer-db:
    image: hyperledger/explorer-db:latest
    container_name: explorer_db
    environment:
      - DATABASE_DATABASE=fabricexplorer
      - DATABASE_USERNAME=hppoc
      - DATABASE_PASSWORD=password
    networks:
      - riskdl-net

  explorer:
    image: hyperledger/explorer:latest
    container_name: explorer
    environment:
      - DATABASE_HOST=explorer-db
      - DATABASE_PORT=5432
      - DATABASE_DATABASE=fabricexplorer
      - DATABASE_USERNAME=hppoc
      - DATABASE_PASSWORD=password
      - DISCOVERY_AS_LOCALHOST=false
    volumes:
      - ./fabric/explorer/config.json:/opt/explorer/app/platform/fabric/config.json
      - ./fabric/explorer/connection-profile.json:/opt/explorer/app/platform/fabric/connection-profile/first-network.json
      - ./fabric/crypto-config:/tmp/crypto
    ports:
      - "8090:8080"
    depends_on:
      - explorer-db
      - fabric-gateway
    networks:
      - riskdl-net
    """
    add_code_block(code_docker_explorer)
    p = doc.add_paragraph("Giải thích các biến số và cơ chế hoạt động của tệp cài đặt:")
    add_explanation_bullet("explorer-db", "Dịch vụ Container", "Chạy một container hệ quản trị cơ sở dữ liệu PostgreSQL làm kho chứa dữ liệu lịch sử được đồng bộ từ blockchain về để phục vụ giao diện Web.")
    add_explanation_bullet("DATABASE_DATABASE=fabricexplorer", "Biến môi trường", "Đặt tên cơ sở dữ liệu là 'fabricexplorer' trong hệ thống PostgreSQL.")
    add_explanation_bullet("explorer", "Dịch vụ Container", "Chạy container giao diện Hyperledger Explorer Web App, chạy trên nền Node.js.")
    add_explanation_bullet("DATABASE_HOST=explorer-db", "Biến môi trường", "Chỉ định địa chỉ máy chủ chứa cơ sở dữ liệu là container 'explorer-db'.")
    add_explanation_bullet("./fabric/explorer/config.json...", "Định tuyến thư mục (Volume mount)", "Kết nối file cấu hình danh sách mạng chuỗi khối cục bộ vào container.")
    add_explanation_bullet("./fabric/explorer/connection-profile.json...", "Định tuyến thư mục (Volume mount)", "Kết nối file hồ sơ kết nối mạng. File này chứa địa chỉ IP và số hiệu cổng gRPC (7051) của Peer, danh sách các channel và cách xác thực TLS.")
    add_explanation_bullet("./fabric/crypto-config:/tmp/crypto", "Định tuyến thư mục (Volume mount)", "Truyền toàn bộ các file khóa bí mật và chứng thư số admin của tổ chức (Org1) đã được sinh ra bởi script khởi tạo vào trong container để Explorer sử dụng làm thông tin xác thực admin kết nối mạng.")
    add_explanation_bullet("ports: - \"8090:8080\"", "Định tuyến cổng mạng", "Mở cổng 8090 trên máy tính thật kết nối vào cổng 8080 bên trong container. Người dùng truy cập http://localhost:8090 từ trình duyệt để xem giao diện quản trị.")

    doc.add_page_break()

    # --- SECTION 6: TECHNICAL SUMMARY ---
    add_custom_heading("6. TỔNG KẾT HỆ THỐNG & CẤU HÌNH CÔNG NGHỆ", level=1)
    p = doc.add_paragraph("Bảng tổng hợp chi tiết vai trò của các thư viện và cấu phần công nghệ tham gia vào hệ thống:")

    table = doc.add_table(rows=11, cols=3)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="CCCCCC")
    
    col_widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Lớp Công nghệ"
    hdr_cells[1].text = "Công cụ / Thư viện"
    hdr_cells[2].text = "Ý nghĩa & Tác dụng"
    
    for cell in hdr_cells:
        set_cell_background(cell, "005A5A") # Deep Teal
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    data = [
        ["Django Client", "requests", "Dùng để gửi dữ liệu và lệnh gọi mạng qua lại giữa các dịch vụ."],
        ["Django Client", "Cơ sở dữ liệu cục bộ", "Lưu trữ thông tin công ty, nhân sự và phiên bản hợp đồng phục vụ nghiệp vụ."],
        ["Blockchain Service", "hashlib (Thuật toán SHA-256)", "Mã hóa băm một chiều nội dung hợp đồng thô thành chuỗi mã bảo mật."],
        ["Blockchain Service", "timezone", "Lấy thời gian thực hành động của người dùng để lưu vào sổ sách."],
        ["Node.js Gateway", "@hyperledger/fabric-gateway", "Bộ công cụ SDK chính thống để giao tiếp và gửi giao dịch lên chuỗi khối."],
        ["Node.js Gateway", "gRPC với bảo mật SSL/TLS", "Đường truyền mạng mã hóa bảo mật kết nối trực tiếp đến Peer0 ở cổng 7051."],
        ["Node.js Gateway", "crypto", "Đọc file khóa bí mật ECDSA để ký xác thực giao dịch."],
        ["Go Chaincode", "contractapi", "Bộ thư viện khung để xây dựng các hàm giao dịch của hợp đồng thông minh."],
        ["Go Chaincode", "shim", "Cung cấp các lệnh đọc ghi dữ liệu sổ cái chuỗi khối (PutState, GetState, GetHistoryForKey)."],
        ["Go Chaincode", "encoding/json", "Chuyển đổi dữ liệu đối tượng sang dạng chuỗi JSON để lưu trữ."]
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            if col_idx == 0:
                set_cell_background(row_cells[col_idx], "F0F8F8")
                row_cells[col_idx].paragraphs[0].runs[0].font.bold = True
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Save document
    output_filename = "Bao_Cao_Ky_Thuat_Blockchain_Integration_Chi_Tiet.docx"
    doc.save(output_filename)
    print(f"[SUCCESS] Re-generated simplified Blockchain Technical report with Explorer saved to file: {output_filename}")

if __name__ == "__main__":
    create_report()
