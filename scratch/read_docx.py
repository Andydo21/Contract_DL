import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    for i, p in enumerate(doc.paragraphs):
        if "Query 7" in p.text or "Compare" in p.text:
            print(f"P{i}: {p.text}")
            # print surrounding paragraphs
            for j in range(max(0, i-2), min(len(doc.paragraphs), i+5)):
                print(f"  P{j}: {doc.paragraphs[j].text}")
            print("-" * 50)

try:
    read_docx(r"d:\Django_project\FSOFT\SQL_Assignment4_Đỗ_Đăng An.docx")
except Exception as e:
    print(f"Error: {e}")
