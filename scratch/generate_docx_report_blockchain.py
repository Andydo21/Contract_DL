import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set cell background color in docx table."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # ------------------ COVER PAGE / TITLE ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO KỸ THUẬT PHÂN HỆ")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue 900
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_p2.add_run("BLOCKCHAIN INTEGRATION FLOW & ENTITY INTERACTIONS")
    title_run2.font.name = 'Arial'
    title_run2.font.size = Pt(24)
    title_run2.font.bold = True
    title_run2.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Phân tích chi tiết quy trình Neo giữ mã băm (Anchoring), Đăng ký danh tính (Identity Registry), Chữ ký số và Tác động đến thực thể CSDL")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Separator Line
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep_run = p_sep.add_run("―" * 40)
    p_sep_run.font.color.rgb = RGBColor(226, 232, 240)
    
    doc.add_page_break()
    
    # ------------------ SECTION 1: KIẾN TRÚC TÍCH HỢP ------------------
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. Kiến Trúc Tích Hợp Hyperledger Fabric")
    r1.font.name = 'Arial'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(30, 58, 138)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Phân hệ Blockchain trong RiskDL được xây dựng trên nền tảng Hyperledger Fabric "
        "(mạng Blockchain dành cho doanh nghiệp với quyền kiểm soát truy cập). Hệ thống giao tiếp thông qua "
        "cổng dịch vụ Fabric Gateway (Node.js API chạy trên cổng 5000). Gateway này gọi đến Chaincode "
        "(Smart Contract) cài đặt trên các Peer để cập nhật sổ cái (Ledger)."
    )
    p.paragraph_format.space_after = Pt(10)
    
    p_fb = doc.add_paragraph(
        "Cơ chế Phòng ngừa Sự cố (Fault-Tolerance): Trong trường hợp Fabric Gateway bị ngoại tuyến, "
        "các dịch vụ Blockchain tự động chuyển sang cơ chế Giả Lập Neo Giữ (Simulated Anchoring). "
        "Hệ thống sẽ tạo ra các mã giao dịch (tx_hash) và số block ngẫu nhiên để ghi nhận vào Database cục bộ, "
        "đảm bảo luồng nghiệp vụ của người dùng không bao giờ bị gián đoạn."
    )
    p_fb.paragraph_format.space_after = Pt(10)
    
    doc.add_page_break()
    
    # ------------------ SECTION 2: CÁC LUỒNG NGHIỆP VỤ ------------------
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. Các Luồng Nghiệp Vụ Blockchain Chi Tiết")
    r2.font.name = 'Arial'
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(30, 58, 138)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    # Flow 1
    p_f1 = doc.add_paragraph()
    p_f1.add_run("Luồng 1: Đăng ký chứng thư số & Danh tính (Identity Registry)").bold = True
    p_f1.paragraph_format.space_before = Pt(6)
    p_f1.paragraph_format.space_after = Pt(4)
    
    p_f1_desc = doc.add_paragraph(
        "Khi một Doanh nghiệp hoặc Người dùng mới được thêm vào hệ thống, thông tin định danh của họ "
        "(như tên doanh nghiệp, mã số thuế, mã người dùng, vai trò) được đẩy lên Smart Contract của Fabric "
        "thông qua API `/company/store` và `/user/store` để đăng ký danh tính hợp pháp. Đồng thời, "
        "hệ thống cấp phát chứng thư số cá nhân (SignatureCertificate) dùng cho việc ký duyệt."
    )
    
    # Flow 2
    p_f2 = doc.add_paragraph()
    p_f2.add_run("Luồng 2: Tạo và Neo Bằng Chứng Hợp Đồng (Proof Generation & Anchoring)").bold = True
    p_f2.paragraph_format.space_before = Pt(6)
    p_f2.paragraph_format.space_after = Pt(4)
    
    p_f2_desc = doc.add_paragraph(
        "Mỗi khi một phiên bản hợp đồng (ContractVersion) được tải lên: "
        "1) Hệ thống băm nhị phân nội dung file bằng SHA-256 tạo ra document_hash. "
        "2) Tính Merkle Root bằng cách nối hash hiện tại với hash của phiên bản liền trước để tạo liên kết chuỗi phiên bản (Version Chain). "
        "3) Lưu trữ thông tin này vào HashProof. "
        "4) Gọi Fabric Gateway gửi giao dịch neo giữ lên sổ cái Blockchain. "
        "5) Nhận tx_hash, block_number trả về và cập nhật vào bản ghi BlockchainTransaction để lưu dấu vết đối soát vĩnh viễn."
    )
    
    # Flow 3
    p_f3 = doc.add_paragraph()
    p_f3.add_run("Luồng 3: Ký duyệt hợp đồng bằng chữ ký số (Digital Signing)").bold = True
    p_f3.paragraph_format.space_before = Pt(6)
    p_f3.paragraph_format.space_after = Pt(4)
    
    p_f3_desc = doc.add_paragraph(
        "Khi quản lý duyệt hợp đồng trong Workflow: "
        "1) Hệ thống kiểm tra tính hợp lệ của chứng thư số cá nhân (chưa hết hạn, chưa bị thu hồi). "
        "2) Kết hợp mã khóa bí mật của người dùng với mã băm tài liệu trong HashProof để tạo ra mã chữ ký số độc nhất. "
        "3) Lưu trữ bản ghi chữ ký vào bảng DigitalSignature kết nối chặt chẽ giữa tài liệu và danh tính người ký."
    )
    
    # Flow 4
    p_f4 = doc.add_paragraph()
    p_f4.add_run("Luồng 4: Kiểm tra và đối soát tính toàn vẹn (Integrity Verification)").bold = True
    p_f4.paragraph_format.space_before = Pt(6)
    p_f4.paragraph_format.space_after = Pt(4)
    
    p_f4_desc = doc.add_paragraph(
        "Để xác thực hợp đồng có bị chỉnh sửa trái phép hay không: "
        "Hệ thống tự động tải file hợp đồng hiện tại lên bộ nhớ, tính toán lại mã băm SHA-256 cục bộ, "
        "rồi gửi yêu cầu truy vấn lên sổ cái Blockchain qua API `/verify` của Fabric Gateway. "
        "Nếu mã băm khớp và thông tin giao dịch neo giữ trùng khớp, hợp đồng được xác nhận toàn vẹn (Integrity Verified)."
    )
    
    doc.add_page_break()
    
    # ------------------ SECTION 3: MÔ TẢ CÁC HÀM BLOCKCHAIN ------------------
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Chi Tiết Lớp Dịch Vụ Blockchain (blockchain/services.py)")
    r3.font.name = 'Arial'
    r3.font.size = Pt(16)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(30, 58, 138)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    funcs_blockchain = [
        ("CertificateService.register_certificate", "user_id, serial_number, issuer, valid_days=365", "Tạo và đăng ký một chứng thư số SignatureCertificate mới cho người dùng."),
        ("ProofService.generate_proof", "version_id, content=None, contract_code, version_number", "Tính toán mã băm SHA-256 từ nội dung file hợp đồng và tính Merkle Root kết nối phiên bản trước."),
        ("BlockchainAnchorService.anchor_proof", "proof_id, network_id=1, smart_contract_id=1", "Thực hiện gửi giao dịch neo giữ bằng chứng lên Blockchain và tạo bản ghi BlockchainTransaction & BlockchainAudit."),
        ("VerificationService.verify_proof", "version_id, content=None, previous_version_id=None", "Thực hiện đối soát mã băm hiện tại của tài liệu với mã băm đã lưu trên Blockchain ledger và ghi lịch sử xác thực."),
        ("SignatureService.verify_and_sign", "step_id, user_id, certificate_id, signature_hash", "Kiểm tra tính hợp lệ của chứng thư số cá nhân và ký duyệt lên tài liệu (tạo bản ghi DigitalSignature)."),
        ("EnterpriseRegistryService.register_company", "company_id, company_name, tax_code", "Gửi yêu cầu lưu trữ thông tin doanh nghiệp lên sổ cái Blockchain thông qua Fabric Gateway."),
        ("EnterpriseRegistryService.register_user", "user_id, username, company_id, role", "Gửi yêu cầu đăng ký thông tin danh tính và phân quyền người dùng lên sổ cái Blockchain.")
    ]
    
    table_bc = doc.add_table(rows=1, cols=3)
    table_bc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_bc.autofit = False
    
    col_widths = [Inches(2.0), Inches(1.8), Inches(2.7)]
    
    hdr_cells = table_bc.rows[0].cells
    headers = ["Class & Tên hàm", "Tham số chính", "Mô tả ý nghĩa chức năng"]
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "1E3A8A")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths[idx]
        
    for name, params, desc in funcs_blockchain:
        row_cells = table_bc.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[1].text = params
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.0)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[2].text = desc
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths[idx]
            
    doc.add_page_break()
    
    # ------------------ SECTION 4: TÁC ĐỘNG TỚI ENTITIES ------------------
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Tác Động Tới Các Thực Thể Cơ Sở Dữ Liệu (Database Entities)")
    r4.font.name = 'Arial'
    r4.font.size = Pt(16)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(30, 58, 138)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Các hoạt động của phân hệ Blockchain tương tác trực tiếp và làm thay đổi trạng thái của các thực thể trong cơ sở dữ liệu như sau:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    entities_impact = [
        ("SignatureCertificate", "Lưu trữ chứng thư số công cộng được cấp phát cho người dùng. Trạng thái ACTIVE/REVOKED được quản lý để kiểm tra tính hợp lệ khi ký số."),
        ("HashProof", "Thực thể trung tâm lưu giữ mã băm SHA-256 của hợp đồng, Merkle Root và trạng thái xác thực (verified = True/False)."),
        ("BlockchainNetwork", "Lưu cấu hình mạng Blockchain kết nối (RPC Endpoint, Chain Type). Được tham chiếu bởi các giao dịch để phân biệt các môi trường."),
        ("SmartContract", "Lưu thông tin Smart Contract/Chaincode (Tên chaincode, Địa chỉ contract, Phiên bản)."),
        ("BlockchainTransaction", "Ghi nhận chi tiết từng giao dịch ghi lên sổ cái (tx_hash, block_number, confirm_time, status CONFIRMED/FAILED)."),
        ("BlockchainAudit", "Nhật ký kiểm toán ghi nhận lịch sử thay đổi trạng thái (ví dụ: chuyển đổi từ UNVERIFIED sang VERIFIED của HashProof)."),
        ("DigitalSignature", "Thực thể liên kết giữa chứng thư số (SignatureCertificate) và bằng chứng băm (HashProof) để xác thực người dùng đã ký duyệt phiên bản này."),
        ("VerificationHistory", "Lưu trữ nhật ký các lần kiểm tra tính toàn vẹn hợp đồng (thời gian kiểm tra, kết quả đúng/sai, lý do mismatch nếu có).")
    ]
    
    table_ent = doc.add_table(rows=1, cols=2)
    table_ent.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ent.autofit = False
    
    col_widths_ent = [Inches(2.2), Inches(4.3)]
    
    hdr_cells = table_ent.rows[0].cells
    hdr_headers = ["Thực thể (Entity / Table)", "Hành động & Mức độ ảnh hưởng"]
    for idx, name in enumerate(hdr_headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "334155") # Slate 700
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths_ent[idx]
        
    for name, desc in entities_impact:
        row_cells = table_ent.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(79, 70, 229)
        
        row_cells[1].text = desc
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths_ent[idx]
            
    # Save document
    output_path = r"d:\Django_project\RiskDL\Bao_Cao_Ky_Thuat_Blockchain_Integration_Chi_Tiet.docx"
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_report()
