import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_future_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # Helper to add headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        return p
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(79, 70, 229)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_bullet(bold_prefix, text_content, indent_level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
        p.paragraph_format.space_after = Pt(3)
        r_bold = p.add_run(bold_prefix)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(30, 41, 59)
        
        r_text = p.add_run(text_content)
        r_text.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # ------------------ COVER PAGE ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(48)
    title_run = title_p.add_run("BÁO CÁO KẾ HOẠCH PHÁT TRIỂN TƯƠNG LAI")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_p2.add_run("LỘ TRÌNH NÂNG CẤP CÔNG NGHỆ AI & PHẦN CỨNG GIẢ LẬP (DỰ ÁN QUẢN LÝ HỢP ĐỒNG)")
    title_run2.font.name = 'Arial'
    title_run2.font.size = Pt(22)
    title_run2.font.bold = True
    title_run2.font.color.rgb = RGBColor(79, 70, 229)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Đặc tả chi tiết giải pháp nâng cấp 4 dịch vụ AI cốt lõi và tích hợp hệ thống phần cứng nhúng biên (IoT) thông qua công cụ mô phỏng Software-in-the-Loop (SIL)")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run("Tác giả: ").bold = True
    p_meta.add_run("Đỗ Đăng An (dodangan)\n")
    p_meta.add_run("Phiên bản tài liệu: ").bold = True
    p_meta.add_run("v1.1.0 (Bổ sung hiện trạng công nghệ)\n")
    p_meta.add_run("Mục tiêu tích hợp: ").bold = True
    p_meta.add_run("Conversational RAG, Hybrid Search, GraphRAG, Agentic Workflows, QEMU, Wokwi/Proteus Simulation\n")
    p_meta.paragraph_format.line_spacing = 1.3
    
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep_run = p_sep.add_run("―" * 50)
    p_sep_run.font.color.rgb = RGBColor(226, 232, 240)
    
    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 1: TỔNG QUAN VÀ ĐỊNH HƯỚNG PHÁT TRIỂN
    # =========================================================================
    add_heading_1("Chương 1. Tổng Quan Định Hướng Phát Triển Dự Án Quản Lý Hợp Đồng")
    
    doc.add_paragraph(
        "Dự án Quản lý hợp đồng phiên bản hiện tại đã thiết lập được khung nền tảng vững chắc với sự kết hợp của Django Backend, "
        "FastAPI AI Service và sổ cái Blockchain Hyperledger Fabric. Để đáp ứng các yêu cầu bảo mật cấp độ doanh nghiệp "
        "và nâng cao độ chính xác cũng như trải nghiệm tương tác thông minh, dự án Quản lý hợp đồng xác định hai trục nâng cấp chính trong tương lai:\n"
        "1. Trục Nâng cấp Trí tuệ Nhân tạo (AI Roadmap): Tối ưu hóa sâu sắc 4 phân hệ dịch vụ AI cốt lõi (Risk Analysis, Semantic Search, AI Extract, AI Summary) dựa trên nền tảng kỹ thuật hiện tại.\n"
        "2. Trục Tích hợp Vật lý - Phần cứng Nhúng biên (IoT Simulation): Xây dựng giải pháp giám sát vật lý, chống đánh tráo hợp đồng giấy gốc bằng tủ hồ sơ thông minh, camera chụp biên và ki-ốt tự phục vụ."
    )
    
    doc.add_paragraph(
        "Tài liệu này tập trung làm rõ phương án kiến trúc, mô tả chi tiết cách phát triển và các kỹ thuật nâng cấp cụ thể cho từng thành phần."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 2: HIỆN TRẠNG CÔNG NGHỆ VÀ TÌNH HÌNH TRIỂN KHAI HIỆN TẠI
    # =========================================================================
    add_heading_1("Chương 2. Hiện Trạng Công Nghệ Và Tình Hình Triển Khai Hiện Tại")
    
    doc.add_paragraph(
        "Trước khi triển khai kế hoạch nâng cấp dài hạn, tình hình triển khai thực tế của các phân hệ thuộc dự án Quản lý Hợp đồng "
        "được tổng kết chi tiết như sau nhằm làm rõ điểm xuất phát công nghệ:"
    )
    
    add_heading_2("2.1 Tình hình hiện tại của các Dịch vụ AI (AI Services Current Status)")
    
    add_bullet(
        "Dịch vụ Phân tích Rủi ro (Risk Analysis Service): ",
        "Đã hoàn thành huấn luyện QLoRA (lượng tử hóa 4-bit) trên nền mô hình Qwen2.5-3B-Instruct bằng tập dữ liệu CUAD và dữ liệu hợp đồng thực tế. "
        "Mô hình đạt độ hỗn loạn (Perplexity) cực thấp là 1.24 và độ chính xác phân loại rủi ro trên tập test đạt 80%. "
        "API phân tích đã được tích hợp qua FastAPI, trả về cấu trúc rủi ro có định dạng rõ ràng nhờ Pydantic validation."
    )
    
    add_bullet(
        "Dịch vụ Tìm kiếm Ngữ nghĩa (Semantic Search Service): ",
        "Đã triển khai thành công mô hình nhúng dense vector Multilingual-E5-Base để tạo các vector đặc trưng 768 chiều cho từng điều khoản hợp đồng. "
        "Sử dụng cơ sở dữ liệu vector ChromaDB chạy cục bộ để lập chỉ mục và thực hiện tìm kiếm tương tự ngữ nghĩa dựa trên độ đo Cosine."
    )
    
    add_bullet(
        "Dịch vụ Trích xuất AI (AI Extract Service): ",
        "Đã xây dựng bộ Regex ClauseSplitter dự phòng và thuật toán bóc tách heuristics để phân tách các hợp đồng tiếng Việt thành các Clause riêng biệt. "
        "Đã tích hợp thư viện PaddleOCR để nhận diện ký tự quang học đối với các hợp đồng dạng ảnh quét hoặc PDF Scanned."
    )
    
    add_bullet(
        "Dịch vụ Tóm tắt AI (AI Summary Service): ",
        "Sử dụng mô hình gốc Qwen2.5-3B-Instruct thông qua các tham số prompt tối ưu để tự động sinh tóm tắt nội dung hợp đồng bằng tiếng Việt dưới dạng executive summary và lưu trữ vào cơ sở dữ liệu."
    )
    
    add_heading_2("2.2 Tình hình hiện tại của Phân hệ Phần cứng nhúng biên (IoT Current Status)")
    
    add_bullet(
        "Thiết bị phần cứng thực tế: ",
        "Do giới hạn về mặt ngân sách đầu tư và chi phí thiết bị ban đầu, dự án chưa mua sắm các thiết bị vật lý thực tế như Raspberry Pi, cảm biến vân tay AS608, đầu đọc NFC PN532 hay chốt khóa Solenoid."
    )
    
    add_bullet(
        "Mô phỏng và giả lập phần mềm: ",
        "Dự án đã thiết lập thành công mô hình cấu trúc bảng cơ sở dữ liệu lưu trữ nhật ký thiết bị trên Django (bảng CabinetAccessLog, CabinetDocumentEvent) "
        "và xây dựng các script giả lập logic phần cứng bằng Python (như mock_hardware.py và các driver giả lập tín hiệu GPIO ảo). "
        "Toàn bộ tính năng phần cứng hiện tại được vận hành và thử nghiệm thông qua các công cụ giả lập để chứng minh tính đúng đắn về mặt logic nghiệp vụ."
    )
    
    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 3: KẾ HOẠCH PHÁT TRIỂN CHO TỪNG DỊCH VỤ AI CHI TIẾT
    # =========================================================================
    add_heading_1("Chương 3. Kế Hoạch Nâng Cấp Chi Tiết Cho Từng Dịch Vụ AI Cốt Lõi")
    
    doc.add_paragraph(
        "Dưới đây là kế hoạch phát triển chi tiết cho 4 dịch vụ AI cốt lõi, chỉ rõ cách kế thừa và nâng cấp các kỹ thuật đã triển khai thực tế trong hệ thống:"
    )

    # ------------------ RISK ANALYSIS ------------------
    add_heading_2("3.1 Dịch vụ Phân tích Rủi ro (Risk Analysis Service)")
    doc.add_paragraph(
        "Hiện trạng đã triển khai: Hệ thống sử dụng một mô hình Qwen2.5-3B-Instruct đã được tinh chỉnh qua kỹ thuật QLoRA (lượng tử hóa 4-bit) "
        "để phân tích các điều khoản và phát hiện rủi ro (RiskFinding) dựa trên một System Prompt chi tiết và ràng buộc Pydantic Schema."
    )
    doc.add_paragraph(
        "Kế hoạch nâng cấp và phương pháp phát triển:"
    )
    
    add_bullet(
        "Tích hợp luồng Conversational RAG hỏi đáp pháp lý: ",
        "Thay vì chỉ cung cấp một danh sách rủi ro tĩnh một chiều (như điểm số overall_score và RiskFinding), hệ thống sẽ mở rộng giao diện hội thoại "
        "cho phép chuyên gia pháp lý đặt câu hỏi chất vấn ngược lại mô hình. Ví dụ: 'Tại sao Điều 5 lại bị chấm điểm rủi ro cao?', 'Đề xuất phương án sửa đổi "
        "điều khoản này để có lợi nhất cho Bên B'. Hệ thống sẽ tự động trích xuất điều khoản liên quan làm ngữ cảnh để LLM trả lời chuẩn xác."
    )
    
    add_bullet(
        "Xây dựng kiến trúc AI Agentic Workflows (Tác tử thông minh): ",
        "Thay thế luồng suy luận tuyến tính một bước bằng mô hình tác tử lập kế hoạch (Planning / ReAct). AI Agent sẽ tự động phân rã quy trình "
        "phân tích một hợp đồng phức tạp thành nhiều bước: (1) Nhận diện loại hợp đồng và điều khoản trọng yếu; (2) Truy vấn các luật rủi ro liên quan; "
        "(3) Chạy phân tích rủi ro từng điều khoản; (4) Tự kiểm duyệt và đối soát chéo kết quả (Self-Verification) để loại bỏ hiện tượng ảo giác (hallucination)."
    )

    add_bullet(
        "Ứng dụng Context Compression (Nén ngữ cảnh thông qua LLMLingua): ",
        "Đối với các hợp đồng dài hàng trăm trang, việc đưa toàn bộ văn bản vào prompt sẽ gây quá tải cửa sổ ngữ cảnh và tăng chi phí token. "
        "Hệ thống sẽ áp dụng thuật toán LLMLingua để loại bỏ các từ ngữ dư thừa, các câu từ mang tính chất mô tả chung chung, chỉ giữ lại các từ khóa và mệnh đề "
        "pháp lý cốt lõi, giúp nén prompt đến 40-50% mà không làm giảm độ chính xác của phân tích rủi ro."
    )

    add_bullet(
        "Áp dụng giao thức Model Context Protocol (MCP): ",
        "Chuẩn hóa giao thức kết nối và truyền tải ngữ cảnh giữa AI Service với các nguồn dữ liệu bên ngoài của doanh nghiệp (như cơ sở dữ liệu luật pháp quốc gia) "
        "để AI Agent có thể truy cập thời gian thực các văn bản pháp luật hiện hành và đưa ra cảnh báo rủi ro khớp với luật sửa đổi mới nhất."
    )

    # ------------------ SEMANTIC SEARCH ------------------
    add_heading_2("3.2 Dịch vụ Tìm kiếm Ngữ nghĩa (Semantic Search Service)")
    doc.add_paragraph(
        "Hiện trạng đã triển khai: Hệ thống sử dụng mô hình SentenceTransformers (Multilingual-E5-Base) để chuyển các điều khoản (Clause) thành vector nhúng dense 768 chiều và lưu trữ trong cơ sở dữ liệu vector ChromaDB để thực hiện tìm kiếm tương tự Cosine (Cosine Similarity)."
    )
    doc.add_paragraph(
        "Kế hoạch nâng cấp và phương pháp phát triển:"
    )

    add_bullet(
        "Thiết lập cơ chế Hybrid Search (Tìm kiếm hỗn hợp): ",
        "Bên cạnh việc so khớp ngữ nghĩa bằng dense vector (ChromaDB), hệ thống sẽ tích hợp thêm bộ tìm kiếm từ khóa tần suất (Sparse Vector sử dụng thuật toán BM25). "
        "Hybrid Search kết hợp thế mạnh của cả hai phương pháp: hiểu được ngữ nghĩa ẩn dụ của các khái niệm pháp lý, đồng thời so khớp chính xác tuyệt đối các thuật ngữ chuyên ngành, mã hiệu thiết bị, hoặc số hiệu văn bản pháp luật đặc thù."
    )

    add_bullet(
        "Tích hợp Cross-Encoder Reranking (Xếp hạng lại): ",
        "Sau khi truy vấn được top 20 điều khoản liên quan nhất từ ChromaDB và BM25, hệ thống sẽ đưa các kết quả này qua một mô hình Cross-Encoder Rerank (ví dụ BGE-Reranker). "
        "Mô hình reranker thực hiện đánh giá mối tương quan trực tiếp giữa câu truy vấn và từng tài liệu đích một cách sâu sắc hơn, giúp sắp xếp lại thứ tự ưu tiên của các đoạn văn bản chính xác nhất trước khi gửi vào prompt của LLM."
    )

    add_bullet(
        "Áp dụng Metadata Filtering (Bộ lọc siêu dữ liệu nâng cao): ",
        "Tối ưu hóa tốc độ và độ chính xác của Vector DB bằng cách gắn thêm các trường siêu dữ liệu (metadata) vào các vector điều khoản (ví dụ: ngày ký, giá trị hợp đồng, mã đối tác, trạng thái phê duyệt). "
        "Khi tìm kiếm, hệ thống có thể áp dụng các bộ lọc cứng (ví dụ: 'chỉ tìm kiếm các điều khoản thanh toán của các hợp đồng ký với đối tác TechVibe từ năm 2025') trước khi thực hiện tìm kiếm vector, giúp giảm đáng kể không gian tìm kiếm."
    )

    add_bullet(
        "Xây dựng GraphRAG (Đồ thị tri thức kết hợp RAG): ",
        "Sử dụng cơ sở dữ liệu đồ thị Neo4j để mô hình hóa các thực thể pháp lý (Công ty, Cá nhân, Hợp đồng, Tài sản bàn giao, Nghĩa vụ thanh toán) và mối liên kết giữa chúng. "
        "GraphRAG cho phép AI truy vấn các mối quan hệ sở hữu chéo, các điều khoản ràng buộc chồng chéo giữa nhiều hợp đồng khác nhau của cùng một tập đoàn đối tác, phát hiện xung đột lợi ích vật chất mà tìm kiếm vector thông thường không thể nhận biết."
    )

    # ------------------ AI EXTRACT ------------------
    add_heading_2("3.3 Dịch vụ Trích xuất AI (AI Extract Service)")
    doc.add_paragraph(
        "Hiện trạng đã triển khai: Sử dụng mô hình AI kết hợp với bộ Regex ClauseSplitter dự phòng và các heuristics cơ bản để bóc tách điều khoản (Clause) và trích xuất các thực thể cơ bản (ExtractedEntity như Bên A, Bên B, giá trị, hiệu lực) với độ tin cậy được tính toán sơ bộ."
    )
    doc.add_paragraph(
        "Kế hoạch nâng cấp và phương pháp phát triển:"
    )

    add_bullet(
        "Ứng dụng LLM Tool Calling / Function Calling: ",
        "Cấu hình mô hình AI khả năng tự động sinh mã lệnh gọi các API đối soát thông tin. Ví dụ, khi trích xuất thông tin doanh nghiệp, AI sẽ tự động kích hoạt gọi API của Tổng cục Thuế để xác thực xem mã số thuế trích xuất được có khớp với tên doanh nghiệp đăng ký hay không, tự động sửa lỗi chính tả hoặc cảnh báo nếu thông tin doanh nghiệp giả mạo."
    )

    add_bullet(
        "Xây dựng Automated ETL & Knowledge Sync Pipelines: ",
        "Phát triển hệ thống kết nối tự động (cổng API) kết nối trực tiếp với các kho lưu trữ đám mây dùng chung của doanh nghiệp như OneDrive, SharePoint, Google Drive. "
        "Hệ thống sẽ chạy một tiến trình ngầm (background worker) định kỳ quét các thư mục được chỉ định, tự động tải các hợp đồng mới lên, giải mã AES, OCR bằng PaddleOCR và thực thi trích xuất thực thể đưa vào CSDL mà không cần người dùng thao tác upload thủ công."
    )

    # ------------------ AI SUMMARY ------------------
    add_heading_2("3.4 Dịch vụ Tóm tắt AI (AI Summary Service)")
    doc.add_paragraph(
        "Hiện trạng đã triển khai: Sử dụng mô hình Qwen2.5-3B-Instruct để sinh ra đoạn tóm tắt tiếng Việt (Executive Summary) cho mỗi phiên bản hợp đồng và lưu vào bảng ContractSummary."
    )
    doc.add_paragraph(
        "Kế hoạch nâng cấp và phương pháp phát triển:"
    )

    add_bullet(
        "Áp dụng Semantic Caching (GPTCache): ",
        "Đối với các doanh nghiệp ký kết lượng hợp đồng mẫu (template) cực lớn, nội dung giữa các hợp đồng thường giống nhau tới 90%. "
        "Hệ thống sẽ sử dụng GPTCache để lưu trữ các câu trả lời tóm tắt của các hợp đồng có độ tương tự ngữ nghĩa rất cao. "
        "Khi một hợp đồng mới có cấu trúc tương đương được tải lên, hệ thống sẽ trả về kết quả tóm tắt ngay lập tức từ cache, tiết kiệm chi phí tính toán và giảm thời gian phản hồi xuống còn mili-giây."
    )

    add_bullet(
        "Triển khai cơ chế SSE Streaming (Server-Sent Events): ",
        "Khi tóm tắt các hợp đồng dung lượng lớn, thời gian chờ mô hình suy luận và sinh văn bản có thể mất từ 10-15 giây. "
        "Hệ thống sẽ nâng cấp API tóm tắt sang dạng stream token sử dụng giao thức SSE. Giao diện người dùng sẽ hiển thị chữ chạy thời gian thực (word-by-word) ngay khi mô hình đang sinh câu trả lời, loại bỏ cảm giác chờ đợi mệt mỏi cho người dùng."
    )

    add_bullet(
        "Tích hợp Interactive Summary Memory (Trí nhớ tóm tắt tương tác): ",
        "Thiết lập cơ sở dữ liệu lưu lịch sử hội thoại tóm tắt (sử dụng Redis hoặc PostgreSQL). "
        "Người dùng có thể yêu cầu AI tóm tắt thu gọn hoặc chi tiết hóa theo từng khía cạnh khác nhau như: 'Hãy tóm tắt lại chỉ tập trung vào các mốc thời gian bàn giao' hoặc 'Hãy viết lại tóm tắt này dưới dạng bảng so sánh quyền lợi hai bên'."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 4: TÍCH HỢP HỆ THỐNG PHẦN CỨNG NHÚNG BIÊN (IOT)
    # =========================================================================
    add_heading_1("Chương 4. Kế Hoạch Tích Hợp Hệ Thống Phần Cứng Nhúng Biên (IoT)")
    
    doc.add_paragraph(
        "Để mở rộng khả năng kiểm soát an toàn vật lý của tài liệu gốc tại các doanh nghiệp, "
        "dự án Quản lý Hợp đồng vạch ra lộ trình tích hợp hệ thống phần cứng biên thông minh bao gồm thiết bị quét ảnh, ki-ốt tự ký kết và tủ lưu trữ bảo mật chống đánh tráo."
    )
    
    add_heading_2("4.1 Thiết bị Quét Hợp đồng Thông minh (Smart Contract Scanner)")
    doc.add_paragraph(
        "Thiết bị quét đặt tại văn phòng sử dụng bộ vi điều khiển Raspberry Pi Zero 2W hoặc Raspberry Pi 4 kết hợp với Camera Module V2 (IMX219) "
        "và vòng đèn LED WS2812B trợ sáng. Tiến trình hoạt động bao gồm:\n"
        "1. Người dùng bấm nút vật lý quét kết nối qua chân GPIO 24.\n"
        "2. Thiết bị bật đèn LED ring (GPIO 18) để chiếu sáng đồng đều mặt giấy hợp đồng và kích hoạt camera chụp ảnh.\n"
        "3. Chạy OpenCV Edge Pipeline trực tiếp trên thiết bị biên để định vị tài liệu: Lọc nhiễu Gaussian, phát hiện cạnh Canny, áp dụng Perspective Warp hiệu chỉnh độ nghiêng, và nhị phân hóa thích ứng (Adaptive Thresholding) để tạo ảnh đen trắng rõ nét.\n"
        "4. Tệp ảnh kết quả được mã hóa đối xứng bằng thuật toán AES-256-GCM cục bộ tại RAM của thiết bị, truyền tải an toàn qua HTTPS POST mTLS lên Django Server API."
    )
    
    add_heading_2("4.2 Ki-ốt Tự dựng và Ký kết Hợp đồng (Contract Self-Builder Kiosk)")
    doc.add_paragraph(
        "Trạm ki-ốt đặt tại phòng họp/văn phòng giao dịch giúp đối tác tự dựng văn bản và thực thi ký số bảo mật phần cứng:\n"
        "1. Đối tác xác thực danh tính 2 yếu tố (2FA) bằng cách quẹt thẻ RFID/NFC PN532 (giao tiếp SPI/I2C) và quét vân tay trên module AS608 (UART Serial).\n"
        "2. Dựng hợp đồng PDF từ template HTML thông qua WeasyPrint và Jinja2 dựa trên thông tin đã điền.\n"
        "3. Ký số ECDSA phần cứng: Trích xuất mã băm SHA-256 của hợp đồng, truyền qua giao tiếp I2C vào chip bảo mật phần cứng chuyên dụng ATECC608A. Chip thực hiện tạo chữ ký số nội bộ bằng khóa riêng (Private Key) được ghi vĩnh viễn trong phần cứng bảo mật mà không bao giờ lộ khóa ra ngoài phần mềm.\n"
        "4. Đồng bộ dữ liệu lên Django Backend và sổ cái Blockchain Hyperledger Fabric, đồng thời in hóa đơn nhiệt chứa QR Code đối soát Blockchain."
    )
    
    add_heading_2("4.3 Tủ Hồ Sơ Hợp Đồng Thông Minh Chống Đánh Tráo (Smart Cabinet)")
    doc.add_paragraph(
        "Hệ thống tủ bảo mật giám sát và quản lý các hợp đồng giấy bản gốc vật lý chống đánh tráo:"
    )
    
    add_bullet(
        "Kiểm soát quyền truy cập mở tủ dựa trên phân quyền (Role-Based Access Control): ",
        "Tủ được trang bị cảm biến vân tay AS608 kết nối với Raspberry Pi điều khiển. Khi người dùng quét vân tay, hệ thống nhúng gửi mã vân tay về Django Server. "
        "Django kiểm tra phân quyền tài khoản (User Role) từ CSDL. Chỉ các tài khoản có vai trò được ủy quyền đặc biệt (MANAGER hoặc LEGAL_OFFICER) mới có quyền mở tủ. "
        "Khi được phê duyệt, Django Server gửi lệnh điều khiển kích hoạt chốt khóa điện từ Solenoid qua GPIO 26 -> UNLOCKED."
    )

    add_bullet(
        "Ghi nhận nhật ký kiểm toán mở/đóng tủ tự động: ",
        "Tủ được trang bị cảm biến hành trình cơ học (hoặc cảm biến từ Reed Switch) ở cánh cửa tủ kết nối GPIO để giám sát trạng thái đóng mở cửa. "
        "Hệ thống tự động ghi nhận chính xác thời gian cửa tủ được mở (opened_at), thời gian cửa đóng lại (closed_at), người thực hiện mở và trạng thái khóa vào bảng CabinetAccessLog trên CSDL."
    )

    add_bullet(
        "Nhận biết và phòng chống đánh tráo hợp đồng (Anti-Tampering & Anti-Swapping): ",
        "Trên bìa của mỗi tập hồ sơ hợp đồng giấy gốc được dán một thẻ RFID/NFC chứa một mã nhận diện UID duy nhất. "
        "Trong hộc tủ lưu trữ được trang bị đầu đọc RFID PN532. Khi người dùng mở tủ và đặt/rút hồ sơ, đầu đọc PN532 sẽ quét mã UID của thẻ RFID này.\n"
        "Thiết bị biên gửi mã UID thẻ RFID này lên máy chủ đối soát chéo với mã băm toàn vẹn HashProof đã neo giữ trên sổ cái Hyperledger Fabric Blockchain.\n"
        "Nếu mã băm hợp đồng hiện tại trùng khớp hoàn toàn với thông tin thẻ đã đăng ký trước đó trên Blockchain -> Xác nhận hồ sơ nguyên bản không bị đánh tráo (VALID).\n"
        "Nếu phát hiện mã thẻ lạ, không tồn tại hoặc mã băm không trùng khớp -> Hệ thống cảnh báo hành vi tráo đổi hồ sơ trái phép (INVALID), lập tức kích hoạt còi báo động GPIO 25 và gửi thông báo đỏ khẩn cấp đến Quản trị viên."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 5: CÔNG CỤ GIẢ LẬP PHẦN CỨNG VÀ KẾ HOẠCH VẬN HÀNH (SIL)
    # =========================================================================
    add_heading_1("Chương 5. Thiết Lập Các Công Cụ Giả Lập và Kế Hoạch Thử Nghiệm (SIL)")
    
    doc.add_paragraph(
        "Để phát triển và kiểm thử độc lập mà không phụ thuộc vào thiết bị phần cứng vật lý, phân hệ nhúng "
        "được cấu hình chạy giả lập hoàn toàn dưới dạng Software-in-the-Loop (SIL) sử dụng các công cụ ảo hóa chuyên dụng:"
    )
    
    add_heading_2("5.1 Các công cụ giả lập phần cứng được áp dụng")
    
    add_bullet(
        "Bộ giả lập hệ điều hành và CPU QEMU (Quick Emulator): ",
        "Cấu hình giả lập phần cứng ARMv8 của Raspberry Pi 4 để chạy hệ điều hành Raspberry Pi OS (Debian) ảo hóa trên máy tính xách tay x86 chạy Windows. "
        "Điều này giúp chạy thử nghiệm các tiến trình Python nhúng trong một hệ thống phân quyền, quản lý tiến trình giống hệt thiết bị thực tế."
    )

    add_bullet(
        "Bộ mô phỏng mạch điện tử Proteus VSM & Wokwi: ",
        "Vẽ sơ đồ nguyên lý mạch điện tử và giả lập thời gian thực các bus truyền thông. "
        "Công cụ mô phỏng tín hiệu điện áp chân GPIO, phản hồi thanh ghi của đầu đọc NFC PN532 thông qua bus I2C/SPI, tập lệnh Hex xác thực vân tay của module AS608 qua UART Serial, và cơ chế đóng mở của chốt Solenoid 12V qua Relay."
    )

    add_bullet(
        "Thư viện ảo hóa Driver trong mã nguồn Python (mock_hardware.py): ",
        "Đảm nhận vai trò đánh chặn (intercept) toàn bộ các lệnh truy cập cổng phần cứng vật lý cấp thấp (RPi.GPIO, smbus2, serial). "
        "Lớp này mô phỏng các trạng thái phản hồi nhị phân ngẫu nhiên, sinh mã băm SHA-256 ảo và chữ ký ECDSA tuân thủ cấu trúc của chip bảo mật ATECC608A."
    )

    add_heading_2("5.2 Kế hoạch chi tiết vận hành và thử nghiệm với bộ giả lập (Simulation Plan)")
    doc.add_paragraph(
        "Kế hoạch kiểm thử và tích hợp được triển khai chi tiết qua 4 giai đoạn khép kín:\n"
        "- Giai đoạn 1. Khởi tạo môi trường ảo QEMU: Xây dựng tệp ảnh đĩa cứng Raspberry Pi OS Lite, cấu hình chia sẻ cổng mạng (Port Forwarding) từ QEMU (cổng 8000) đến máy chủ Windows chạy Django Backend để thiết lập kênh mTLS an toàn.\n"
        "- Giai đoạn 2. Giả lập phần cứng ngoại vi và đo đạc xung nhịp: Trên Proteus VSM/Wokwi, chạy các kịch bản kiểm định thời gian trễ (latency) của bus I2C/SPI khi truyền tải gói dữ liệu lớn và kịch bản sụt áp nguồn cấp của chốt khóa Solenoid khi kích hoạt mở tủ.\n"
        "- Giai đoạn 3. Thực thi kịch bản kiểm thử luồng tích hợp Cabinet & Kiosk: Khởi chạy cabinet_simulator.py và kiosk_app.py trong QEMU, thực hiện các ca kiểm thử giả lập: (a) Quét vân tay hợp lệ của MANAGER/LEGAL_OFFICER -> kích hoạt GPIO 26 mở khóa; (b) Quét vân tay không hợp lệ -> từ chối và ghi log; (c) Quét thẻ RFID hợp lệ -> xác minh toàn vẹn Blockchain; (d) Quét thẻ giả mạo -> còi hú báo động GPIO 25 và gửi cảnh báo đỏ.\n"
        "- Giai đoạn 4. Kiểm duyệt nhật ký kiểm toán và đối soát: Trích xuất tệp cabinet_audit_log.json ra khỏi QEMU, kiểm duyệt cấu trúc JSON và dữ liệu thời gian mở tủ (opened_at) / đóng tủ (closed_at) để đảm bảo đồng bộ hoàn chỉnh trên Django Admin."
    )

    # Save document
    output_path = r"d:\Django_project\RiskDL\Bao_Cao_Ke_Hoach_Tuong_Lai_Quan_Ly_Hop_Dong.docx"
    try:
        doc.save(output_path)
        print(f"Future Plan Text Report successfully saved to: {output_path}")
    except PermissionError:
        output_path = r"d:\Django_project\RiskDL\Bao_Cao_Ke_Hoach_Tuong_Lai_Quan_Ly_Hop_Dong_new.docx"
        doc.save(output_path)
        print(f"Future Plan Text Report successfully saved to: {output_path} (original was locked)")

if __name__ == "__main__":
    create_future_report()
