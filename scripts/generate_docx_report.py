import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color definitions
    c_primary = RGBColor(0, 51, 102)     # Deep Blue
    c_secondary = RGBColor(70, 130, 180) # Steel Blue
    c_text = RGBColor(51, 51, 51)        # Off-black
    c_gray = RGBColor(128, 128, 128)     # Gray
    
    # Configure default style font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = c_text
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    def add_custom_heading(text, level, space_before=12, space_after=6):
        h = doc.add_heading('', level=level)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        
        run = h.add_run(text)
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = c_primary
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = c_secondary
        elif level == 3:
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = c_gray
        return h

    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr.append(borders)

    def add_code_block(code_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        
        cell = table.cell(0, 0)
        set_cell_background(cell, "F8F8F8")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        set_table_borders(table, color="D3D3D3", sz="6")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        run = p.add_run(code_text.strip())
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(45, 45, 45)
        
        # Add spacing after table
        empty_p = doc.add_paragraph()
        empty_p.paragraph_format.space_before = Pt(4)
        empty_p.paragraph_format.space_after = Pt(4)
        empty_p.paragraph_format.line_spacing = 1.0

    def add_explanation_bullet(code_line, explanation):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        
        run_code = p.add_run(code_line)
        run_code.font.name = 'Consolas'
        run_code.font.size = Pt(9.0)
        run_code.font.bold = True
        run_code.font.color.rgb = RGBColor(200, 40, 40)
        
        p.add_run(" : ")
        
        run_exp = p.add_run(explanation)
        run_exp.font.name = 'Calibri'
        run_exp.font.size = Pt(10.5)

    def add_detailed_bullet(title, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.bold = True
        p.add_run(" " + text)

    # --- TITLE PAGE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("BÁO CÁO KỸ THUẬT & TOÀN BỘ MÃ NGUỒN HỆ THỐNG AI")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = c_primary

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(120)
    run_sub = p_sub.add_run("Phân Tích Nghiệp Vụ, Ý Nghĩa Kỹ Thuật, Ưu Điểm, Bất Lợi, Vấn Đề Fine-tune Hiện Tại, Hướng Phát Triển Tương Lai Để Cải Thiện Kết Quả Trả Ra Và Từng Dòng Lệnh Mã Nguồn Cho 2 Mô Hình AI Trực Thuộc RiskDL:\n1. Mô hình trích xuất thông tin hợp đồng (Qwen2.5-3B-Instruct QLoRA)\n2. Mô hình tìm kiếm nội dung hợp đồng (Multilingual-E5-Base MNRL, Gradio, HfApi)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = c_secondary

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(0)
    run_meta = p_meta.add_run("Hệ thống quản lý hợp đồng thông minh RiskDL\nBản cập nhật đầy đủ chức năng lọc dữ liệu, huấn luyện và suy luận")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = c_gray

    doc.add_page_break()

    # --- SECTION 1: SYSTEM OVERVIEW (BUSINESS MEANING, PROS & CONS) ---
    add_custom_heading("1. TỔNG QUAN HỆ THỐNG AI TRÊN RiskDL (RiskDL AI SYSTEM OVERVIEW)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống AI cốt lõi hỗ trợ nền tảng quản trị rủi ro pháp lý RiskDL được cấu thành từ hai mô hình chuyên biệt: Mô hình trích xuất thông tin hợp đồng (Qwen2.5-3B-Instruct) và Mô hình tìm kiếm nội dung hợp đồng (Multilingual-E5-Base). Dưới đây là phân tích chi tiết về vai trò nghiệp vụ, ưu thế kỹ thuật cũng như những hạn chế tồn tại của từng giải pháp.")

    add_custom_heading("1.1 Mô hình Trích xuất Thông tin Hợp đồng (Qwen2.5-3B-Instruct)", level=2)
    
    p = doc.add_paragraph()
    p.add_run("• Ý nghĩa và Công dụng nghiệp vụ:")
    p = doc.add_paragraph("Mô hình đóng vai trò là một 'Chuyên gia pháp lý ảo' tự động hóa. Khi tiếp nhận tài liệu hợp đồng, mô hình thực hiện bóc tách cấu trúc, đọc hiểu ngữ nghĩa sâu sắc của từng điều khoản con để tự động phân nhóm điều khoản vào 8 danh mục pháp lý trọng điểm (thanh toán, lương thưởng, bảo mật thông tin, giới hạn trách nhiệm, thời hạn và chấm dứt hợp đồng, bảo hiểm, giải quyết tranh chấp và nhóm mặc định khác). Song song với đó, mô hình thực hiện trích xuất thông tin hợp đồng và tóm tắt chi tiết các nghĩa vụ pháp lý ràng buộc hoặc các điều khoản bất lợi tiềm ẩn dưới định dạng dữ liệu có cấu trúc JSON để đưa thẳng vào hệ thống quản lý cơ sở dữ liệu. Việc này giúp doanh nghiệp giảm thiểu tới 70% thời gian rà soát thủ công và triệt tiêu rủi ro bỏ sót điều khoản.")
    
    p = doc.add_paragraph()
    p.add_run("• Ưu điểm và Lợi ích kỹ thuật (Pros):")
    add_detailed_bullet("Tối ưu hóa tài nguyên phần cứng vượt trội:", "Việc áp dụng kỹ thuật QLoRA (lượng tử hóa 4-bit với phân phối NF4 và lượng tử hóa kép Double Quantization) giúp nén dung lượng mô hình trên VRAM từ ~6GB xuống chỉ còn ~1.55GB. Điều này cho phép doanh nghiệp tự vận hành (self-host) hoàn toàn mô hình trên một GPU giá rẻ duy nhất (như NVIDIA T4 hoặc P100) mà không cần đầu tư máy chủ đắt đỏ.")
    add_detailed_bullet("Bảo mật thông tin và dữ liệu thương mại tuyệt đối:", "Do mô hình chạy cục bộ (on-premise/local) trong hạ tầng của RiskDL, toàn bộ dữ liệu hợp đồng thương mại nhạy cảm và điều khoản bảo mật của khách hàng không bị truyền ra ngoài internet hoặc gửi lên các dịch vụ API bên thứ ba, tránh hoàn toàn các rủi ro rò rỉ thông tin pháp lý.")
    add_detailed_bullet("Hội tụ huấn luyện tối ưu:", "Nhờ đắp các adapters thích ứng thứ hạng thấp (LoRA với rank r=16 và alpha=32), mô hình học được các thuật ngữ pháp lý và phong cách viết hợp đồng mới rất nhanh chóng mà không gặp hiện tượng quên lãng thảm họa (catastrophic forgetting).")

    p = doc.add_paragraph()
    p.add_run("• Nhược điểm, Bất lợi và Hạn chế (Cons):")
    add_detailed_bullet("Độ trễ giải nén lượng tử hóa (Dequantization Overhead):", "Khi suy luận (inference), các trọng số lượng tử hóa 4-bit cần được giải nén liên tục về kiểu float16 để thực hiện tính toán ma trận. Quá trình dequantization liên tục này tạo ra một khoản chi phí tính toán phụ trợ (computational overhead), có thể làm tăng nhẹ độ trễ tạo token (generation latency) trên mỗi ký tự so với mô hình FP16 nguyên bản nếu không sử dụng các kernel tối ưu đặc thù.")
    add_detailed_bullet("Giới hạn dung lượng tham số (3B Parameters):", "Để đáp ứng yêu cầu tiết kiệm bộ nhớ, mô hình được chọn có kích thước 3 tỷ tham số. Kích thước này hạn chế khả năng lập luận logic đa bước đối với các hợp đồng có cấu trúc câu lồng nhau quá phức tạp, có thể dẫn đến việc tóm tắt thiếu ý hoặc hiểu sai lệch nhẹ cấu trúc câu phủ định kép trong văn bản luật.")
    add_detailed_bullet("Tỷ lệ sinh ảo (Hallucination):", "Tương tự như mọi mô hình ngôn ngữ lớn khác, mô hình vẫn có một tỷ lệ nhỏ sinh ra thông tin ảo (bịa đặt thông tin hoặc diễn giải sai lệch so với văn bản gốc) nếu câu lệnh prompt đầu vào không được thiết kế chặt chẽ và không có cơ chế đối chiếu hậu xử lý.")

    add_custom_heading("1.2 Mô hình Tìm kiếm Nội dung Hợp đồng (Multilingual-E5-Base)", level=2)
    
    p = doc.add_paragraph()
    p.add_run("• Ý nghĩa và Công dụng nghiệp vụ:")
    p = doc.add_paragraph("Mô hình chịu trách nhiệm chuyển đổi toàn bộ các đoạn văn bản (chunks) của hợp đồng sang không gian vector biểu diễn đặc trưng (dense embeddings) 768 chiều để lập chỉ mục (indexing) vào cơ sở dữ liệu Vector Database ChromaDB. Công dụng cốt lõi là cung cấp giải pháp tìm kiếm nội dung hợp đồng song ngữ và chéo ngôn ngữ (Cross-lingual Search) vượt trội. Người dùng có thể sử dụng các câu truy vấn bằng tiếng Việt tự nhiên để tìm ra chính xác các điều khoản tương đương viết bằng tiếng Anh trong hợp đồng đối tác và ngược lại.")
    
    p = doc.add_paragraph()
    p.add_run("• Ưu điểm và Lợi ích kỹ thuật (Pros):")
    add_detailed_bullet("Khả năng hiểu ngữ nghĩa vượt trội:", "Khác biệt hoàn toàn với công cụ tìm kiếm từ khóa truyền thống (như BM25 của Elasticsearch dễ bỏ sót từ đồng nghĩa hoặc cách diễn đạt khác), E5 ánh xạ câu chữ dựa trên mối quan hệ ngữ nghĩa bản chất, cho phép tìm kiếm chính xác các khái niệm pháp lý tương tự ngay cả khi không có bất kỳ từ khóa nào trùng khớp trực tiếp.")
    add_detailed_bullet("Tốc độ phản hồi thời gian thực siêu tốc:", "Với kích thước nhỏ gọn chỉ khoảng 278 triệu tham số, mô hình thực hiện encode và truy vấn vector trong chưa đầy 10 mili-giây trên CPU hoặc GPU tầm trung, tiêu thụ cực ít tài nguyên phần cứng.")
    add_detailed_bullet("Độ tương thích đa ngôn ngữ cao:", "Được huấn luyện sẵn trên kho dữ liệu khổng lồ gồm hơn 94 ngôn ngữ, mô hình có khả năng căn chỉnh không gian vector rất tốt giữa tiếng Anh và tiếng Việt, tạo điều kiện thuận lợi cho việc xử lý các giao dịch thương mại quốc tế.")

    p = doc.add_paragraph()
    p.add_run("• Nhược điểm, Bất lợi và Hạn chế (Cons):")
    add_detailed_bullet("Nhiễu do dịch thuật máy tự động trong dữ liệu train:", "Do tập dữ liệu huấn luyện CUAD gốc được viết hoàn toàn bằng tiếng Anh, quy trình xây dựng tập huấn luyện song ngữ phải phụ thuộc vào công cụ dịch máy tự động (Google Translator). Sự thiếu chính xác của dịch máy khi xử lý các thuật ngữ pháp lý chuyên ngành phức tạp (như 'indemnification', 'severability') vô hình trung tạo ra các nhãn nhiễu (noisy labels), làm giảm độ sắc bén của không gian vector khi biểu diễn các truy vấn tiếng Việt chuyên sâu.")
    add_detailed_bullet("Hạn chế về kích thước Batch Size khi huấn luyện:", "Hàm mất mát MultipleNegativesRankingLoss (MNRL) phụ thuộc rất mạnh vào kích thước Batch Size lớn để cung cấp nhiều mẫu phủ định chéo (in-batch negatives). Giới hạn phần cứng của môi trường thử nghiệm bắt buộc phải huấn luyện với batch_size=8, điều này làm giảm bớt khả năng phân biệt tinh tế của mô hình đối với các điều khoản có nội dung gần giống nhau nhưng mang ý nghĩa pháp lý khác nhau (hard negatives).")
    add_detailed_bullet("Cắt rời ngữ cảnh do Phân đoạn (Chunking limitations):", "Hàm phân đoạn theo điều khoản (chunk_by_section) chia cắt văn bản hợp đồng thành nhiều đoạn độc lập. Nếu một câu hỏi đòi hỏi thông tin tổng hợp chéo giữa nhiều điều khoản nằm xa nhau (ví dụ: liên hệ giữa thời hạn thanh toán ở điều 5 và lãi suất quá hạn ở điều 9), mô hình truy xuất vector đơn lẻ sẽ không thể trả về kết quả hoàn chỉnh mà chỉ trích xuất được từng đoạn rời rạc.")

    # --- NEW SECTION 1.3: PROBLEMS IN CURRENT FINETUNING APPROACH ---
    add_custom_heading("1.3 Phân tích Vấn đề và Hạn chế của Phương pháp Fine-tune Hiện tại", level=2)
    p = doc.add_paragraph()
    p.add_run("Mặc dù phương pháp huấn luyện hiện tại mang lại kết quả bước đầu khả quan, việc đưa hệ thống vào vận hành thực tế vẫn vấp phải một số vấn đề kỹ thuật cố hữu sau:")
    
    add_detailed_bullet("Sai số do Lượng tử hóa tích lũy (Quantization Loss):", "Đối với Model Trích xuất Thông tin Hợp đồng, việc nén cứng các trọng số gốc về 4-bit (QLoRA) dẫn đến việc mô hình mất đi tính liên tục trong không gian phân phối số thực. Sai số lượng tử hóa tích lũy qua 32 lớp Transformer của Qwen có thể làm suy giảm nhẹ khả năng suy luận logic đa bước khi phải đọc các hợp đồng dịch thuật có cấu trúc câu lồng nhau quá phức tạp.")
    add_detailed_bullet("Lệch phân phối độ dài hội thoại (Sequence Length Mismatch):", "Tập dữ liệu huấn luyện được giới hạn ở max_seq_length=512 tokens để tiết kiệm VRAM. Trong thực tế, các văn bản hợp đồng pháp lý của doanh nghiệp thường có độ dài từ 2,000 đến 10,000 tokens. Sự chênh lệch lớn này khiến mô hình dễ bị mất phương hướng (context window confusion), dẫn đến hiện tượng trích xuất thiếu sót thông tin khi xử lý các điều khoản nằm ở cuối văn bản siêu dài.")
    add_detailed_bullet("Hiện tượng Nhiễu dữ liệu Nhãn do Dịch máy tự động:", "Với Model Tìm kiếm Nội dung Hợp đồng, việc dùng công cụ dịch máy tự động (Google Translator) để tạo dữ liệu huấn luyện song ngữ Anh-Việt mang lại nguồn dữ liệu lớn nhưng đi kèm tỷ lệ nhiễu cao. Nhiều thuật ngữ pháp lý chuyên ngành mang tính bảo thủ (như 'indemnification' dịch thành 'bồi thường thiệt hại' nhưng đôi khi bị dịch sai thành 'sự miễn trừ') làm phân tán sự căn chỉnh các vector ngôn ngữ, khiến độ chính xác tìm kiếm chéo bị giảm sút đối với các điều khoản đặc thù.")
    add_detailed_bullet("Hạn chế của Hàm mất mát MNRL với Batch Size nhỏ:", "Hàm loss MultipleNegativesRankingLoss phụ thuộc trực tiếp vào kích thước Batch Size để tạo ra các mẫu phủ định ngẫu nhiên (in-batch negatives). Do giới hạn phần cứng, batch size được đặt ở mức 8 (chỉ có 7 mẫu phủ định chéo). Số lượng mẫu phủ định quá ít khiến mô hình dễ gặp hiện tượng 'dễ dãi', chỉ cần phân biệt các đoạn văn hoàn toàn khác nhau mà chưa học được cách phân biệt các đoạn văn có từ khóa cực kỳ giống nhau nhưng mang nghĩa pháp lý trái ngược (hard negatives).")

    doc.add_page_break()

    # --- SECTION 2 ---
    add_custom_heading("2. CHI TIẾT KỸ THUẬT MÔ HÌNH TRÍCH XUẤT THÔNG TIN HỢP ĐỒNG (QWEN2.5-3B-INSTRUCT)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Phần này trình bày chi tiết về mã nguồn xử lý dữ liệu đầu vào, cấu hình lượng tử hóa, thiết lập ma trận adapters và tham số tối ưu hóa huấn luyện cho mô hình Qwen.")

    add_custom_heading("2.1 Quy trình Lọc dữ liệu, Ánh xạ Danh mục & Cân bằng mẫu", level=2)
    p = doc.add_paragraph("Để giải quyết sự mất cân bằng dữ liệu nặng nề trong bộ dữ liệu CUAD và loại bỏ các mẫu không chứa thông tin hữu ích, chúng tôi áp dụng quy trình lọc sau:")

    add_custom_heading("A. Bộ quy tắc ánh xạ câu hỏi sang 8 danh mục nghiệp vụ chính", level=3)
    p = doc.add_paragraph("Ánh xạ từ 41 câu hỏi chi tiết của CUAD về 8 danh mục tổng quát bằng kỹ thuật Keyword Matching:")
    
    code_filter_map = """
CATEGORY_KEYWORDS = {
    "salary": ["salary", "wage", "compensation", "remuneration", "bonus", "annual base salary"],
    "payment": ["payment", "fee", "price", "revenue share", "royalty", "pricing", "invoice", "cost"],
    "confidentiality": ["confidential", "non-disclosure", "nda", "proprietary", "trade secret"],
    "liability": ["liability", "indemnif", "limitation of liability", "damages", "cap on liability"],
    "termination": ["terminat", "expir", "cancel", "renewal", "notice period", "end date"],
    "insurance": ["insurance", "coverage", "insurer", "policy"],
    "dispute_resolution": ["dispute", "arbitration", "governing law", "jurisdiction", "litigation"],
}

def map_question_to_category(question: str) -> str:
    q = question.lower()
    for category, keywords in CATEGORY_KEYWORDS.items():
        if any(kw in q for kw in keywords):
            return category
    return "other"
    """
    add_code_block(code_filter_map)
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("CATEGORY_KEYWORDS = {...}", "Thiết lập các nhóm từ khóa pháp lý. Giúp mô hình định hướng tốt hơn khi gom cụm các khái niệm ngôn ngữ đồng dạng.")
    add_explanation_bullet("map_question_to_category(question)", "Phân loại câu hỏi thô từ CUAD về các nhãn nghiệp vụ rõ ràng, giúp chuẩn hóa nhãn đích (target label) cho mô hình học tập.")
    add_explanation_bullet('return "other"', "Ánh xạ các câu hỏi nằm ngoài danh sách đặc thù về lớp 'other', tránh làm nhiễu loạn không gian đặc trưng của các lớp quan trọng khác.")

    add_custom_heading("B. Thuật toán lọc mẫu trống và chuyển đổi định dạng hội thoại", level=3)
    p = doc.add_paragraph("Mã nguồn lọc các dòng dữ liệu không có nhãn hoặc rỗng ngữ cảnh, đồng thời chuyển đổi dữ liệu sang định dạng Qwen Chat Template (ChatML):")
    
    code_filter_convert = """
def convert_cuad(dataset_split):
    converted, skipped = [], 0
    for sample in dataset_split:
        context  = sample.get("context", "").strip()
        question = sample.get("question", "").strip()
        answers  = sample.get("answers", {})
        texts    = answers.get("text", [])

        # BƯỚC LỌC QUAN TRỌNG: Loại bỏ mẫu không có câu trả lời hoặc context rỗng
        if not texts or not context:
            skipped += 1
            continue

        category = map_question_to_category(question)
        summary  = make_summary(texts[0], context)
        text     = build_chat_text(context, category, summary)

        converted.append({"text": text, "category": category})
    return converted, skipped
    """
    add_code_block(code_filter_convert)
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("sample.get(...)", "Trích xuất thông tin thô. Đảm bảo mô hình được tiếp nhận đầy đủ 3 thành phần ngữ cảnh (context), câu hỏi (question) và nhãn đúng (answers).")
    add_explanation_bullet("if not texts or not context:", "Bộ lọc mẫu dữ liệu trống. Ý nghĩa đối với mô hình: Ngăn chặn mô hình tiếp thu các mẫu không có thông tin phản hồi dương tính, tránh việc mô hình học cách sinh ra các văn bản rỗng hoặc dự đoán sai lệch khi gặp ngữ cảnh dài.")
    add_explanation_bullet("skipped += 1", "Ghi nhận số lượng mẫu nhiễu bị loại bỏ, giúp theo dõi chất lượng phân phối dữ liệu đầu vào trước khi cấp cho mô hình học.")
    add_explanation_bullet("make_summary(texts[0], context)", "Tạo nội dung tóm tắt rủi ro làm dữ liệu nhãn chuẩn (ground truth) để mô hình học cách tóm lược thông tin chính xác.")
    add_explanation_bullet("build_chat_text(...)", "Đóng gói dữ liệu theo cấu trúc hội thoại ChatML (<|im_start|>system... user... assistant...<|im_end|>). Ý nghĩa đối với mô hình: Giúp mô hình phân biệt rõ ràng đâu là chỉ thị hệ thống, đâu là ngữ cảnh đầu vào của người dùng và đâu là định dạng JSON đầu ra cần trích xuất.")

    add_custom_heading("C. Giới hạn số lượng mẫu (Capping) để cân bằng Class", level=3)
    p = doc.add_paragraph("Do nhóm 'other' chiếm số lượng áp đảo dẫn đến mô hình bị lệch hướng học tập, mã nguồn thực hiện gom nhóm và giới hạn mẫu tối đa cho mỗi danh mục:")
    
    code_cap = """
MAX_PER_CLASS = 400

# Nhóm dữ liệu theo danh mục
by_cat = {}
for item in all_data:
    by_cat.setdefault(item["category"], []).append(item)

# Cân bằng dữ liệu (Capping & Shuffle)
balanced = []
for cat in sorted(by_cat):
    items = by_cat[cat]
    random.shuffle(items)
    selected = items[:MAX_PER_CLASS] # Lấy tối đa 400 mẫu cho mỗi class
    balanced.extend(selected)

random.shuffle(balanced)
    """
    add_code_block(code_cap)
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("MAX_PER_CLASS = 400", "Ngưỡng khống chế mẫu tối đa. Ý nghĩa đối với mô hình: Ngăn chặn hiện tượng thiên vị lớp (majority class bias) - nơi mô hình có xu hướng dự đoán mọi điều khoản đều thuộc nhóm 'other' do tần suất xuất hiện của nhóm này quá lớn trong dữ liệu huấn luyện.")
    add_explanation_bullet("by_cat.setdefault(...).append(item)", "Gom dữ liệu theo nhóm nhãn để chuẩn bị cho quá trình lấy mẫu cân bằng.")
    add_explanation_bullet("random.shuffle(items)", "Xáo trộn ngẫu nhiên dữ liệu trước khi cắt giảm, đảm bảo mô hình được học từ một tập con mang tính đại diện cao và khách quan nhất.")
    add_explanation_bullet("selected = items[:MAX_PER_CLASS]", "Cắt giảm số lượng mẫu (undersampling). Ý nghĩa đối với mô hình: Ép phân phối của các nhãn trong tập huấn luyện trở về trạng thái phân phối đều (uniform distribution), buộc mô hình phải tập trung học các đặc trưng ngữ nghĩa riêng biệt của từng lớp pháp lý đặc thù.")
    add_explanation_bullet("balanced.extend(selected)", "Tích lũy các lớp đã cân bằng thành tập dữ liệu huấn luyện thống nhất.")
    add_explanation_bullet("random.shuffle(balanced)", "Xáo trộn chéo toàn bộ dữ liệu huấn luyện cuối cùng. Ý nghĩa đối với mô hình: Tránh việc mô hình học theo chuỗi (sequence bias) - ví dụ học liên tục 400 mẫu thanh toán rồi mới học bảo mật - điều này có thể làm mô hình bị quên cục bộ các lớp học trước đó.")

    doc.add_page_break()

    # --- REST OF THE DOCUMENT ---
    add_custom_heading("2.2 Cấu hình lượng tử hóa BitsAndBytesConfig", level=2)
    p = doc.add_paragraph("Mã nguồn cấu hình nén mô hình để chạy trên GPU NVIDIA Tesla T4 hoặc P100 (tiết kiệm VRAM tối đa):")
    
    code_bnb = """
bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.float16,
    bnb_4bit_use_double_quant=True,
)
    """
    add_code_block(code_bnb)
    
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("load_in_4bit=True", "Kích hoạt lượng tử hóa 4-bit. Ý nghĩa đối với mô hình: Nén toàn bộ ma trận trọng số (weight matrices) của mô hình từ định dạng Float16 xuống 4-bit. Điều này trực tiếp làm thô không gian tham số của mô hình để tiết kiệm VRAM từ ~6GB xuống ~1.55GB, cho phép nạp được mô hình 3B lên các GPU phổ thông.")
    add_explanation_bullet('bnb_4bit_quant_type="nf4"', "Sử dụng kiểu lượng tử hóa NormalFloat 4 (NF4). Ý nghĩa đối với mô hình: Đây là kiểu phân phối số tối ưu hóa toán học riêng cho các trọng số mạng neural tuân theo phân phối chuẩn (Gauss). Đối với mô hình, NF4 giúp giảm thiểu sai số lượng tử hóa (quantization error) xuống mức thấp nhất, bảo toàn khả năng lập luận ngôn ngữ tốt hơn rất nhiều so với định dạng FP4 thông thường.")
    add_explanation_bullet("bnb_4bit_compute_dtype=torch.float16", "Đặt kiểu dữ liệu tính toán là Float16. Ý nghĩa đối với mô hình: Khi mô hình thực hiện phép nhân ma trận lúc lan truyền xuôi và lan truyền ngược, các trọng số 4-bit sẽ được giải nén tạm thời (dequantize) sang Float16 để thực hiện tính toán trên Tensor Cores, giúp mô hình giữ được độ chính xác số học và tránh hiện tượng tràn số.")
    add_explanation_bullet("bnb_4bit_use_double_quant=True", "Kích hoạt lượng tử hóa kép. Ý nghĩa đối với mô hình: Lượng tử hóa tiếp tục các tham số tỷ lệ lượng tử hóa (quantization constants) từ 32-bit xuống 8-bit. Giúp mô hình tiết kiệm thêm trung bình 0.4 bit cho mỗi tham số mà không gây ảnh hưởng tiêu cực đến chất lượng học tập của trọng số.")

    add_custom_heading("2.3 Thiết lập cấu hình thích ứng thứ hạng thấp LoraConfig", level=2)
    p = doc.add_paragraph("Cấu hình thiết lập ma trận adapters LoRA đắp lên các lớp tuyến tính của Qwen2.5:")
    
    code_lora_cfg = """
lora_config = LoraConfig(
    task_type=TaskType.CAUSAL_LM,
    r=16,
    lora_alpha=32,
    lora_dropout=0.05,
    bias="none",
    target_modules=[
        "q_proj", "k_proj", "v_proj", "o_proj",
        "gate_proj", "up_proj", "down_proj",
    ],
)
    """
    add_code_block(code_lora_cfg)
    
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("task_type=TaskType.CAUSAL_LM", "Thiết lập nhiệm vụ Causal Language Modeling. Ý nghĩa đối với mô hình: Định hình cấu trúc tính toán mất mát ở đầu ra của mô hình là hàm dự đoán token tiếp theo tự hồi quy (autoregressive loss), phù hợp với mô hình dạng Decoder-only.")
    add_explanation_bullet("r=16", "Rank của adapter LoRA. Ý nghĩa đối với mô hình: Xác định thứ hạng (chiều rộng) của các ma trận phân rã thích ứng thứ hạng thấp bổ sung. r=16 nghĩa là mô hình chỉ được phép học các hướng cập nhật trọng số quan trọng nhất trong một không gian con 16 chiều. Điều này hạn chế tối đa số tham số huấn luyện mới, giúp mô hình học nhanh và tránh hiện tượng 'quên lãng thảm họa' (catastrophic forgetting) các tri thức nền tảng đã có.")
    add_explanation_bullet("lora_alpha=32", "Hệ số tỷ lệ LoRA. Ý nghĩa đối với mô hình: Hệ số phóng đại dùng để điều chỉnh sức ảnh hưởng của trọng số adapter khi cộng vào trọng số mô hình gốc. Công thức cập nhật là: Trọng_số_mới = Trọng_số_gốc + (lora_alpha / r) * Trọng_số_LoRA. Ở đây hệ số phóng đại bằng 32/16 = 2, giúp cân bằng hoàn hảo giữa việc học các thuật ngữ pháp lý mới và việc giữ nguyên văn phong tự nhiên có sẵn của Qwen gốc.")
    add_explanation_bullet("lora_dropout=0.05", "Tỷ lệ dropout của adapter. Ý nghĩa đối với mô hình: Tắt ngẫu nhiên 5% các kết nối trong ma trận adapter trong mỗi bước huấn luyện. Điều này ép mô hình không được phụ thuộc vào bất kỳ đường truyền trọng số cụ thể nào, tăng cường khả năng khái quát hóa và chống hiện tượng quá khớp (overfitting).")
    add_explanation_bullet('bias="none"', "Đóng băng hoàn toàn các trọng số bias, giúp giảm thiểu tối đa số lượng tham số cần tối ưu hóa và tiết kiệm bộ nhớ.")
    add_explanation_bullet("target_modules=[...]", "Các lớp chiếu đích. Ý nghĩa đối với mô hình: Đắp các adapter LoRA lên toàn bộ các lớp chiếu của cơ chế Self-Attention (q, k, v, o_proj) và các lớp biến đổi phi tuyến tính MLP (gate, up, down_proj). Việc này cho phép mô hình điều chỉnh linh hoạt cả khả năng tập trung ngữ cảnh của các token lẫn khả năng biểu diễn tri thức logic chuyên ngành pháp lý.")

    add_custom_heading("2.4 Khởi tạo mô hình và Chuẩn bị huấn luyện k-bit", level=2)
    p = doc.add_paragraph("Nạp mô hình gốc 4-bit và đóng gói thích ứng LoRA:")
    
    code_init = """
model = AutoModelForCausalLM.from_pretrained(
    MODEL_NAME,
    quantization_config=bnb_config,
    device_map="auto",
    trust_remote_code=True,
    torch_dtype=torch.float16,
)
model = prepare_model_for_kbit_training(model, use_gradient_checkpointing=True)
model = get_peft_model(model, lora_config)
model.config.use_cache = False
    """
    add_code_block(code_init)

    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("AutoModelForCausalLM.from_pretrained(...)", "Tải mô hình Qwen từ HuggingFace Hub với các tham số cấu hình tùy chỉnh để bắt đầu nạp kiến trúc mạng neural.")
    add_explanation_bullet("quantization_config=bnb_config", "Nạp cấu hình nén. Ý nghĩa đối với mô hình: Áp dụng trực tiếp bộ lọc nén lượng tử hóa 4-bit lên các ma trận trọng số ngay khi tải mô hình vào VRAM hệ thống.")
    add_explanation_bullet('device_map="auto"', "Tự động phân bổ các lớp (layers) của mô hình lên GPU có sẵn để tối ưu hóa tốc độ tính toán song song.")
    add_explanation_bullet("trust_remote_code=True", "Cho phép chạy mã nguồn cấu trúc tùy chỉnh của dòng Qwen, giúp mô hình tự động tối ưu các toán tử attention đặc thù (như FlashAttention) trên GPU.")
    add_explanation_bullet("torch_dtype=torch.float16", "Đặt kiểu dữ liệu tải các layer hỗ trợ là Float16 để tương thích với GPU.")
    add_explanation_bullet("prepare_model_for_kbit_training(...)", "Chuẩn bị mô hình cho huấn luyện k-bit. Ý nghĩa đối với mô hình: Ép kiểu dữ liệu của một số layer nhạy cảm như LayerNorm (chuẩn hóa lớp) và lm_head (lớp dự đoán token đầu ra) sang Float32 để đảm bảo độ chính xác tính toán gradient không bị suy giảm trong quá trình tối ưu hóa 4-bit.")
    add_explanation_bullet("use_gradient_checkpointing=True", "Gradient Checkpointing. Ý nghĩa đối với mô hình: Trong quá trình lan truyền xuôi, mô hình sẽ không lưu giữ các giá trị kích hoạt (activations) của các lớp trung gian vào bộ nhớ GPU nữa. Thay vào đó, mô hình sẽ tính toán lại các activations này khi thực hiện lan truyền ngược. Kỹ thuật này trực tiếp giảm tới 60-70% bộ nhớ động cần thiết cho việc lưu trữ activations.")
    add_explanation_bullet("get_peft_model(model, lora_config)", "Khởi tạo mô hình PEFT. Ý nghĩa đối với mô hình: Đóng băng (freeze) 100% trọng số của mô hình Qwen gốc và chỉ mở quyền cập nhật (requires_grad=True) đối với các tham số thuộc ma trận LoRA. Từ thời điểm này, mô hình chỉ học thông qua ~29.9 triệu tham số LoRA, giảm thiểu chi phí tính toán gradient cực lớn.")
    add_explanation_bullet("model.config.use_cache = False", "Tắt tính năng Key-Value Cache. Ý nghĩa đối với mô hình: KV Cache lưu lại các giá trị Key-Value của các token trước đó để tăng tốc sinh từ lúc suy luận, nhưng tính năng này không tương thích với Gradient Checkpointing trong quá trình huấn luyện và sẽ gây lỗi tính toán lan truyền ngược nếu không tắt.")

    add_custom_heading("2.5 Cấu hình SFTConfig và Khởi tạo SFTTrainer", level=2)
    p = doc.add_paragraph("Thiết lập cấu hình huấn luyện chi tiết và khởi chạy Trainer:")
    
    code_sft = """
sft_config = SFTConfig(
    output_dir=CKPT_DIR,
    num_train_epochs=3,
    per_device_train_batch_size=4,
    per_device_eval_batch_size=4,
    gradient_accumulation_steps=4,
    learning_rate=2e-4,
    weight_decay=0.01,
    optim="paged_adamw_8bit",
    lr_scheduler_type="cosine",
    warmup_steps=50,
    max_grad_norm=0.3,
    gradient_checkpointing=True,
    fp16=USE_FP16,
    max_seq_length=512,
    dataset_text_field="text",
)

trainer = SFTTrainer(
    model=model,
    args=sft_config,
    train_dataset=train_ds,
    eval_dataset=val_ds,
    processing_class=tokenizer,
)
    """
    add_code_block(code_sft)

    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("output_dir=CKPT_DIR", "Chỉ định thư mục đích để lưu trữ các checkpoints trọng số thích ứng LoRA (.safetensors) khi mô hình đạt điểm tối ưu.")
    add_explanation_bullet("num_train_epochs=3", "Số Epoch huấn luyện. Ý nghĩa đối với mô hình: Ép mô hình học lặp đi lặp lại 3 lần trên toàn bộ tập dữ liệu huấn luyện để tối đa hóa khả năng nhận diện các đặc trưng pháp lý mà không lo bị quá khớp (overfitting).")
    add_explanation_bullet("per_device_train_batch_size=4", "Kích thước batch huấn luyện thực tế cấp cho mô hình là 4 mẫu văn bản trong một bước tính toán gradient.")
    add_explanation_bullet("gradient_accumulation_steps=4", "Tích lũy gradient. Ý nghĩa đối với mô hình: Mô hình sẽ thực hiện tính toán lan truyền ngược trên 4 batch nhỏ liên tiếp để cộng dồn gradient rồi mới thực hiện cập nhật trọng số một lần. Điều này giúp mô hình giả lập một kích thước batch lớn hơn (4 * 4 = 16 mẫu) mà không làm tăng dung lượng VRAM tiêu thụ trong một thời điểm, giúp bước đi tối ưu hóa ổn định hơn.")
    add_explanation_bullet("learning_rate=2e-4", "Tốc độ học tối đa. Ý nghĩa đối với mô hình: Kích thước bước đi tối ưu hóa (step size) của gradient trên bề mặt hàm mất mát. Giá trị 2e-4 là mức chuẩn tối ưu giúp các ma trận LoRA hội tụ nhanh và mượt mà nhất.")
    add_explanation_bullet("weight_decay=0.01", "Suy giảm trọng số. Ý nghĩa đối với mô hình: Áp dụng phạt L2 Regularization lên độ lớn của các trọng số LoRA. Giúp mô hình giữ các trọng số ở biên độ nhỏ, tránh việc mô hình học vẹt (quá khắt khe) trên dữ liệu CUAD.")
    add_explanation_bullet('optim="paged_adamw_8bit"', "Sử dụng Optimizer Paged AdamW 8-bit. Ý nghĩa đối với mô hình: Optimizer AdamW cần lưu trữ các trạng thái gradient (moments) làm tăng gấp đôi bộ nhớ mô hình. Lượng tử hóa 8-bit giúp nén trạng thái này, đồng thời cơ chế phân trang (paging) tự động đẩy các trạng thái thừa từ VRAM sang RAM hệ thống khi bị tràn, ngăn chặn lỗi Out-Of-Memory (OOM) của mô hình.")
    add_explanation_bullet('lr_scheduler_type="cosine"', "Bộ điều chỉnh tốc độ học Cosine. Ý nghĩa đối với mô hình: Giảm dần tốc độ học theo đồ thị hàm số Cosine từ mức cực đại 2e-4 về gần 0 ở giai đoạn cuối huấn luyện. Giúp mô hình có những bước đi lớn và nhanh ở giai đoạn đầu, sau đó đi những bước rất nhỏ và mịn ở cuối để tìm ra điểm cực trị toàn cục tốt nhất của hàm mất mát.")
    add_explanation_bullet("warmup_steps=50", "Số bước khởi động. Ý nghĩa đối với mô hình: Trong 50 bước huấn luyện đầu tiên, tốc độ học sẽ tăng dần từ 0 lên 2e-4. Giúp mô hình tránh được các cú sốc gradient (gradient shock) ở các bước đầu tiên khi adapter LoRA chưa có tri thức gì, giữ cho quá trình huấn luyện ổn định.")
    add_explanation_bullet("max_grad_norm=0.3", "Cắt tỉa gradient tối đa là 0.3. Ý nghĩa đối với mô hình: Nếu độ dài vector gradient vượt quá 0.3, nó sẽ bị scale nhỏ lại. Kỹ thuật này giúp mô hình chống lại hiện tượng bùng nổ gradient (gradient explosion) khi gặp các mẫu dữ liệu bất thường.")
    add_explanation_bullet("fp16=USE_FP16", "Sử dụng Mixed Precision FP16 giúp tăng tốc độ tính toán cho mô hình trên các GPU hỗ trợ Tensor Cores.")
    add_explanation_bullet("max_seq_length=512", "Độ dài chuỗi tối đa. Ý nghĩa đối với mô hình: Giới hạn độ dài ngữ cảnh xử lý tối đa là 512 tokens để khống chế kích thước của ma trận attention, ngăn bộ nhớ tăng theo hàm mũ.")
    add_explanation_bullet('dataset_text_field="text"', "Chỉ thị trường dữ liệu chứa chuỗi văn bản ChatML để nạp vào tokenizer chuyển đổi thành các chuỗi ID tokens tương ứng.")
    add_explanation_bullet("trainer = SFTTrainer(...)", "Khởi tạo đối tượng quản lý quy trình huấn luyện Supervised Fine-Tuning với mô hình, dữ liệu và các cấu hình tối ưu đã thiết lập.")

    doc.add_page_break()

    # --- SECTION 3 ---
    add_custom_heading("3. CHI TIẾT KỸ THUẬT MÔ HÌNH TÌM KIẾM NỘI DUNG HỢP ĐỒNG (MULTILINGUAL-E5-BASE)", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Mô hình thứ hai là tìm kiếm thông tin nội dung hợp đồng đa ngôn ngữ (Anh - Việt chéo) sử dụng mô hình mã hóa E5 và lập chỉ mục vào Vector Database ChromaDB. Dưới đây là các câu lệnh chi tiết và giải thích cụ thể.")

    add_custom_heading("3.1 Chi tiết ánh xạ câu hỏi nghiệp vụ (Query Mapping)", level=2)
    p = doc.add_paragraph("Để ánh xạ các nhãn thuộc tính thô từ bộ dữ liệu CUAD sang câu hỏi đầy đủ bằng tiếng Anh và tiếng Việt, cấu trúc `query_mapping` được xây dựng như sau:")
    
    code_mapping = """
query_mapping = {
    "Document Name":                     ("What is the name of this contract document?",              "Tên tài liệu hợp đồng này là gì?"),
    "Parties":                           ("Who are the parties involved in this contract?",            "Các bên tham gia hợp đồng này là ai?"),
    "Agreement Date":                    ("What is the agreement date of this contract?",              "Ngày ký kết hợp đồng là khi nào?"),
    "Effective Date":                    ("What is the effective date of this contract?",              "Ngày có hiệu lực của hợp đồng là khi nào?"),
    "Expiration Date":                   ("When does this contract expire?",                           "Hợp đồng này hết hạn vào ngày nào?"),
    "Renewal Term":                      ("What is the renewal term of this contract?",                "Điều khoản gia hạn hợp đồng là gì?"),
    "Notice Period To Terminate Renewal":("What is the notice period required to terminate renewal?", "Thời gian thông báo để chấm dứt gia hạn là bao lâu?"),
    "Governing Law":                     ("What is the governing law of this contract?",               "Luật điều chỉnh hợp đồng này là luật nào?"),
    "Most Favored Nation":               ("Does this contract contain a most favored nation clause?",  "Hợp đồng có điều khoản tối huệ quốc không?"),
    "Non-Compete":                       ("Does this contract contain a non-compete clause?",          "Hợp đồng có điều khoản không cạnh tranh không?"),
    "Exclusivity":                       ("Is there an exclusivity clause in this contract?",          "Hợp đồng có điều khoản độc quyền không?"),
    "No-Solicit Of Customers":           ("Does this contract restrict solicitation of customers?",    "Hợp đồng có cấm tiếp cận khách hàng không?"),
    "Termination For Convenience":       ("Can this contract be terminated for convenience?",          "Hợp đồng có thể chấm dứt theo ý muốn không?"),
    "Ip Ownership Assignment":           ("Who owns the intellectual property in this contract?",      "Ai sở hữu tài sản trí tuệ trong hợp đồng này?"),
    "Audit Rights":                      ("Does this contract include audit rights?",                  "Hợp đồng có điều khoản quyền kiểm toán không?"),
    "Cap On Liability":                  ("What is the cap on liability in this contract?",            "Giới hạn trách nhiệm pháp lý trong hợp đồng là bao nhiêu?"),
    "Warranty Duration":                 ("What is the warranty duration in this contract?",           "Thời hạn bảo hành trong hợp đồng là bao lâu?"),
    "Insurance":                         ("What are the insurance requirements in this contract?",     "Yêu cầu bảo hiểm trong hợp đồng là gì?"),
}
    """
    add_code_block(code_mapping)
    p = doc.add_paragraph("Ý nghĩa đối với mô hình: Thay vì bắt mô hình học biểu diễn các từ khóa nhãn thô cực kỳ ngắn ngủi (như 'Governing Law') vốn rất xa lạ với ngôn ngữ tự nhiên, chúng ta định nghĩa các câu hỏi hoàn chỉnh bằng cả hai ngôn ngữ. Ý nghĩa đối với mô hình: Cung cấp các cấu trúc cú pháp phong phú và tự nhiên nhất, giúp mô hình học cách ánh xạ vector của các câu hỏi thực tế của người dùng vào cùng khu vực không gian vector với điều khoản hợp đồng tương ứng.")

    add_custom_heading("3.2 Tạo tập dữ liệu song ngữ chéo (Cross-lingual Dataset Expansion)", level=2)
    p = doc.add_paragraph("Mã nguồn thực hiện dịch máy các mẫu dữ liệu huấn luyện sang tiếng Việt để mô hình học cách map không gian vector giữa hai ngôn ngữ:")
    
    code_translate = """
translator = GoogleTranslator(source='en', target='vi')
vi_pairs = []

for i, row in sample_vi.iterrows():
    try:
        query_vi   = translator.translate(row["query"][:200])
        passage_vi = translator.translate(row["passage"][:400])
        
        # Thêm 3 luồng dữ liệu tương tác chéo ngôn ngữ
        vi_pairs.append({"query": query_vi,     "passage": row["passage"]})
        vi_pairs.append({"query": query_vi,     "passage": passage_vi})
        vi_pairs.append({"query": row["query"], "passage": passage_vi})
    except:
        continue
    """
    add_code_block(code_translate)

    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("GoogleTranslator(source='en', target='vi')", "Khởi tạo API dịch máy để chuyển ngữ dữ liệu thô.")
    add_explanation_bullet('translator.translate(row["query"][:200])', "Dịch câu hỏi sang tiếng Việt, tạo nguồn truy vấn bản địa.")
    add_explanation_bullet('translator.translate(row["passage"][:400])', "Dịch đoạn văn điều khoản hợp đồng sang tiếng Việt.")
    add_explanation_bullet('vi_pairs.append({"query": query_vi, "passage": row["passage"]})', "Cặp dữ liệu chéo (Query tiếng Việt - Passage tiếng Anh). Ý nghĩa đối với mô hình: Đây là dòng dữ liệu cốt lõi thúc đẩy mô hình học cách ánh xạ các khái niệm ngữ nghĩa tiếng Việt của người dùng với văn bản hợp đồng thực tế bằng tiếng Anh, thiết lập sự tương thích chéo ngôn ngữ (Cross-lingual retrieval).")
    add_explanation_bullet('vi_pairs.append({"query": query_vi, "passage": passage_vi})', "Cặp dữ liệu bản địa (Query tiếng Việt - Passage tiếng Việt). Ý nghĩa đối với mô hình: Giúp mô hình duy trì và cải thiện năng lực tìm kiếm đơn ngữ thuần tiếng Việt (Monolingual Search).")
    add_explanation_bullet('vi_pairs.append({"query": row["query"], "passage": passage_vi})', "Cặp dữ liệu chéo ngược (Query tiếng Anh - Passage tiếng Việt). Ý nghĩa đối với mô hình: Đa dạng hóa không gian vector của mô hình, giúp không gian nhúng của tiếng Anh và tiếng Việt được căn chỉnh chồng khít lên nhau.")

    add_custom_heading("3.3 Huấn luyện SentenceTransformer với MultipleNegativesRankingLoss", level=2)
    p = doc.add_paragraph("Cấu hình huấn luyện đối sánh dùng mô hình E5 đa ngôn ngữ:")
    
    code_mnrl = """
train_examples = [
    InputExample(texts=[f"query: {r['query']}", f"passage: {r['passage']}"])
    for _, r in combined.iterrows()
]

model = SentenceTransformer("intfloat/multilingual-e5-base")
train_dataloader = DataLoader(train_examples, shuffle=True, batch_size=8)
train_loss = losses.MultipleNegativesRankingLoss(model=model)
warmup_steps = int(len(train_dataloader) * EPOCHS * 0.1)

model.fit(
    train_objectives=[(train_dataloader, train_loss)],
    epochs=EPOCHS,
    warmup_steps=warmup_steps,
    show_progress_bar=True,
    output_path="/kaggle/working/contract-e5-v2-final",
)
    """
    add_code_block(code_mnrl)

    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("InputExample(texts=[f\"query: {r['query']}\", f\"passage: {r['passage']}\"])", "Thêm tiền tố định hướng. Ý nghĩa đối với mô hình: Kiến trúc E5 yêu cầu bắt buộc thêm tiền tố 'query: ' và 'passage: ' để cơ chế Self-Attention của mô hình biết được đâu là câu hỏi ngắn cần tìm kiếm và đâu là đoạn văn dài chứa ngữ cảnh thông tin, giúp mô hình tối ưu hóa vector biểu diễn tốt nhất.")
    add_explanation_bullet('SentenceTransformer("intfloat/multilingual-e5-base")', "Tải mô hình nền tảng E5 đa ngôn ngữ có cấu trúc Encoder-only (nhận vào chuỗi văn bản và trả về vector đặc trưng 768 chiều).")
    add_explanation_bullet("DataLoader(train_examples, shuffle=True, batch_size=8)", "Bộ nạp dữ liệu với batch size là 8. Ý nghĩa đối với mô hình: Cung cấp đồng thời 8 cặp dữ liệu dương tính trong một bước huấn luyện.")
    add_explanation_bullet("losses.MultipleNegativesRankingLoss(model=model)", "Khởi tạo hàm mất mát MNRL. Ý nghĩa đối với mô hình: Đây là cốt lõi của huấn luyện đối sánh (contrastive learning). Với mỗi cặp (q_i, p_i) trong batch, hàm loss này coi p_i là nhãn đúng (positive sample) của q_i, đồng thời coi 7 đoạn văn bản còn lại (p_j với j != i) là các mẫu phủ định (in-batch negatives). Hàm loss sẽ tính toán điểm cosine similarity giữa q_i và tất cả p, sau đó tối đa hóa xác suất dự đoán đúng p_i bằng cách kéo vector q_i lại gần p_i nhất có thể, đồng thời đẩy vector q_i ra thật xa các vector p_j khác. Điều này giúp mô hình xây dựng các vùng ngữ nghĩa phân tách cực kỳ sắc bén trong không gian 768 chiều.")
    add_explanation_bullet("warmup_steps = int(...)", "Tính toán số bước khởi động để tăng dần tốc độ học ở giai đoạn đầu, bảo vệ mô hình khỏi hiện tượng nhiễu loạn gradient.")
    add_explanation_bullet("model.fit(...)", "Hàm thực thi fine-tune. Ý nghĩa đối với mô hình: Cho phép lan truyền ngược cập nhật toàn bộ các trọng số của các lớp Encoder E5 dựa trên lỗi tính toán từ hàm mất mát MNRL, định hình lại không gian vector phù hợp với lĩnh vực pháp lý.")

    add_custom_heading("3.4 Phân đoạn văn bản theo điều khoản (Section Chunking)", level=2)
    p = doc.add_paragraph("Hàm cắt văn bản thông minh để lập chỉ mục chính xác mà không làm mất ngữ cảnh:")
    
    code_chunk_detail = """
def chunk_by_section(text):
    # Nhận diện vị trí bắt đầu điều khoản bằng regex dạng số điều khoản
    matches = list(re.finditer(r'\\d+\\.\\d+\\s+\\w', text))
    
    chunks = []
    for i, match in enumerate(matches):
        start = match.start()
        end   = matches[i+1].start() if i+1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        
        # Chỉ giữ các đoạn hợp lệ có từ 5 từ trở lên
        if len(chunk.split()) >= 5:
            chunks.append(chunk)
            
    return chunks
    """
    add_code_block(code_chunk_detail)

    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("matches = list(re.finditer(r'\\d+\\.\\d+\\s+\\w', text))", "Sử dụng regex định vị tiêu đề các điều khoản con (X.Y).")
    add_explanation_bullet("start = match.start()", "Xác định điểm bắt đầu của đoạn điều khoản.")
    add_explanation_bullet("end = matches[i+1].start() ...", "Xác định điểm kết thúc tại vị trí bắt đầu của điều khoản kế tiếp.")
    add_explanation_bullet("chunk = text[start:end].strip()", "Cắt văn bản gốc thành các đoạn điều khoản độc lập. Ý nghĩa đối với mô hình: Việc cắt nhỏ theo đúng cấu trúc logic của văn bản pháp lý (thay vì cắt theo độ dài ký tự cố định) giúp mô hình E5 nhận diện trọn vẹn một ý nghĩa pháp lý thống nhất trong một đoạn nhúng duy nhất, tránh việc các câu chữ bị cắt nửa chừng làm mất ngữ cảnh bản chất.")
    add_explanation_bullet("if len(chunk.split()) >= 5:", "Lọc bỏ các đoạn quá ngắn hoặc ký tự rác để tránh nạp các vector nhiễu không mang giá trị ngữ nghĩa vào cơ sở dữ liệu.")

    add_custom_heading("3.5 Lập chỉ mục và truy vấn với Vector Database ChromaDB", level=2)
    p = doc.add_paragraph("Lập chỉ mục các vector 768 chiều vào ChromaDB và thực hiện tìm kiếm:")
    
    code_db = """
client = chromadb.Client()
collection = client.create_collection("contract_chunks")

embeddings = model_ft.encode(
    [f"passage: {c}" for c in chunks_sec],
    normalize_embeddings=True,
    show_progress_bar=True
).tolist()

collection.add(
    embeddings=embeddings,
    documents=chunks_sec,
    ids=[f"sec_{i}" for i in range(len(chunks_sec))],
    metadatas=[{"chunk_index": i} for i in range(len(chunks_sec))]
)
    """
    add_code_block(code_db)

    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("chromadb.Client()", "Khởi tạo cơ sở dữ liệu vector ChromaDB để quản lý và lập chỉ mục các điểm dữ liệu không gian.")
    add_explanation_bullet('client.create_collection("contract_chunks")', "Tạo bảng lưu trữ vector cho hợp đồng hiện tại.")
    add_explanation_bullet("model_ft.encode([f\"passage: {c}\" for c in chunks_sec], ...)", "Tính toán vector nhúng. Ý nghĩa đối với mô hình: Chuyển đổi các chuỗi ký tự tự nhiên thành các vector biểu diễn số học nằm trong không gian 768 chiều chứa tri thức ngữ nghĩa mà mô hình đã được học.")
    add_explanation_bullet("normalize_embeddings=True", "Chuẩn hóa vector L2. Ý nghĩa đối với mô hình: Thực hiện phép chia vector đầu ra cho độ dài Euclid của chính nó để ép độ dài vector (norm) về đúng 1.0. Trực tiếp đưa mọi vector biểu diễn văn bản lên một mặt cầu đơn vị. Ý nghĩa toán học là: Khi hai vector có độ dài bằng 1, khoảng cách Cosine Similarity giữa chúng chính xác bằng tích vô hướng của hai vector (Dot Product), giúp quá trình truy xuất thông tin cực kỳ nhanh và chuẩn xác.")
    add_explanation_bullet("collection.add(...)", "Lưu trữ vector và metadata vào ChromaDB phục vụ cho các phép toán so sánh khoảng cách không gian thời gian thực.")

    doc.add_page_break()

    # --- SECTION 4 ---
    add_custom_heading("4. CÁC CHỨC NĂNG PHỤ TRỢ (DEPLOYMENT, EVALUATION & GRADIO DEMO)", level=1)
    p = doc.add_paragraph("Để hoàn thiện vòng đời của mô hình AI, notebook cũng chứa các chức năng cực kỳ quan trọng dùng để tải mô hình lên HuggingFace Hub, đánh giá chất lượng độ tương đồng vector, và xây dựng giao diện demo tương tác nhanh Gradio.")

    add_custom_heading("4.1 Đẩy mô hình lên HuggingFace Hub bằng HfApi", level=2)
    p = doc.add_paragraph("Đoạn mã sau thực hiện việc tải trực tiếp thư mục checkpoints mô hình lên HuggingFace Hub thông qua API token:")
    
    code_upload = """
from huggingface_hub import HfApi

api = HfApi()
api.upload_folder(
    folder_path="/kaggle/working/contract-e5-v2-final",
    repo_id="d90nqm/contract-search-e5-v2",
    repo_type="model",
    token="hf_xxxxxxxxxxxxxxxxxxxxxxxxxxxxxx" # Thay thế bằng token ghi thực tế
)
    """
    add_code_block(code_upload)
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("from huggingface_hub import HfApi", "Import API của HuggingFace để quản lý vòng đời triển khai (deployment) của mô hình.")
    add_explanation_bullet("api.upload_folder(...)", "Tải toàn bộ các file lưu trạng thái trọng số (weights), cấu hình (config.json) và tokenizer lên đám mây HuggingFace Hub. Ý nghĩa đối với mô hình: Đóng gói và lưu trữ không đổi toàn bộ cấu trúc mạng neural đã tối ưu để có thể tải lại bất kỳ lúc nào phục vụ sản xuất (production) mà không cần huấn luyện lại.")

    add_custom_heading("4.2 Đánh giá và so sánh chất lượng mô hình v1 vs v2", level=2)
    p = doc.add_paragraph("Để chứng minh tính hiệu quả của mô hình v2 sau khi fine-tune song ngữ, đoạn mã so sánh trực quan khoảng cách vector giữa hai mô hình được lập trình:")
    
    code_compare = """
test_cases = [
    ("governing law",             "This Agreement shall be governed by the laws of Illinois."),
    ("luật điều chỉnh hợp đồng", "This Agreement shall be governed by the laws of Illinois."),
    ("điều khoản chấm dứt",      "Either party may terminate this Agreement upon 30 days notice."),
]

for query, passage in test_cases:
    # Đo điểm số tương đồng trên Model v1
    s1 = float(model_v1.encode([f"query: {query}"], normalize_embeddings=True) @
               model_v1.encode([f"passage: {passage}"], normalize_embeddings=True).T)
    # Đo điểm số tương đồng trên Model v2
    s2 = float(model_v2.encode([f"query: {query}"], normalize_embeddings=True) @
               model_v2.encode([f"passage: {passage}"], normalize_embeddings=True).T)
    diff = s2 - s1
    flag = "✅" if diff > 0 else "❌"
    print(f"{query:<35} {s1:>6.3f} {s2:>6.3f}  {flag} {diff:>+.3f}")
    """
    add_code_block(code_compare)
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("test_cases = [...]", "Tạo các cặp kịch bản đánh giá năng lực thực tế của mô hình trước và sau khi học.")
    add_explanation_bullet("model.encode(...)", "Sinh vector nhúng chuẩn hóa L2 cho các câu hỏi truy vấn và đoạn văn hợp đồng kiểm thử.")
    add_explanation_bullet("@ ... .T", "Phép nhân vô hướng (matrix multiplication). Ý nghĩa đối với mô hình: Tính toán trực tiếp tích vô hướng giữa hai vector đã chuẩn hóa L2. Phép toán này phản ánh chính xác điểm Cosine Similarity (điểm tương đồng ngữ nghĩa nằm trong đoạn từ -1 đến 1). Đối với mô hình, điểm số này càng lớn nghĩa là mô hình nhận định hai câu văn có mối quan hệ ngữ nghĩa càng khăng khít.")
    add_explanation_bullet("diff = s2 - s1", "Đo lường mức độ dịch chuyển khoảng cách vector ngữ nghĩa. Điểm số duy nhất (diff > 0) chứng minh rằng sau quá trình huấn luyện song ngữ, mô hình đã kéo hai câu khác ngôn ngữ (như câu tiếng Việt và điều khoản tiếng Anh tương ứng) lại gần nhau hơn trong không gian 768 chiều.")

    add_custom_heading("4.3 Giao diện Demo tương tác nhanh Gradio", level=2)
    p = doc.add_paragraph("Mã nguồn xây dựng ứng dụng web tương tác nhanh để thử nghiệm kết quả tìm kiếm ngữ nghĩa song ngữ:")
    
    code_gradio = """
def search_final(query: str, top_k: int = 3):
    if not query.strip():
        return "⚠️ Vui lòng nhập câu hỏi"
    q_emb   = model_ft.encode([f"query: {query}"], normalize_embeddings=True).tolist()
    results = collection.query(query_embeddings=q_emb, n_results=top_k)
    
    output  = f"### Kết quả cho: *\\"{query}\\"*\\n\\n📄 `{contract_name}`\\n\\n---\\n"
    for i, (doc, dist) in enumerate(zip(results["documents"][0], results["distances"][0])):
        score  = round(1 - dist, 3)
        bar    = "█" * int(score * 20)
        sec_no = results["metadatas"][0][i]["chunk_index"] + 1
        output += f"**Section #{sec_no}** · score: `{score}` {bar}\\n\\n> {doc[:500]}\\n\\n---\\n"
    return output

gr.Interface(
    fn=search_final,
    inputs=[
        gr.Textbox(label="Câu hỏi (tiếng Việt hoặc tiếng Anh)",
                   placeholder="VD: governing law / điều khoản chấm dứt / payment terms",
                   lines=2),
        gr.Slider(1, 5, value=3, step=1, label="Số đoạn trả về")
    ],
    outputs=gr.Markdown(label="Kết quả"),
    title="📄 Contract Content Search v2",
    description="Model v2 — fine-tuned với data EN + VI · Score > 0.7"
).launch(share=True)
    """
    add_code_block(code_gradio)
    p = doc.add_paragraph("Giải thích chi tiết từng dòng lệnh & Ý nghĩa trực tiếp đối với mô hình:")
    add_explanation_bullet("q_emb = model_ft.encode(...)", "Mô hình thực hiện mã hóa câu truy vấn thời gian thực của người dùng thành vector đại diện trong không gian nhúng.")
    add_explanation_bullet("collection.query(...)", "ChromaDB thực hiện thuật toán tìm kiếm vector lân cận gần nhất (k-nearest neighbors) để tìm ra top_k đoạn văn bản có khoảng cách không gian nhỏ nhất đối với vector câu hỏi.")
    add_explanation_bullet("score = round(1 - dist, 3)", "Quy đổi khoảng cách L2 (dist) về điểm số tương đồng ngữ nghĩa Cosine Similarity (score) giúp hiển thị trực quan độ chính xác cho người dùng kiểm thử.")
    add_explanation_bullet("gr.Interface(...).launch(share=True)", "Khởi chạy UI tương tác và mở cổng chia sẻ link ngrok, cho phép gửi các truy vấn trực tiếp từ trình duyệt đến mô hình đang chạy trong backend.")

    doc.add_page_break()

    # --- NEW SECTION 5: FUTURE ROADMAP & IMPROVEMENT DIRECTIONS ---
    add_custom_heading("5. HƯỚNG PHÁT TRIỂN MÔ HÌNH VÀ LỘ TRÌNH TƯƠNG LAI (FUTURE ROADMAP)", level=1)
    p = doc.add_paragraph()
    p.add_run("Để khắc phục triệt để các hạn chế của phương pháp fine-tune hiện tại và nâng cấp chất lượng đầu ra của cả hai mô hình, lộ trình phát triển được định hướng thông qua các trụ cột kỹ thuật cốt lõi sau:")
    
    add_custom_heading("5.1 Giải pháp cải tiến chất lượng và kiến trúc Model", level=2)
    add_detailed_bullet("Tăng cường dữ liệu đối kháng (Hard Negative Mining):", "Đối với Model Tìm kiếm Nội dung Hợp đồng, thay vì chỉ dựa vào các mẫu phủ định ngẫu nhiên trong batch, chúng tôi sẽ sử dụng mô hình hiện tại để truy xuất ra các đoạn văn bản có điểm số tương đồng cao nhất nhưng bị gán nhãn sai (false positives). Đưa trực tiếp các đoạn này vào làm mẫu Hard Negatives để mô hình học cách phân biệt tinh tế các điều khoản pháp lý gần giống nhau về mặt từ khóa nhưng khác biệt về mặt nghĩa luật.")
    add_detailed_bullet("Ứng dụng phương pháp huấn luyện căn chỉnh DPO / RLHF:", "Thu thập dữ liệu phản hồi (chấm điểm hoặc viết lại câu trả lời chưa chuẩn) trực tiếp từ các chuyên gia pháp chế và luật sư. Sau đó, tiến hành huấn luyện tinh chỉnh căn chỉnh DPO (Direct Preference Optimization) để phạt các câu trả lời sai lệch/ảo giác của mô hình trích xuất, đồng thời tăng điểm thưởng cho các câu trả lời chính xác, giúp chất lượng đầu ra đạt chuẩn văn phong luật Việt Nam.")
    add_detailed_bullet("Nâng cấp lên kiến trúc mô hình lớn hơn với cửa sổ ngữ cảnh mở rộng:", "Thay vì sử dụng kích thước 3B tham số, lộ trình sẽ nâng cấp lên các dòng mô hình lớn hơn như Qwen2.5-7B-Instruct hoặc Qwen2.5-14B-Instruct. Điều này giúp nâng cao đáng kể năng lực lập luận logic đa bước và mở rộng cửa sổ ngữ cảnh lên tới 128,000 tokens, cho phép nạp trực tiếp toàn bộ hợp đồng siêu dài mà không cần cắt nhỏ, giải quyết hoàn toàn lỗi lệch phân phối độ dài hội thoại.")
    add_detailed_bullet("Sử dụng kỹ thuật Full Fine-Tuning hoặc các thuật toán nâng cao (GaLore / LoRA+):", "Loại bỏ hoàn toàn lỗi lượng tử hóa VRAM bằng cách chạy huấn luyện đầy đủ tham số (Full Fine-tuning) trên hệ thống GPU cụm (như A100/H100), hoặc ứng dụng các phương pháp tối ưu hóa adapters tiên tiến như GaLore (Gradient Low-Rank Projection) và LoRA+ để cân bằng tốc độ học của các ma trận chiếu khác nhau, tăng khả năng tiếp thu tri thức sâu sắc của mô hình.")

    add_custom_heading("5.2 Giải pháp tối ưu hóa hạ tầng và tích hợp hệ thống", level=2)
    add_detailed_bullet("Xây dựng hệ thống RAG nâng cao (Advanced Retrieval-Augmented Generation):", "Kết hợp chặt chẽ mô hình tìm kiếm nội dung hợp đồng làm Retriever để định vị nhanh các điều khoản và đưa ngữ cảnh tìm được làm đầu vào (context) cho mô hình trích xuất thông tin hợp đồng phân tích chi tiết. Cơ chế này giúp triệt tiêu hiện tượng sinh ảo và vượt qua giới hạn độ dài ngữ cảnh của LLM.")
    add_detailed_bullet("Tối ưu hóa hiệu năng suy luận (Inference Optimization):", "Chuyển dịch runtime suy luận từ thư viện PyTorch tiêu chuẩn sang các framework chuyên biệt như vLLM, TensorRT-LLM hoặc Ollama. Việc tích hợp kỹ thuật PagedAttention giúp tăng tốc độ suy luận gấp 3-4 lần và giảm thiểu độ trễ sinh từ do lượng tử hóa 4-bit gây ra.")
    add_detailed_bullet("Tích hợp Tìm kiếm kết hợp (Hybrid Search):", "Kết hợp song song công nghệ tìm kiếm nội dung hợp đồng Dense Vector (E5) với tìm kiếm từ khóa Sparse Vector truyền thống (BM25 của Elasticsearch). Việc này giúp cải thiện độ chính xác đối với các truy vấn đặc thù chứa mã số điều khoản hoặc các danh từ riêng pháp lý.")

    doc.add_page_break()

    # --- SECTION 6 ---
    add_custom_heading("6. BẢNG SO SÁNH & TỔNG KẾT KỸ THUẬT", level=1)
    
    p = doc.add_paragraph("Dưới đây là bảng tổng hợp so sánh các kỹ thuật, mục đích sử dụng và các thông số then chốt giữa hai mô hình được tích hợp trong RiskDL:")

    # Add comparison table
    table = doc.add_table(rows=7, cols=3)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="CCCCCC")
    
    headers = ["Đặc tính / Mô hình", "1. Mô hình Trích xuất Thông tin Hợp đồng (LLM)", "2. Mô hình Tìm kiếm Nội dung Hợp đồng (Embedding)"]
    col_widths = [Inches(1.8), Inches(2.35), Inches(2.35)]
    
    hdr_cells = table.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "003366")
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        
    data = [
        ["Kiến trúc gốc", "Qwen2.5-3B-Instruct (Decoder-Only)", "Multilingual-E5-Base (Encoder-Only)"],
        ["Nhiệm vụ chính", "Trích xuất thông tin hợp đồng và tóm tắt rủi ro", "Tìm kiếm nội dung hợp đồng chéo Anh - Việt song song"],
        ["Định dạng Output", "JSON cấu trúc: {category, summary}", "Vector biểu diễn đặc trưng 768 chiều"],
        ["Kỹ thuật nén", "QLoRA 4-bit (NF4, Double Quantization)", "Không nén (chạy trực tiếp FP16/FP32)"],
        ["Hàm mất mát (Loss)", "Cross-Entropy Loss tiêu chuẩn", "MultipleNegativesRankingLoss (MNRL)"],
        ["Tham số huấn luyện", "Adapters LoRA r=16 (29.9M / 3.1B tham số)", "Cập nhật toàn bộ trọng số (Full fine-tune)"]
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            if col_idx == 0:
                set_cell_background(row_cells[col_idx], "F9F9F9")
                row_cells[col_idx].paragraphs[0].runs[0].font.bold = True
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
            
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    p_end = doc.add_paragraph()
    p_end.paragraph_format.space_before = Pt(24)
    run_end = p_end.add_run("Kết luận: ")
    run_end.bold = True
    p_end.add_run("Việc kết hợp cả hai mô hình này trong nền tảng RiskDL mang lại giải pháp hoàn chỉnh từ việc tìm kiếm nội dung hợp đồng thông minh cho đến việc đi sâu trích xuất thông tin hợp đồng chi tiết từng điều khoản, đáp ứng đầy đủ các tiêu chuẩn nghiệp vụ và kỹ thuật trong quản lý hợp đồng.")

    # Save document
    output_filename = "Bao_Cao_Ky_Thuat_Hai_Model_AI_Chi_Tiet.docx"
    doc.save(output_filename)
    print(f"[SUCCESS] Deeply detailed report saved to file: {output_filename}")

if __name__ == "__main__":
    create_report()
