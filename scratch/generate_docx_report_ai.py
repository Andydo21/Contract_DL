import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set cell background color in docx table."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # ------------------ COVER PAGE / TITLE ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO KỸ THUẬT PHÂN HỆ")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue 900
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_p2.add_run("AI SUMMARY & AI INFERENCE SERVICE")
    title_run2.font.name = 'Arial'
    title_run2.font.size = Pt(24)
    title_run2.font.bold = True
    title_run2.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Chi tiết cơ chế Dual Inference, cấu trúc API, mã nguồn Kaggle Notebook và giải pháp tối ưu VRAM GPU")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Separator Line
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep_run = p_sep.add_run("―" * 40)
    p_sep_run.font.color.rgb = RGBColor(226, 232, 240)
    
    doc.add_page_break()
    
    # ------------------ SECTION 1: KIẾN TRÚC HẠ TẦNG AI ------------------
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. Tổng Quan Kiến Trúc Hạ Tầng AI")
    r1.font.name = 'Arial'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(30, 58, 138)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Hệ thống RiskDL triển khai cơ chế Chạy kép (Dual Inference) linh hoạt để tương thích tối đa với cấu hình phần cứng doanh nghiệp. "
        "Dịch vụ AI phân rã thành hai module độc lập:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    # List of modes
    modes = [
        ("Chế độ cục bộ (Local GPU Mode):", " Hệ thống tự động phát hiện GPU và tải mô hình Qwen 3B kết hợp với PEFT LoRA adapter. Để chạy mượt mà trên các máy trạm có VRAM thấp (từ 4GB), mô hình được tải dưới dạng định dạng lượng hóa 4-bit (NF4)."),
        ("Chế độ đám mây (Kaggle Cloud Mode):", " Nếu cấu hình máy chủ cục bộ không có card đồ họa, hệ thống chuyển tiếp toàn bộ yêu cầu qua Ngrok tunnel kết nối trực tiếp đến môi trường Kaggle Notebook chạy GPU Tesla T4 kép (32GB VRAM). Chế độ này sử dụng phiên bản merged model ở dạng float16 có độ chính xác cao nhất.")
    ]
    
    for title, desc in modes:
        bp = doc.add_paragraph(style='List Bullet')
        bp_title = bp.add_run(title)
        bp_title.bold = True
        bp_title.font.color.rgb = RGBColor(79, 70, 229)
        bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)
        
    doc.add_page_break()
    
    # ------------------ SECTION 2: AI INFERENCE SERVICE ------------------
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. Dịch Vụ Phân Tích Rủi Ro (ai_service)")
    r2.font.name = 'Arial'
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(30, 58, 138)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Được triển khai bằng FastAPI, ai_service làm nhiệm vụ chính là thực thi mô hình LLM để đánh giá và chấm điểm rủi ro của từng điều khoản hợp đồng."
    )
    p.paragraph_format.space_after = Pt(10)
    
    funcs_ai_service = [
        ("analyze_contract", "POST /api/v1/analyze", "Điểm đầu cuối nhận vào danh sách các điều khoản pháp lý, thực thể bóc tách cơ bản và các luật rủi ro cần kiểm tra. Gọi hàm run_ai_analysis để xử lý."),
        ("run_ai_analysis", "clauses, entities, risk_rules", "Hàm điều phối luồng suy luận. Nếu KAGGLE_AI_URL được cấu hình, nó tự động chuyển tiếp request sang Kaggle. Nếu chạy local, nó sẽ lặp qua từng điều khoản và gọi _infer_local."),
        ("_infer_local", "messages", "Khởi chạy quá trình sinh văn bản từ mô hình LLM trên GPU cục bộ bằng tokenizer và cấu hình tham số suy luận cực kỳ chặt chẽ (temperature=0.1)."),
        ("_forward_to_kaggle", "payload", "Gửi yêu cầu phân tích thông qua requests POST đến URL ngrok của Kaggle với thời gian chờ tối đa (timeout) là 5 phút."),
        ("clean_and_parse_json", "text", "Hàm regex phân tách và làm sạch mã JSON được trả về từ LLM (vượt qua các ký tự markdown ```json) giúp ứng dụng Django phân tích dữ liệu một cách an toàn."),
        ("health_check", "GET /health", "Trả về trạng thái hoạt động của mô hình và báo cáo xem model đã load xong hay chưa.")
    ]
    
    table_ai = doc.add_table(rows=1, cols=3)
    table_ai.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ai.autofit = False
    
    col_widths = [Inches(1.8), Inches(1.8), Inches(2.9)]
    
    hdr_cells = table_ai.rows[0].cells
    headers = ["Tên hàm", "API Endpoint / Tham số", "Mô tả ý nghĩa chức năng"]
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "1E3A8A")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths[idx]
        
    for name, params, desc in funcs_ai_service:
        row_cells = table_ai.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[1].text = params
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[2].text = desc
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths[idx]
            
    doc.add_page_break()
    
    # ------------------ SECTION 3: AI SUMMARY SERVICE ------------------
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Dịch Vụ Tóm Tắt & Trích Xuất Thực Thể (ai_summary)")
    r3.font.name = 'Arial'
    r3.font.size = Pt(16)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(30, 58, 138)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Module ai_summary đóng vai trò là một Proxy trung gian gọn nhẹ chạy trên cổng 8004. Module này nhận các yêu cầu xử lý từ Django "
        "và phân phối sang máy chủ Kaggle chạy mô hình chuyên dụng cho nhiệm vụ tóm tắt văn bản và bóc tách thực thể pháp lý."
    )
    p.paragraph_format.space_after = Pt(10)
    
    funcs_ai_summary = [
        ("summarize_contract", "POST /api/v1/summarize", "Nhận vào danh sách các điều khoản, gọi dịch vụ AISummaryService.summarize để sinh Tóm tắt điều hành bằng tiếng Việt (150-250 từ)."),
        ("extract_entities", "POST /api/v1/extract_entities", "Nhận văn bản thô bất kỳ và bóc tách các thực thể COMPANY_NAME, TAX_CODE, CONTRACT_VALUE, DATE_EFFECTIVE, DATE_EXPIRE dưới dạng JSON."),
        ("health_check", "GET /health", "Đo lường tính sẵn sàng của Proxy và ping trực tiếp đầu cuối kiểm tra sức khỏe (/health) của Kaggle Notebook.")
    ]
    
    table_sum = doc.add_table(rows=1, cols=3)
    table_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sum.autofit = False
    
    hdr_cells = table_sum.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "4F46E5")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths[idx]
        
    for name, params, desc in funcs_ai_summary:
        row_cells = table_sum.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[1].text = params
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[2].text = desc
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths[idx]

    doc.add_page_break()
    
    # ------------------ SECTION 4: KAGGLE INFRASTRUCTURE ------------------
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Máy Chủ Kaggle Notebook & Phơi Lộ API")
    r4.font.name = 'Arial'
    r4.font.size = Pt(16)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(30, 58, 138)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Kaggle Notebook đóng vai trò máy chủ GPU hiệu năng cao chạy mô hình ngôn ngữ lớn Qwen 3B. "
        "Môi trường này được thiết lập thông qua hai tập tin cấu hình chính:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Subsection 4.1: kaggle_server.py
    p_srv = doc.add_paragraph()
    p_srv.add_run("4.1 Kịch bản Phân Tích Rủi Ro (ai_service/kaggle_server.py):").bold = True
    p_srv.paragraph_format.space_before = Pt(6)
    p_srv.paragraph_format.space_after = Pt(4)
    
    p_srv_desc = doc.add_paragraph(
        "Được dán trực tiếp vào các Cell của Kaggle. Tập lệnh thực hiện tải mô hình đã gộp (Merged Model) "
        "ở định dạng dấu phẩy động 16-bit (float16) nhờ tận dụng 2x GPU Tesla T4 có sẵn. "
        "Sau đó, khởi chạy một luồng Uvicorn ngầm chạy FastAPI trên port 8000 và sử dụng thư viện pyngrok "
        "để tạo một đường truyền an toàn (ngrok tunnel) sinh ra liên kết URL công khai để dán vào .env của hệ thống."
    )
    
    # Subsection 4.2: kaggle_notebook.py
    p_nb = doc.add_paragraph()
    p_nb.add_run("4.2 Kịch bản Tóm Tắt & Trích Xuất (ai_summary/kaggle_notebook.py):").bold = True
    p_nb.paragraph_format.space_before = Pt(6)
    p_nb.paragraph_format.space_after = Pt(4)
    
    p_nb_desc = doc.add_paragraph(
        "Chịu trách nhiệm chạy mô hình phamthanhfd/contract-analysis-qwen2.5-3b lượng hóa 4-bit. "
        "Kịch bản này cung cấp hai đầu cuối API chính cho việc sinh executive summary (tối đa 15 điều khoản đầu tiên để tránh tràn cửa sổ ngữ cảnh) "
        "và bóc tách thực thể pháp lý. Cổng dịch vụ 8001 được ngrok ánh xạ ra URL ngoài."
    )
    
    # Save document
    output_path = r"d:\Django_project\RiskDL\Bao_Cao_Ky_Thuat_AI_Summary_AI_Extract.docx"
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_report()
