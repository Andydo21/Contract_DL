import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set cell background color in docx table."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # ------------------ COVER PAGE / TITLE ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO KỸ THUẬT PHÂN HỆ")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue 900
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2 = title_p2.add_run("CONTRACT SERVICE & AI EXTRACT SERVICE")
    title_run2.font.name = 'Arial'
    title_run2.font.size = Pt(24)
    title_run2.font.bold = True
    title_run2.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Chi tiết Kiến trúc, Luồng Dữ liệu, Mô tả Hàm và Tác động tới Database Entities")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Separator Line
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep_run = p_sep.add_run("―" * 40)
    p_sep_run.font.color.rgb = RGBColor(226, 232, 240)
    
    doc.add_page_break()
    
    # ------------------ SECTION 1: TỔNG QUAN & LUỒNG DỰ ÁN ------------------
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. Tổng Quan & Luồng Dữ Liệu Dự Án")
    r1.font.name = 'Arial'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(30, 58, 138)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Hệ thống RiskDL được thiết kế dựa trên kiến trúc hướng dịch vụ (Service-Oriented Architecture). "
        "Nhiệm vụ cốt lõi là xử lý và lưu trữ thông tin hợp đồng an toàn, tự động phân tích cấu trúc điều khoản "
        "bằng trí tuệ nhân tạo (AI) và bảo chứng tính toàn vẹn thông qua công nghệ Blockchain."
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Bullet points for flow description
    p_flow_intro = doc.add_paragraph("Luồng xử lý tài liệu khi tải lên gồm các bước chính:")
    p_flow_intro.paragraph_format.space_after = Pt(6)
    
    flow_steps = [
        ("Mã hóa tài liệu (AES-256-GCM):", " File PDF/DOCX tải lên sẽ được mã hóa nhị phân trước khi lưu xuống ổ đĩa, đảm bảo bảo mật dữ liệu ở trạng thái nghỉ (data-at-rest)."),
        ("Trích xuất nội dung văn bản (OCR/Text Extraction):", " Giải mã file tạm thời trong bộ nhớ và chuyển qua DocumentService. Nếu là file PDF scan hoặc ảnh, hệ thống chạy PaddleOCR để chuyển văn bản hình ảnh thành văn bản tiếng Việt có dấu."),
        ("Lưu trữ ngữ cảnh phân trang (ContractContext):", " Lưu toàn bộ văn bản thô theo từng trang của hợp đồng để làm nguồn phân tích ngữ cảnh RAG."),
        ("Phân tách điều khoản (Clause Splitting):", " Gửi văn bản thô tới dịch vụ AI (Kaggle/LLM). Nếu AI mất kết nối hoặc quá tải, hệ thống tự động fallback sử dụng ClauseSplitter chạy Regex cục bộ để bóc tách điều mục."),
        ("Bảo chứng Blockchain (Anchoring):", " Tạo mã băm SHA-256 duy nhất của file hợp đồng và đăng ký (anchor) mã băm lên mạng Blockchain ảo. Lưu mã giao dịch (tx_hash) và số block để làm bằng chứng đối soát tính toàn vẹn sau này.")
    ]
    
    for title, desc in flow_steps:
        bp = doc.add_paragraph(style='List Bullet')
        bp_title = bp.add_run(title)
        bp_title.bold = True
        bp_title.font.color.rgb = RGBColor(79, 70, 229)
        bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)
        
    doc.add_page_break()
    
    # ------------------ SECTION 2: CONTRACT SERVICE ------------------
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. Chi Tiết Phân Hệ Contract Service (contracts/services.py)")
    r2.font.name = 'Arial'
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(30, 58, 138)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Lớp ContractService là đầu mối điều phối chính của ứng dụng Django, thực hiện quản trị vòng đời "
        "hợp đồng và giao tiếp trực tiếp với các dịch vụ bên ngoài."
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Table of functions
    funcs_cs = [
        ("list_all_contracts", "company=None", "Trả về danh sách hợp đồng, hỗ trợ lọc theo công ty quản lý."),
        ("get_contract_details", "contract_id, version_id=None", "Lấy toàn bộ chi tiết hợp đồng, danh sách các phiên bản, điều khoản và phân tích rủi ro tương ứng."),
        ("create_and_analyze_contract", "code, title, contract_type, start_date, end_date, contract_value, file_obj=None, raw_content=None, company=None", "Tạo mới hợp đồng ở trạng thái nháp DRAFT, mã hóa AES-256-GCM file tải lên và gọi bóc tách điều khoản ban đầu."),
        ("create_new_version", "contract_id, file_obj=None, raw_content=None, change_summary=''", "Tải lên và mã hóa phiên bản cập nhật tiếp theo (v2, v3...), cập nhật danh sách file liên kết."),
        ("analyze_contract", "contract_id, version_id=None", "Xóa các phân tích cũ, gửi dữ liệu văn bản thô của hợp đồng sang AI Service để chạy đánh giá rủi ro chuyên sâu."),
        ("manual_extract_contract", "contract_id, version_id=None", "Chạy trích xuất điều khoản bằng Regex cục bộ bỏ qua cuộc gọi AI."),
        ("extract_and_save_clauses_via_processor", "version, force_rule_based=False", "Gọi ClauseExtractService để bóc tách văn bản của phiên bản hợp đồng."),
        ("generate_blockchain_proof", "contract_id, version_id=None", "Tạo mã băm và đăng ký Proof ID trên Blockchain Service."),
        ("anchor_blockchain_proof", "contract_id, proof_id, network_id", "Xác nhận neo bằng chứng lên Blockchain và lưu tx_hash, block_number vào DB."),
        ("verify_blockchain_proof", "contract_id, version_id=None", "Đối soát file hiện tại bằng cách tính lại mã băm và đối chiếu với bản ghi gốc trên Blockchain.")
    ]
    
    table_cs = doc.add_table(rows=1, cols=3)
    table_cs.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cs.autofit = False
    
    # Widths
    col_widths = [Inches(1.8), Inches(1.8), Inches(2.9)]
    
    # Header Row
    hdr_cells = table_cs.rows[0].cells
    headers = ["Tên hàm", "Tham số truyền vào", "Mô tả ý nghĩa chức năng"]
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "1E3A8A") # Navy Blue
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths[idx]
        
    for name, params, desc in funcs_cs:
        row_cells = table_cs.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[1].text = params
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[2].text = desc
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths[idx]
            
    doc.add_page_break()
    
    # ------------------ SECTION 3: AI EXTRACT SERVICE ------------------
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Chi Tiết Phân Hệ AI Extract Service (ai_extract/services.py)")
    r3.font.name = 'Arial'
    r3.font.size = Pt(16)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(30, 58, 138)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "AI Extract Service cung cấp kết nối trung gian và xử lý ngôn ngữ tự nhiên thông qua hai lớp dịch vụ:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Subsection 3.1: ClauseExtractService
    sub31 = doc.add_paragraph()
    r31 = sub31.add_run("3.1 Lớp ClauseExtractService (Tách điều khoản)")
    r31.font.name = 'Arial'
    r31.font.size = Pt(13)
    r31.font.bold = True
    r31.font.color.rgb = RGBColor(79, 70, 229)
    
    p_31 = doc.add_paragraph(
        "Nhiệm vụ chính là gửi văn bản thô tới mô hình AI trên máy chủ phụ trợ để phân rã tài liệu thành các điều khoản riêng lẻ."
    )
    
    funcs_31 = [
        ("extract_version", "version, re_extract=False, force_rule_based=False", "Quy trình giải mã file, chạy DocumentService trích xuất văn bản thô lưu vào ContractContext, gửi yêu cầu tách điều khoản lên AI Service. Hỗ trợ tự động fallback sang Regex local nếu AI lỗi hoặc offline."),
        ("_get_raw_text", "version", "Khôi phục lại toàn bộ chuỗi văn bản thô của hợp đồng bằng cách ghép nối nội dung từ ContractContext (hoặc từ các điều khoản cũ nếu không có ngữ cảnh trang).")
    ]
    
    table_31 = doc.add_table(rows=1, cols=3)
    table_31.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_31.autofit = False
    
    hdr_cells = table_31.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "4F46E5") # Indigo
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths[idx]
        
    for name, params, desc in funcs_31:
        row_cells = table_31.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[1].text = params
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[2].text = desc
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths[idx]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Subsection 3.2: ExtractEntityService
    sub32 = doc.add_paragraph()
    r32 = sub32.add_run("3.2 Lớp ExtractEntityService (Nhận diện thực thể)")
    r32.font.name = 'Arial'
    r32.font.size = Pt(13)
    r32.font.bold = True
    r32.font.color.rgb = RGBColor(79, 70, 229)
    
    p_32 = doc.add_paragraph(
        "Chịu trách nhiệm trích xuất các thực thể quan trọng từ từng điều khoản đã bóc tách bằng cách gọi API AI `/api/v1/extract_entities`."
    )
    
    funcs_32 = [
        ("extract_version", "version, re_extract=False", "Duyệt qua từng Clause của phiên bản hợp đồng, gửi nội dung điều khoản sang AI Service và lưu các thực thể được chuẩn hóa vào DB thông qua Repository."),
        ("extract_from_text", "text, clause=None", "Hàm đa năng trích xuất thực thể từ chuỗi văn bản tự do bất kỳ. Nếu có truyền đối tượng Clause thì sẽ tự động lưu kết quả vào DB.")
    ]
    
    table_32 = doc.add_table(rows=1, cols=3)
    table_32.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_32.autofit = False
    
    hdr_cells = table_32.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "4F46E5") # Indigo
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths[idx]
        
    for name, params, desc in funcs_32:
        row_cells = table_32.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[1].text = params
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        
        row_cells[2].text = desc
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths[idx]
            
    doc.add_page_break()
    
    # ------------------ SECTION 4: TÁC ĐỘNG TỚI ENTITIES ------------------
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Tác Động Tới Các Thực Thể Cơ Sở Dữ Liệu (Database Entities)")
    r4.font.name = 'Arial'
    r4.font.size = Pt(16)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(30, 58, 138)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Các hoạt động của hai lớp dịch vụ làm thay đổi trạng thái của các thực thể trong cơ sở dữ liệu như sau:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    entities_impact = [
        ("Contract", "Tạo bản ghi mới khi đăng ký hợp đồng. Cập nhật trạng thái (DRAFT -> ANALYZING -> ACTIVE) trong quá trình xử lý."),
        ("ContractVersion", "Tạo mới khi hợp đồng được khởi tạo hoặc khi tải lên phiên bản cập nhật. Cập nhật các trường tx_hash, block_number từ Blockchain."),
        ("ContractFile", "Tạo liên kết tệp tin mã hóa với phiên bản hợp đồng khi tải lên file."),
        ("ContractContext", "Xóa toàn bộ bản ghi cũ của phiên bản đang phân tích và tạo mới các bản ghi phân trang lưu văn bản thô (raw_text)."),
        ("Clause", "Xóa toàn bộ các điều khoản cũ của phiên bản hợp đồng trước khi tiến hành ghi đè kết quả bóc tách mới (từ AI hoặc Regex fallback)."),
        ("ExtractedEntity", "Tạo mới các thực thể bóc tách tương ứng với từng Clause. Tự động bị xóa sạch (Cascade delete) khi Clause bị xóa."),
        ("RiskAnalysis", "Xóa các phân tích rủi ro cũ của phiên bản hợp đồng và tạo mới sau khi mô hình AI trả về kết quả đánh giá rủi ro chi tiết.")
    ]
    
    table_ent = doc.add_table(rows=1, cols=2)
    table_ent.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_ent.autofit = False
    
    col_widths_ent = [Inches(2.2), Inches(4.3)]
    
    hdr_cells = table_ent.rows[0].cells
    hdr_headers = ["Thực thể (Entity / Table)", "Hành động & Mức độ ảnh hưởng"]
    for idx, name in enumerate(hdr_headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "334155") # Slate 700
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        hdr_cells[idx].width = col_widths_ent[idx]
        
    for name, desc in entities_impact:
        row_cells = table_ent.add_row().cells
        row_cells[0].text = name
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 58, 138)
        
        row_cells[1].text = desc
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100)
            cell.width = col_widths_ent[idx]
            
    # Save document
    output_path = r"d:\Django_project\RiskDL\Bao_Cao_Ky_Thuat_Contract_AI_Services.docx"
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_report()
