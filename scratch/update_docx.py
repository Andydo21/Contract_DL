import docx

def update_docx(file_path):
    doc = docx.Document(file_path)
    found = False
    
    for i, p in enumerate(doc.paragraphs):
        if "Subquery seem has less execute time" in p.text:
            print(f"Found target paragraph at index {i}: '{p.text}'")
            
            # Clear the old paragraph text
            p.text = ""
            
            # Add structured response
            p.paragraph_format.space_before = docx.shared.Pt(6)
            p.paragraph_format.space_after = docx.shared.Pt(6)
            
            run_intro = p.add_run("Although both queries produce the same result, their execution plans reveal key differences in optimization and efficiency:\n\n")
            run_intro.font.name = 'Calibri'
            run_intro.font.size = docx.shared.Pt(11)
            
            # We will insert paragraphs after this one for structured points
            # Let's insert new paragraphs below this one
            
            p1 = doc.add_paragraph()
            p1.paragraph_format.left_indent = docx.shared.Inches(0.25)
            p1.paragraph_format.space_after = docx.shared.Pt(4)
            run = p1.add_run("1. Logical Operation (Semi-Join vs. Inner Join): ")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = docx.shared.Pt(11)
            run_text = p1.add_run(
                "The subquery using EXISTS is optimized into a Left Semi-Join. A semi-join is highly efficient "
                "because it stops processing the inner tables (SalesPerson and Employee) as soon as the first match "
                "is found (short-circuit evaluation). The JOIN query, conversely, executes an Inner Join that physically "
                "merges all three tables before filtering."
            )
            run_text.font.name = 'Calibri'
            run_text.font.size = docx.shared.Pt(11)
            
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = docx.shared.Inches(0.25)
            p2.paragraph_format.space_after = docx.shared.Pt(4)
            run = p2.add_run("2. Data Projection and Memory Overhead: ")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = docx.shared.Pt(11)
            run_text = p2.add_run(
                "Because EXISTS only checks for existence, the database engine does not need to carry over or "
                "project columns from SalesPerson or Employee into the execution pipeline. This reduces the memory grant "
                "and CPU cycles compared to the JOIN query, which passes data fields from all tables through the join operators."
            )
            run_text.font.name = 'Calibri'
            run_text.font.size = docx.shared.Pt(11)

            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = docx.shared.Inches(0.25)
            p3.paragraph_format.space_after = docx.shared.Pt(6)
            run = p3.add_run("3. Execution Time & Plan Cost: ")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = docx.shared.Pt(11)
            run_text = p3.add_run(
                "As observed, the subquery approach can have a lower execution cost and time. This is because SQL Server "
                "can filter SalesPerson first (identifying the few salesmen with Bonus > 5000) and then perform targeted "
                "Clustered Index Seeks on Person.Person, minimizing row scans."
            )
            run_text.font.name = 'Calibri'
            run_text.font.size = docx.shared.Pt(11)

            # Move these new paragraphs to be right after p
            # In python-docx, new paragraphs are added at the end, so we need to move them in the XML tree.
            p_element = p._p
            p_parent = p_element.getparent()
            
            # Get index of p_element
            p_idx = p_parent.index(p_element)
            
            # Insert p1, p2, p3 right after p
            p_parent.insert(p_idx + 1, p1._p)
            p_parent.insert(p_idx + 2, p2._p)
            p_parent.insert(p_idx + 3, p3._p)
            
            found = True
            break
            
    if found:
        doc.save(file_path)
        print("[SUCCESS] Updated document with professional execution plan comparison.")
    else:
        print("[ERROR] Target paragraph not found.")

try:
    update_docx(r"d:\Django_project\FSOFT\SQL_Assignment4_Đỗ_Đăng An.docx")
except Exception as e:
    print(f"Error: {e}")
